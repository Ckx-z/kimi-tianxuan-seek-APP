"""
bridge/_docx_styles.py
========================
v8 样式工具 - 抽出来供所有 docx 生成器复用
"""
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def setup_styles(doc):
    """统一设置文档样式"""
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')

    h1 = doc.styles['Heading 1']
    h1.font.name = '黑体'
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.line_spacing = 1.5
    rPr1 = h1.element.get_or_add_rPr()
    rFonts1 = rPr1.find(qn('w:rFonts'))
    if rFonts1 is None:
        rFonts1 = OxmlElement('w:rFonts')
        rPr1.append(rFonts1)
    rFonts1.set(qn('w:eastAsia'), '黑体')

    h2 = doc.styles['Heading 2']
    h2.font.name = '黑体'
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x2E, 0x5C, 0x8A)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.line_spacing = 1.5

    h3 = doc.styles['Heading 3']
    h3.font.name = '黑体'
    h3.font.size = Pt(11.5)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.line_spacing = 1.5


def set_cell_bg(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def set_para_bg(para, color_hex):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    pPr.append(shd)


def add_run(p, text, size=10.5, bold=False, italic=False, color=None, font='宋体'):
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    return run


def add_h1(doc, text):
    p = doc.add_heading('', level=1)
    add_run(p, text, size=16, bold=True, color=RGBColor(0x1F, 0x3A, 0x5F), font='黑体')
    return p


def add_h2(doc, text):
    p = doc.add_heading('', level=2)
    add_run(p, text, size=13, bold=True, color=RGBColor(0x2E, 0x5C, 0x8A), font='黑体')
    return p


def add_h3(doc, text):
    p = doc.add_heading('', level=3)
    add_run(p, text, size=11.5, bold=True, color=RGBColor(0xC0, 0x39, 0x2B), font='黑体')
    return p


def add_p(doc, text, indent=True, size=10.5, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    add_run(p, text, size=size, bold=bold, italic=italic, color=color)
    return p


def add_bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.first_line_indent = Cm(0.5)
    add_run(p, text, size=size)
    return p


def add_callout(doc, label, content, color_hex='F0F7FF'):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    set_para_bg(p, color_hex)
    add_run(p, f'【{label}】', size=10.5, bold=True, color=RGBColor(0xC0, 0x39, 0x2B))
    add_run(p, content, size=10.5)
    return p


def add_table(doc, header, rows, header_bg='1F3A5F', col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = w

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_bg(hdr_cells[i], header_bg)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        bg = 'F8F9FA' if r_idx % 2 == 0 else 'FFFFFF'
        for i, v in enumerate(row_data):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            add_run(p, str(v), size=9.5)
            set_cell_bg(cells[i], bg)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return table


def add_page_break(doc):
    doc.add_page_break()