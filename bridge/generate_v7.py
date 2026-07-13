"""
bridge/generate_v7.py
=====================
生成 v7 详细版 docx (v3.9 风格)
- D 系: TFPT + TFMB, 基于 R-101-1 (Tp+Pa, DCM/water 室温) + L-73ca3ac9b632 (TpTFMB 原位生长)
- A1 系: TAPT + A6, 基于 R-563-6 (Tp+Tta, 液-液界面) + L-0ca3cbf692f8 (Science Advances 智能膜)
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJ = r'C:\Users\ckx\Desktop\minimax'
OUT = os.path.join(PROJ, 'experiment', 'proposals',
                   'COF-TFPT-TAPT-2026-07-13-D-A1-v7.docx')


def set_cn_font(run, size=11, bold=False, name='宋体'):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_cn_font(run, size=16, bold=True)


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_cn_font(run, size=14, bold=True)


def add_p(doc, text, indent=True, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_cn_font(run, size=size)


def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_cn_font(run, size=size)


def add_table(doc, header, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # Header
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        set_cn_font(run, size=10, bold=True)
    # Rows
    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for i, v in enumerate(row_data):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            run = p.add_run(str(v))
            set_cn_font(run, size=10)


def main():
    doc = Document()
    # 全局字体 (默认宋体)
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(11)

    # ===== 封面 =====
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.line_spacing = 1.5
    run = title.add_run('含氟亚胺键 COF 膜合成方案 v7')
    set_cn_font(run, size=18, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.line_spacing = 1.5
    run = sub.add_run('D 系 (TFPT + TFMB) + A1 系 (TAPT + A6) 新方案')
    set_cn_font(run, size=14, bold=True)

    add_p(doc, '')
    add_p(doc, '方案编号: COF-TFPT-TAPT-2026-07-13-D-A1-v7', indent=False)
    add_p(doc, '生成日期: 2026-07-13', indent=False)
    add_p(doc, '迭代依据: 5 个 ABCDEF 失败 + 8 个 A1-D9 失败 (自反应法普遍失败)', indent=False)
    add_p(doc, '知识库检索: GraphRAG v2 (9066 节点 / 23920 边 / 954 文献 + multi-hop + 4-模态重排)', indent=False)
    add_p(doc, '参考文献基线: 侯盛怀等, Angew. Chem. Int. Ed. 2025, 64, e202421555', indent=False)

    # ===== 摘要 =====
    add_h1(doc, '摘要')
    add_p(doc, '本次方案基于 GraphRAG v2 检索到的关键文献, 针对 D 系 (TFPT + TFMB) 和 A1 系 (TAPT + A6) 重新设计。核心发现:')
    add_bullet(doc, '【D 系】GraphRAG 找到 L-73ca3ac9b632 (ACS AMI) — TpTFMB 2D COF 通过席夫碱反应原位生长在石英毛细管内壁, 证明 TFMB 能形成膜, 但需要特定基底。')
    add_bullet(doc, '【D 系】R-101-1 (Tp+Pa, water/DCM, 室温) 液-液界面法室温成膜是普遍可行的方法。')
    add_bullet(doc, '【A1 系】R-563-6 (Tp+Tta, DCM/water, 室温) 同样为液-液界面法, 三胺+三醛室温直接成膜。')
    add_bullet(doc, '【A1 系】L-0ca3cbf692f8 (Science Advances) Tp-TAPT 智能 COF 膜, 界面合成 + 溶剂响应。')
    add_bullet(doc, '【方法论】自反应法 (A1, A2, A5, D3) 普遍失败, 因缺界面模板 + 三嗪单体活性过强导致沉淀。')

    # ===== 1. 引言与背景 =====
    add_h1(doc, '一、研究背景与科学问题')
    add_p(doc, '本次方案针对实验 ABCDEF 与 A1-D9 共 14 个失败实验的迭代, 核心目标:')
    add_p(doc, '(1) D 系 (TFPT + TFMB): 5 个实验 (D/D3/D4/D7/D9) 全部失败 (Class E), 但 GraphRAG 找到 TpTFMB 体系 (类似结构) 能成膜。需要重新设计 TFPT+TFMB 的合成路径。')
    add_p(doc, '(2) A1 系 (TAPT + A6): 5 个实验 (A/A1/A2/A5/A8) 全部失败 (Class E), 自反应法+新加料序都不奏效。需要液-液界面法或表面引发法。')
    add_p(doc, '(3) 通用方法学问题: 含氟三嗪单体 (TFPT/TAPT) 反应性强, 标准溶剂热法易沉淀 → 需要界面法 + 动力学控制。')

    add_p(doc, '侯盛怀等 (Angew. Chem. Int. Ed. 2025, 64, e202421555) 的扩散/调制剂双介导固-液/气界面合成策略是基础方法, 但本次方案更进一步:')
    add_p(doc, '(a) 引入液-液界面 (DCM/water 两相) 实现真正的"二维聚合 + 室温成膜" (R-101-1, R-563-6, R-106-1) - 室温, 不需要油浴。')
    add_p(doc, '(b) 引入表面引发法 (SI-SBMAP, JACS 2024 L-c8cc6ff417dd) 预形成醛基化单分子层作为模板 - 这是"易脱落基底膜"思路的工业化方法。')

    # ===== 2. 历史失败总结 =====
    add_h1(doc, '二、历史失败总结 (14 条反馈)')
    add_table(doc,
              ['ID', '节点', '醛', '概率', 'Class', '现象摘要', '根因'],
              [
                  ['D', 'TFPT', 'TFMB', '0.596', 'E', '白色固体 0.2cm 粗糙', '操作错: 乙酸 0.4mL 12M + 苯胺量错'],
                  ['D3', 'TFPT', 'TFMB', '0.596', 'E', '自反应膜反复溶解 + 操作失误爆沸', '操作+单体'],
                  ['D4', 'TFPB', 'TFMB', '0.596', 'A', '#aborted 未冷却爆沸', '操作错'],
                  ['D7', 'TFPT', 'TFMB', '0.596', 'E', '自反应 12h 后加 TFMB 不连续粗糙', '单体'],
                  ['D9', 'TFPT', 'TFMB', '0.596', 'E', '放大 3x + 新序, 膜碳化', '操作+单体'],
                  ['A', 'TAPB', 'A6', '0.819', 'E', '黄褐色固体, 不连续粗糙', '操作错'],
                  ['A1', 'TAPT', 'A6', '0.819', 'E', '自反应膜 + 加醛后破碎 (一碰就碎)', '操作+单体'],
                  ['A2', 'TAPB', 'A6', '0.819', 'E', 'TAPB 自反应更差, 不连续深褐色', '操作+单体'],
                  ['A5', 'TAPT', 'A6', '0.819', 'E', '氯仿微调无效', '单体'],
                  ['A8', 'TAPT', 'A6', '0.819', 'E', '新加料序仍粗糙不连续', '操作+单体'],
              ])

    add_p(doc, '共同根因分析:')
    add_bullet(doc, '操作错 (6/10): 加料顺序反 + 乙酸过量 18M/12M + 苯胺量错 v3.9 → 通过严格按 v3.9 SOP 可解决。')
    add_bullet(doc, '单体+操作 (5/10): 自反应法 (A1, A2, A5, D3) 普遍失败 → 12h 自反应膜不稳定, 重启加热后反复溶解-再形成, 加醛后破碎。')
    add_bullet(doc, 'TFMB 反应性中等 (0.596 概率) 但实验 D 系列反复失败 → 暗示体系本身需重新设计。')

    # ===== 3. GraphRAG 检索结果 =====
    add_h1(doc, '三、GraphRAG v2 检索结果')

    add_h2(doc, '3.1 D 系 (TFPT + TFMB) 关键文献')
    add_p(doc, '【TOP 1 文献】L-73ca3ac9b632 (ACS Appl. Mater. Interfaces): TpTFMB 2D COF 体系, 由 2,2\'-双(三氟甲基)联苯胺(TFMB) 与 1,3,5-三甲酰基间苯三酚(Tp) 通过席夫碱反应原位生长在石英毛细管内壁。')
    add_p(doc, '   关键启示: TFMB 能形成膜, 但需要用 Tp (三醛但 β-酮烯胺互变) 而非 TFPT, 且需要在固体基底 (石英毛细管) 上原位生长。')
    add_p(doc, '【TOP 2 文献】L-c8cc6ff417dd (JACS): sp² 碳共轭 COFs 薄膜 (TFPT-TMT), SI-SBMAP 方法预形成醛基化单分子层。')
    add_p(doc, '   关键启示: 表面引发法 (SI) 是"易脱落基底膜"思路的工业化实现。')

    add_h2(doc, '3.2 D 系 (TFPT + TFMB) 关键反应')
    add_p(doc, '【TOP 1】R-101-1 (Tp + Pa): 1,3,5-triformylphloroglucinol + p-phenylenediamine, 溶剂 water/CH₂Cl₂ 两相, 室温 25°C, 产物 film。')
    add_p(doc, '   启示: Tp + 二胺在 DCM/water 界面成膜, 不需要 120°C 高温。')
    add_p(doc, '【关键发现】R-371 系列 (TFPT + 肼类), 水溶剂, 产物 powder → TFPT 与肼类在水相中沉淀, 需避开水。')

    add_h2(doc, '3.3 A1 系 (TAPT + A6) 关键文献')
    add_p(doc, '【TOP 1 文献】L-0ca3cbf692f8 (Science Advances): Tp-TAPT 亚胺键 COF 体系, 界面合成制备高结晶度 2D COF 薄膜 + 溶剂触发的层间滑移效应实现孔径自调控。')
    add_p(doc, '   关键启示: Tp + TAPT 能在界面成膜, 但需要"界面合成"工艺 (不是传统的溶剂热法)。')

    add_h2(doc, '3.4 A1 系 (TAPT + A6) 关键反应')
    add_p(doc, '【TOP 1】R-563-6 (Tp + Tta): Tp + TAPT, 溶剂 DCM (Tp 侧) / 水 (Tta 侧), 室温 25°C, 产物 film, 异相合成 (液-液界面)。')
    add_p(doc, '   启示: 三胺+三醛在 DCM/water 液-液界面室温直接成膜。')

    # ===== 4. D 系新方案 =====
    doc.add_page_break()
    add_h1(doc, '四、D 系新方案 (TFPT + TFMB)')

    add_h2(doc, '4.1 设计原理')
    add_p(doc, '基于 GraphRAG 检索结果, D 系失败的核心原因是 TFPT 三醛反应性过强 + 标准溶剂热法缺界面模板, 导致沉淀/破碎膜而非连续膜。本方案提出三种新策略:')
    add_p(doc, '【策略 1: 液-液界面法】借鉴 R-101-1 (Tp+Pa, water/CH₂Cl₂, 室温), 把 TFPT 溶于有机相 (DCM), TFMB 溶于水相 (或反向), 在界面室温聚合。')
    add_p(doc, '【策略 2: 表面引发法】借鉴 L-c8cc6ff417dd (JACS, SI-SBMAP), 先在玻璃壁上预形成醛基化单分子层作为模板, 再加 TFMB 反应。')
    add_p(doc, '【策略 3: 借鉴 TpTFMB 体系】参考 L-73ca3ac9b632 (ACS AMI), 用 Tp 替代 TFPT (Tp 有 β-酮烯胺互变, 反应性较弱但更可控), 仍配 TFMB。')

    add_h2(doc, '4.2 实验方案 (D-S1 ~ D-S3)')

    add_h2(doc, '【D-S1: TFPT + TFMB, 液-液界面法, 室温】')
    add_p(doc, '设计依据: 借鉴 R-101-1 (Tp+Pa, DCM/water, 室温 film)。DCM/water 界面提供二维限域, 室温聚合避免热扰动。')
    add_p(doc, '步骤:')
    add_p(doc, '(1) 35 mL Pyrex 管准备: 丙酮/乙醇/水超声各 10 min, 120°C 干燥。')
    add_p(doc, '(2) 水相: 称取 TFMB (28.8 mg, 0.09 mmol) 溶于去离子水 (3.0 mL), 超声 10 min 完全溶解 (TFMB 微溶于水, 可加 1 滴乙酸助溶, pH ~5)。')
    add_p(doc, '(3) 有机相: 称取 TFPT (11.8 mg, 0.03 mmol) 溶于 DCM (3.0 mL), 超声 5 min 澄清 (TFPT 在 DCM +)。')
    add_p(doc, '(4) 界面组装: 用滴管沿管壁缓慢将有机相加入水相上层, 形成清晰 DCM/water 界面 (上层 DCM, 下层水)。')
    add_p(doc, '(5) 沿界面加入苯胺 (8.2 μL, 0.09 mmol, TFPT:PhNH₂=1:3) 滴入有机相, 摇匀。')
    add_p(doc, '(6) 沿界面加入 6.0 M 乙酸水溶液 (0.20 mL) 滴入水相, 静置。')
    add_p(doc, '(7) 室温 (25°C) 静置反应 72 h, 避光, 避免振动。')
    add_p(doc, '(8) 观察: 期望在 DCM/water 界面形成连续薄膜。')
    add_p(doc, '(9) 后处理: 用金属刮刀小心从界面剥离膜, 用丙酮和 THF 各洗 3 次, 室温干燥。')

    add_p(doc, '预期结果: 类似 R-101-1, 在 DCM/water 界面形成连续薄膜, 颜色浅黄/橙色, 厚度 ~100-300 nm。')
    add_p(doc, '失败预测:')
    add_bullet(doc, '【预测 1: 界面无膜, 两相清亮】原因: 室温反应速率太慢, TFPT+TFMB 反应活性低。修正: 升温至 40°C, 或延长至 120 h。')
    add_bullet(doc, '【预测 2: 水相浑浊, 有沉淀】原因: TFPT 易扩散到水相并水解。修正: 用 1.0 M NaCl 水溶液代替纯水 (盐析效应抑制扩散)。')
    add_bullet(doc, '【预测 3: 膜形成但易碎】原因: 室温聚合反应不完全。修正: 加催化剂 (乙酸浓度从 6M 提至 12M), 或先 25°C 反应 24h 再 60°C 反应 48h。')

    add_h2(doc, '【D-S2: TFPT + TFMB, 表面引发法 (SI-SBMAP)】')
    add_p(doc, '设计依据: 借鉴 L-c8cc6ff417dd (JACS, sp²c-COF 薄膜 SI-SBMAP)。表面引发法先在玻璃壁上预形成醛基化单分子层, 再加第二种单体反应。')
    add_p(doc, '步骤:')
    add_p(doc, '(1) 玻璃管预处理: 35 mL Pyrex 管用 3-氨丙基三乙氧基硅烷 (APTES, 5 vol% 甲醇溶液) 浸泡 2 h, 120°C 干燥 30 min, 表面氨基化。')
    add_p(doc, '(2) 醛基化: 加入 TFPT (11.8 mg, 0.03 mmol) + DCM (3.0 mL), 25°C 反应 12 h, 玻璃壁上的 -NH₂ 与 TFPT 醛基反应形成醛基化单分子层 (SI-SBMAP 第一步)。')
    add_p(doc, '(3) 清洗: 倒掉 TFPT 溶液, 用 DCM 洗 3 次, 玻璃壁上保留醛基化单分子层。')
    add_p(doc, '(4) 加入 TFMB 溶液: TFMB (28.8 mg, 0.09 mmol) + 甲苯 (0.6 mL) + 氯仿 (0.4 mL) + 苯胺 (8.2 μL) + 6.0 M 乙酸 (0.20 mL)。')
    add_p(doc, '(5) 超声混合 10 min, 加入 Pyrex 管中。')
    add_p(doc, '(6) 120°C 油浴反应 72 h, 醛基化单分子层作为模板引导 TFMB 在壁上聚合。')
    add_p(doc, '(7) 后处理: 丙酮 + THF 洗涤, 室温干燥, 期望得到连续膜。')

    add_p(doc, '预期结果: 在 APTES-TFPT 单分子层上生长出连续 TFPT-TFMB COF 膜, 厚度 200-500 nm。')
    add_p(doc, '失败预测:')
    add_bullet(doc, '【预测 1: 膜不连续】原因: APTES 修饰不均。修正: 增加 APTES 处理时间至 6 h, 或用等离子清洗玻璃管后再 APTES。')
    add_bullet(doc, '【预测 2: 无膜】原因: TFPT 单分子层太稀, 模板失效。修正: 增加 TFPT 浓度至 0.06 mmol (2 倍), 醛基化时间延长至 24 h。')
    add_bullet(doc, '【预测 3: 膜太厚 (>1 μm)】原因: TFPT 单分子层作为聚合核心引发过多聚合。修正: 降低 TFMB 浓度至 0.045 mmol, 缩短反应时间至 48 h。')

    add_h2(doc, '【D-S3: Tp + TFMB, 借鉴 ACS AMI 文献 (L-73ca3ac9b632)】')
    add_p(doc, '设计依据: L-73ca3ac9b632 证明 Tp+TFMB 能在石英毛细管内壁原位生长成膜。本方案把 Tp 替代 TFPT, 在 Pyrex 管壁尝试。')
    add_p(doc, '步骤:')
    add_p(doc, '(1) 35 mL Pyrex 管准备: 标准清洗。')
    add_p(doc, '(2) 称取 Tp (1,3,5-triformylphloroglucinol, 12.6 mg, 0.06 mmol) 溶于 mesitylene/1,4-dioxane (1:1, 3.0 mL) 混合溶剂, 超声 10 min 完全溶解。')
    add_p(doc, '(3) 称取 TFMB (28.8 mg, 0.09 mmol) 加入上述溶液, 加苯胺 (8.2 μL)。')
    add_p(doc, '(4) 加 6.0 M 乙酸 (0.20 mL), 超声 5 min。')
    add_p(doc, '(5) 120°C 油浴反应 72 h, 期望在玻璃壁上原位生长 TpTFMB COF 膜。')
    add_p(doc, '(6) 后处理: 丙酮 + THF 洗涤, 室温干燥。')

    add_p(doc, '预期结果: 类似 L-73ca3ac9b632, 形成连续 TpTFMB 膜。')
    add_p(doc, '失败预测:')
    add_bullet(doc, '【预测 1: 粉末而非膜】原因: Tp 在均相中沉淀。修正: 借鉴文献加 0.5 vol% 水促进界面。')
    add_bullet(doc, '【预测 2: 膜颜色深 (深红/棕)】原因: Tp 易氧化。修正: 反应前鼓氮气 15 min, 全程氮气保护。')

    add_h2(doc, '4.3 D 系对照实验汇总表')
    add_table(doc,
              ['编号', '节点', '胺', '方法', '温度', '溶剂', '时间', '核心变量'],
              [
                  ['D-S1', 'TFPT', 'TFMB', '液-液界面', '室温 25°C', 'DCM/water', '72 h', '新策略: 界面法'],
                  ['D-S2', 'TFPT', 'TFMB', '表面引发 (SI)', '120°C', '甲苯/氯仿', '72 h', 'APTES + TFPT 单分子层'],
                  ['D-S3', 'Tp', 'TFMB', '借鉴 ACS AMI', '120°C', 'mesitylene/dioxane', '72 h', 'Tp 替代 TFPT'],
              ])

    # ===== 5. A1 系新方案 =====
    doc.add_page_break()
    add_h1(doc, '五、A1 系新方案 (TAPT + A6)')

    add_h2(doc, '5.1 设计原理')
    add_p(doc, '基于 GraphRAG 检索, A1 系失败核心是 TAPT 三胺反应性强 + A6 三联苯骨架大位阻 + 自反应法扰动。本方案提出两种新策略:')
    add_p(doc, '【策略 1: 液-液界面法 (LL)】借鉴 R-563-6 (Tp+Tta, DCM/water, 室温) 直接成膜的成功案例。')
    add_p(doc, '【策略 2: 表面引发法 (SI)】借鉴 L-c8cc6ff417dd (JACS, sp²c-COF) + Tp-TAPT 智能膜 (Science Advances)。')

    add_h2(doc, '5.2 实验方案 (A1-S1 ~ A1-S3)')

    add_h2(doc, '【A1-S1: TAPT + A6, 液-液界面法, 室温】')
    add_p(doc, '设计依据: 借鉴 R-563-6 (Tp+Tta, DCM/water, 室温 film)。把 A6 溶于有机相 (DCM), TAPT 溶于水相 (用乙酸助溶), 界面聚合。')
    add_p(doc, '步骤:')
    add_p(doc, '(1) 35 mL Pyrex 管准备: 标准清洗。')
    add_p(doc, '(2) 有机相: 称取 A6 (19.0 mg, 0.045 mmol) 溶于 DCM (3.0 mL), 超声 10 min 完全溶解。')
    add_p(doc, '(3) 水相: 称取 TAPT (10.6 mg, 0.03 mmol) 分散于去离子水 (3.0 mL) + 1 滴 6M 乙酸, 超声 15 min (TAPT 不溶于水但能分散, 加酸助悬)。')
    add_p(doc, '(4) 界面组装: 沿管壁缓慢将有机相加入水相上层, 形成清晰 DCM/water 界面 (DCM 密度大, 在下层)。')
    add_p(doc, '(5) 加苯胺 (8.2 μL, A6:PhNH₂=1:2) 至有机相。')
    add_p(doc, '(6) 加 6.0 M 乙酸 (0.20 mL) 至水相。')
    add_p(doc, '(7) 室温 (25°C) 避光静置 72 h。')
    add_p(doc, '(8) 后处理: 从界面剥离膜, 丙酮 + THF 洗涤。')

    add_p(doc, '预期结果: 在 DCM/water 界面形成浅黄色 TAPT-A6 COF 膜。')
    add_p(doc, '失败预测:')
    add_bullet(doc, '【预测 1: TAPT 在水相中絮凝沉淀】原因: TAPT 几乎不溶于水。修正: 改用 TAPT 在 DCM (上层), A6 在水相 (用 1% SDS 助悬)。')
    add_bullet(doc, '【预测 2: 室温反应 72 h 后无膜】原因: 反应速率太慢。修正: 升温至 40°C, 加 0.05 mL 6M 乙酸作为催化剂。')

    add_h2(doc, '【A1-S2: TAPT + A6, 表面引发法】')
    add_p(doc, '设计依据: 借鉴 L-c8cc6ff417dd (JACS, SI-SBMAP)。APTES 修饰玻璃壁, A6 醛基化单分子层作为模板, 加 TAPT 反应。')
    add_p(doc, '步骤:')
    add_p(doc, '(1) 玻璃管预处理: 5 vol% APTES 甲醇溶液浸泡 2 h, 120°C 干燥 30 min。')
    add_p(doc, '(2) 醛基化: A6 (19.0 mg, 0.045 mmol) + DCM (3.0 mL), 25°C 反应 12 h, 形成醛基化单分子层。')
    add_p(doc, '(3) 倒掉 A6 溶液, DCM 洗 3 次。')
    add_p(doc, '(4) 加 TAPT (10.6 mg, 0.03 mmol) + 甲苯 (0.3 mL) + 氯仿 (0.7 mL) + 苯胺 (8.2 μL) + 6.0 M 乙酸 (0.20 mL)。')
    add_p(doc, '(5) 超声 10 min, 120°C 油浴 72 h。')
    add_p(doc, '(6) 后处理: 丙酮 + THF 洗涤。')

    add_p(doc, '预期结果: 在 APTES-A6 单分子层上生长 TAPT-A6 COF 膜。')
    add_p(doc, '失败预测:')
    add_bullet(doc, '【预测 1: 膜粗糙不连续】原因: A6 三联苯骨架位阻大, 聚合速度慢。修正: 加 0.05 mL 6M 乙酸催化, 延长至 96 h。')
    add_bullet(doc, '【预测 2: 膜太薄】原因: 单分子层太稀。修正: A6 醛基化浓度增至 0.06 mmol。')

    add_h2(doc, '【A1-S3: Tp + TAPT + A6 (三组分, 借鉴 Science Advances)】')
    add_p(doc, '设计依据: L-0ca3cbf692f8 (Science Advances) 报道 Tp-TAPT 智能 COF 膜。本方案在 TAPT + A6 体系中加入少量 Tp (10% mol) 作为模板剂, 促进界面聚合。')
    add_p(doc, '步骤:')
    add_p(doc, '(1) 称取 TAPT (10.6 mg, 0.03 mmol) + A6 (19.0 mg, 0.045 mmol) + Tp (1.3 mg, 0.006 mmol, 10% mol) 混合。')
    add_p(doc, '(2) 加 mesitylene (0.3 mL) + 1,4-dioxane (0.7 mL) + 苯胺 (8.2 μL)。')
    add_p(doc, '(3) 超声 10 min, 加 6.0 M 乙酸 (0.20 mL)。')
    add_p(doc, '(4) 120°C 油浴 72 h。')
    add_p(doc, '(5) 后处理: 丙酮 + THF 洗涤。')

    add_p(doc, '预期结果: Tp 作为"晶种"促进 TAPT+A6 在界面有序聚合, 形成混合 COF 膜。')
    add_p(doc, '失败预测:')
    add_bullet(doc, '【预测 1: Tp 单独成膜不均匀】原因: Tp 量不足。修正: 增加 Tp 至 20% mol。')
    add_bullet(doc, '【预测 2: 三组分反应产物复杂】原因: 三个醛源混合。修正: 先 Tp+A6 反应 12 h, 再加 TAPT 反应 60 h (分步加料)。')

    add_h2(doc, '5.3 A1 系对照实验汇总表')
    add_table(doc,
              ['编号', '节点', '醛', '方法', '温度', '溶剂', '时间', '核心变量'],
              [
                  ['A1-S1', 'TAPT', 'A6', '液-液界面', '室温 25°C', 'DCM/water', '72 h', '新策略: 界面法'],
                  ['A1-S2', 'TAPT', 'A6', '表面引发 (SI)', '120°C', '甲苯/氯仿', '72 h', 'APTES + A6 单分子层'],
                  ['A1-S3', 'Tp+TAPT', 'A6', '三组分', '120°C', 'mesitylene/dioxane', '72 h', 'Tp 作为模板剂'],
              ])

    # ===== 6. 关键参数说明 =====
    doc.add_page_break()
    add_h1(doc, '六、关键参数说明')

    add_h2(doc, '6.1 液-液界面法 (LL) 的关键')
    add_bullet(doc, '【两相密度匹配】DCM (ρ=1.33) > 水 (ρ=1.00), DCM 自然在下层。若用 mesitylene (ρ=0.86), 则在水上层。')
    add_bullet(doc, '【界面清晰度】两相体积比 1:1, 加料时沿管壁缓慢注入避免湍流混合。')
    add_bullet(doc, '【室温反应】25°C 是关键, 高温 (>40°C) 会加剧分子扩散破坏界面。')
    add_bullet(doc, '【时间】72-120 h, 比高温法长但结晶度通常更高 (R-101-1 数据)。')

    add_h2(doc, '6.2 表面引发法 (SI) 的关键')
    add_bullet(doc, '【APTES 处理】5 vol% 甲醇溶液, 2 h 浸泡 + 120°C 干燥。APTES 在玻璃表面形成 -NH₂ 单分子层。')
    add_bullet(doc, '【醛基化时间】12-24 h, 不能太长 (会形成寡聚物而非单分子层)。')
    add_bullet(doc, '【清洗】倒掉醛单体溶液后必须用 DCM 洗 3 次, 去除物理吸附的醛单体 (否则会引发 bulk 聚合)。')

    add_h2(doc, '6.3 与 v3.9 标准方法的差异')
    add_p(doc, '本次方案突破 v3.9 的"扩散/调制剂双介导固-液/气界面"框架, 引入:')
    add_p(doc, '(a) 液-液界面: 真正的两相界面 (DCM/water), 不是油-气-固三相界面。')
    add_p(doc, '(b) 表面引发: 主动构建模板 (APTES + 醛基化), 不是被动等待膜在玻璃上生长。')
    add_p(doc, '(c) 室温聚合: 部分方案 25°C, 不需要 120°C 油浴。')

    # ===== 7. 失败排查 =====
    add_h1(doc, '七、失败排查与修正')

    add_h2(doc, '7.1 D 系通用排查')

    add_p(doc, '▍症状: 液-液界面无膜形成, 两相清亮')
    add_p(doc, '可能原因: (a) TFPT/TFMB 反应活性低, 室温聚合太慢; (b) 界面被湍流破坏; (c) 苯胺量不足。')
    add_p(doc, '修正方法: (a) 升温至 40°C; (b) 重新加料时更缓慢, 用注射泵 0.1 mL/min; (c) 苯胺加至 16.4 μL (TFPT:PhNH₂=1:6)。')

    add_p(doc, '▍症状: 水相浑浊, 有沉淀')
    add_p(doc, '可能原因: TFPT 扩散到水相并水解。')
    add_p(doc, '修正方法: 加 1.0 M NaCl 到水相 (盐析效应抑制扩散), 或改用饱和 NaHCO₃ 溶液 (碱性抑制醛基水解)。')

    add_p(doc, '▍症状: 表面引发法 (D-S2) 无膜')
    add_p(doc, '可能原因: (a) APTES 修饰失败; (b) A6 单分子层未形成; (c) 模板被清洗掉。')
    add_p(doc, '修正方法: (a) 用 XPS 验证 APTES-N 峰; (b) 增加 A6 醛基化时间至 24 h; (c) 减少 DCM 清洗次数到 1 次。')

    add_h2(doc, '7.2 A1 系通用排查')

    add_p(doc, '▍症状: TAPT 在水相絮凝')
    add_p(doc, '可能原因: TAPT 极性低 (LogP=2.9), 不溶于纯水。')
    add_p(doc, '修正方法: 改用反相界面 (TAPT 在 DCM 上层, A6 在水相下层), 或在水相加 1% SDS 助悬。')

    add_p(doc, '▍症状: 膜形成但易碎 (一碰就碎)')
    add_p(doc, '可能原因: A6 三联苯骨架大位阻, 聚合不完全。')
    add_p(doc, '修正方法: (a) 加 0.05 mL 6M 乙酸作为二次催化; (b) 升温至 60°C; (c) 延长反应至 96 h。')

    add_p(doc, '▍症状: 三组分 (A1-S3) 反应产物复杂')
    add_p(doc, '可能原因: Tp + A6 醛基竞争反应。')
    add_p(doc, '修正方法: 改用分步加料 - 先 Tp+A6 反应 12 h, 再加 TAPT 反应 60 h。')

    # ===== 8. 表征方案 =====
    add_h1(doc, '八、表征方案')

    add_table(doc,
              ['表征', '目的', '关键判定', '数据对比'],
              [
                  ['PXRD', '结晶度', 'fwhm < 0.16° 为高结晶', '与 AMCOF-1 (TpPa) 对比'],
                  ['SEM', '膜表面 + 截面', '连续覆盖, 厚度均匀', '膜厚 100-500 nm'],
                  ['AFM', '粗糙度 Ra', 'Ra < 5 nm 为均匀', '与文献对照'],
                  ['FT-IR', '亚胺键 C=N', '~1631 cm⁻¹ 峰', '醛基 1670 cm⁻¹ 应消失'],
                  ['BET', '比表面积', '> 800 m²/g', 'COF 标准'],
                  ['XPS', '元素组成 + APTES', 'N 1s (APTES) + C=N 峰', '验证表面引发'],
                  ['接触角', '疏水性', '含氟膜 > 100°', 'A6 体系应 > 110°'],
              ])

    # ===== 9. 参考文献 =====
    add_h1(doc, '九、参考文献')

    add_p(doc, '1. 侯盛怀等. "扩散/调制剂双介导固-液/气界面合成策略". Angew. Chem. Int. Ed. 2025, 64, e202421555.')
    add_p(doc, '   [本地路径: 实验/文章/侯老师实验/侯盛怀德国应化.pdf]')
    add_p(doc, '   [本地路径: 实验/文献阅读/侯老师/原文/侯盛怀德国应化.pdf]')

    add_p(doc, '2. L-73ca3ac9b632 (ACS Appl. Mater. Interfaces). TpTFMB 2D COF 体系, 原位生长在石英毛细管内壁.')
    add_p(doc, '   [GraphRAG 检索]')

    add_p(doc, '3. L-c8cc6ff417dd (J. Am. Chem. Soc.). SI-SBMAP 方法, sp²c-COF 薄膜 (TFPT-TMT).')

    add_p(doc, '4. L-0ca3cbf692f8 (Science Advances). Tp-TAPT 智能 COF 膜, 界面合成 + 溶剂响应.')

    add_p(doc, '5. R-101-1 (GraphRAG). Tp + Pa, water/CH₂Cl₂ 两相, 室温, film.')

    add_p(doc, '6. R-563-6 (GraphRAG). Tp + Tta, DCM/water 液-液界面, 室温, film, 异相合成.')

    add_p(doc, '7. R-161-1 (GraphRAG). TFTA + TAPT, mesitylene:1,4-dioxane (1:1), 120°C, film, 异相 (Fe3O4 磁核).')

    add_p(doc, '8. R-60-6 (GraphRAG). FPDA + TAPT, 1,4-dioxane:mesitylene (1:1), 120°C, film.')

    add_p(doc, '9. R-106-1 (GraphRAG). Tp + Pa-CF3, DCM/water, 室温, film (含氟二胺界面法).')

    add_p(doc, '10. R-484-1 (GraphRAG). DHTA + TAPT, 微波辅助, 硅片表面, film (表面引发先例).')

    add_p(doc, '11. R-62 系列 (GraphRAG). Tp + Boc 保护胺, dioxane/water, 120°C, film (Boc 保护策略).')

    add_p(doc, '12. R-419ccb6946af (GraphRAG). 含氟 COP 膜, 有机相/水相自由界面瞬时形成, 室温 (含氟界面膜通用方法).')

    # ===== 10. 附录: 与 v3.9 通用步骤的对照 =====
    add_h1(doc, '十、附录: 本方案与 v3.9 通用步骤的对照')

    add_table(doc,
              ['步骤', 'v3.9 通用步骤', 'v7 新方案变化'],
              [
                  ['1. Pyrex管准备', '丙酮/乙醇/水超声各10 min', '同 v3.9'],
                  ['2. 加料 (单相法)', '醛 + 苯胺 → 立即胺 → 最后酸', '改为: 两相分开加 (DCM + water)'],
                  ['3. 界面组装', 'N/A (单相混合)', '新增: 沿管壁缓慢注入另一相'],
                  ['4. 催化剂', '6.0 M 乙酸 0.20 mL', '同 v3.9, 但分两相加'],
                  ['5. 超声', '10-15 min', '同 v3.9, 但分相超声'],
                  ['6. 加热', '120°C 油浴 48-72 h', 'LL 法改为 25°C 室温 72-120 h'],
                  ['7. 监测', '观察膜形成', '同 v3.9, 重点观察界面'],
                  ['8. 收集', '刮刀剥离', 'LL 法改为从界面吸出'],
                  ['9. 洗涤', '丙酮 + THF 各 3 次', '同 v3.9'],
                  ['10. 干燥', '室温或 60°C 真空', '同 v3.9'],
              ])

    add_p(doc, '')
    add_p(doc, '--- 本方案结束 ---')
    add_p(doc, '生成工具: bridge/generate_v7.py (基于 GraphRAG v2 + 人工精修)')
    add_p(doc, '数据来源: tianxuan-seek/data/ (954 yaml + 6197 反应 + 1059 单体)')
    add_p(doc, 'GraphRAG v2 模块: nl2graph + router + multimodal + importance + community + reasoning')

    doc.save(OUT)
    print(f'✓ 已生成: {OUT}')
    print(f'  大小: {os.path.getsize(OUT)/1024:.1f} KB')


if __name__ == '__main__':
    main()