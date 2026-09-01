"""分组导出实验记录 Word 测试：POST /api/records/export-bundle。

覆盖：多组内容（组标题/记录正文与单条导出口径一致）、无记录组标注
「暂无实验记录」、空列表 400、收藏不存在 404、文件名与 content-type。
数据目录打到临时目录，不碰真实数据。
"""

from __future__ import annotations

import io

import pytest
from fastapi.testclient import TestClient

from api.main import app

client = TestClient(app)

_DOCX_MEDIA = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


@pytest.fixture
def isolated_data(tmp_path, monkeypatch):
    from favorites import store as fav_store
    from records import store as rec_store
    monkeypatch.setattr(fav_store, "FAVORITES_DIR", tmp_path / "favs")
    monkeypatch.setattr(rec_store, "RECORDS_DIR", tmp_path / "recs")
    monkeypatch.setattr(rec_store, "ATTACHMENTS_DIR", tmp_path / "atts")
    return fav_store, rec_store


def _create_favorite(ald: str, amine: str,
                     ald_name: str, amine_name: str,
                     score: float | None = 0.66) -> str:
    r = client.post("/api/favorites", json={
        "aldehyde_smiles": ald, "amine_smiles": amine})
    assert r.status_code == 201
    fid = r.json()["id"]
    from favorites import store as fav_store
    # 显式补名称，保证组标题断言确定性（内置库未必收录测试单体）
    fav_store.update_favorite(fid, aldehyde={
        "smiles": ald, "cas": "", "name": ald_name}, amine={
        "smiles": amine, "cas": "", "name": amine_name})
    if score is not None:
        fav_store.update_prediction_snapshot(fid, {
            "score": score, "std": 0.02, "ood": "none"})
    return fid


def _create_record(fid: str, experiment_no: str, process: str,
                   summary: str = "", mistakes: str = "") -> dict:
    r = client.post("/api/records", json={
        "favorite_id": fid,
        "experiment_no": experiment_no,
        "outcome": "film",
        "operator": "张三",
        "conditions": {"solvent_1": "甲苯"},
        "process_notes": process,
        "timeline": [{"time_label": "2026-01-01 09:00",
                      "description": f"{experiment_no} 投料"}],
        "self_summary": summary,
        "mistakes": mistakes,
    })
    assert r.status_code == 201
    return r.json()


def _docx_text(content: bytes) -> str:
    import docx
    doc = docx.Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                parts.append(c.text)
    for s in doc.sections:
        parts.extend(p.text for p in s.footer.paragraphs)
    return "\n".join(parts)


class TestExportBundle:
    def test_multi_group_content(self, isolated_data):
        fid1 = _create_favorite("O=Cc1ccccc1", "Nc1ccccc1", "苯甲醛", "苯胺")
        fid2 = _create_favorite("O=CC=O", "NCCN", "乙二醛", "乙二胺")
        rec1 = _create_record(fid1, "A5", "投料→陈化→干燥",
                              summary="总体顺利", mistakes="加料太快")
        rec2 = _create_record(fid2, "B1", "回流两小时")

        r = client.post("/api/records/export-bundle",
                        json={"favorite_ids": [fid1, fid2]})
        assert r.status_code == 200
        assert r.headers["content-type"].startswith(_DOCX_MEDIA)
        cd = r.headers["content-disposition"]
        assert "attachment" in cd
        assert "filename*=" in cd
        assert "%E5%AE%9E%E9%AA%8C%E8%AE%B0%E5%BD%95%E6%B1%87%E6%80%BB" in cd
        assert r.content[:2] == b"PK"

        text = _docx_text(r.content)
        # 封面
        for kw in ("实验记录汇总导出", "导出时间：", "软件版本：v"):
            assert kw in text, f"封面缺少: {kw}"
        # 组标题（醛+胺）
        assert "苯甲醛 + 苯胺" in text
        assert "乙二醛 + 乙二胺" in text
        # 组内记录内容与单条导出口径一致
        for kw in ("A5", rec1["record_id"], "投料→陈化→干燥",
                   "A5 投料", "总体顺利", "加料太快", "0.660",
                   "B1", rec2["record_id"], "回流两小时", "B1 投料",
                   "由 COF 科研助手导出"):
            assert kw in text, f"汇总文档缺少: {kw}"
        # 组序与请求顺序一致
        assert text.index("苯甲醛 + 苯胺") < text.index("乙二醛 + 乙二胺")

    def test_empty_group_marked(self, isolated_data):
        fid1 = _create_favorite("O=Cc1ccccc1", "Nc1ccccc1", "苯甲醛", "苯胺")
        fid2 = _create_favorite("O=CC=O", "NCCN", "乙二醛", "乙二胺")
        _create_record(fid1, "A5", "投料")
        r = client.post("/api/records/export-bundle",
                        json={"favorite_ids": [fid1, fid2]})
        assert r.status_code == 200
        text = _docx_text(r.content)
        assert "乙二醛 + 乙二胺" in text
        assert "暂无实验记录" in text

    def test_empty_ids_400(self, isolated_data):
        assert client.post("/api/records/export-bundle",
                           json={"favorite_ids": []}).status_code == 400
        assert client.post("/api/records/export-bundle",
                           json={"favorite_ids": ["  "]}).status_code == 400

    def test_missing_favorite_404(self, isolated_data):
        fid = _create_favorite("O=Cc1ccccc1", "Nc1ccccc1", "苯甲醛", "苯胺")
        r = client.post("/api/records/export-bundle",
                        json={"favorite_ids": [fid, "fav_20990101_999"]})
        assert r.status_code == 404

    def test_duplicate_ids_deduped(self, isolated_data):
        fid = _create_favorite("O=Cc1ccccc1", "Nc1ccccc1", "苯甲醛", "苯胺")
        _create_record(fid, "A5", "投料")
        r = client.post("/api/records/export-bundle",
                        json={"favorite_ids": [fid, fid]})
        assert r.status_code == 200
        text = _docx_text(r.content)
        # 组标题出现在「目录」与正文各 1 次；去重验证看记录正文只出现 1 次
        assert text.count("苯甲醛 + 苯胺") == 2
        assert text.count("实验记录 A5") == 1
