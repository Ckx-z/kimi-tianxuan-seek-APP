"""实验记录 Word 导出测试：GET /api/records/{id}/export。

覆盖：200 + content-type + RFC 5987 中文文件名 + python-docx 读回关键字段；
无打分快照 / 无 LLM（性质卡生成失败）降级路径；记录不存在 404。
数据目录打到临时目录，不碰真实数据。
"""

from __future__ import annotations

import io

import pytest
from fastapi.testclient import TestClient

from api.main import app

client = TestClient(app)

_DOCX_MEDIA = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


@pytest.fixture
def isolated_data(tmp_path, monkeypatch):
    from favorites import store as fav_store
    from records import store as rec_store
    monkeypatch.setattr(fav_store, "FAVORITES_DIR", tmp_path / "favs")
    monkeypatch.setattr(rec_store, "RECORDS_DIR", tmp_path / "recs")
    monkeypatch.setattr(rec_store, "ATTACHMENTS_DIR", tmp_path / "atts")
    return fav_store, rec_store


def _create_favorite(with_snapshot: bool = True) -> str:
    r = client.post("/api/favorites", json={
        "aldehyde_smiles": "O=Cc1ccccc1", "amine_smiles": "Nc1ccccc1"})
    assert r.status_code == 201
    fid = r.json()["id"]
    from favorites import store as fav_store
    # 内置库未必收录测试单体：显式补上名称/CAS，保证导出断言确定性
    fav_store.update_favorite(fid, aldehyde={
        "smiles": "O=Cc1ccccc1", "cas": "100-52-7", "name": "苯甲醛"}, amine={
        "smiles": "Nc1ccccc1", "cas": "62-53-3", "name": "苯胺"})
    if with_snapshot:
        fav_store.update_prediction_snapshot(fid, {
            "score": 0.82, "std": 0.03, "ood": "none",
            "tree_score": 0.8, "gnn_score": 0.84})
    return fid


def _create_record(fid: str) -> dict:
    r = client.post("/api/records", json={
        "favorite_id": fid,
        "experiment_no": "A5",
        "outcome": "film",
        "strength": "可成片",
        "operator": "张三",
        "notes": "常规界面法",
        "conditions": {"solvent_1": "甲苯", "solvent_2": "氯仿",
                       "eluent": "石油醚", "temperature_c": "120"},
        "process_notes": "投料→陈化→干燥",
        "timeline": [
            {"time_label": "2026-01-03 10:00", "description": "第二天观察"},
            {"time_label": "2026-01-01 09:00", "description": "投料开始"},
        ],
        "self_summary": "总体顺利",
        "mistakes": "加料太快",
    })
    assert r.status_code == 201
    return r.json()


def _docx_text(content: bytes) -> str:
    """python-docx 读回全文（段落 + 表格 + 页脚）。"""
    import docx
    doc = docx.Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                parts.append(c.text)
    for s in doc.sections:
        parts.extend(p.text for p in s.footer.paragraphs)
    return "\n".join(parts)


class TestExportEndpoint:
    def test_export_200_and_headers(self, isolated_data):
        rec = _create_record(_create_favorite())
        r = client.get(f"/api/records/{rec['record_id']}/export")
        assert r.status_code == 200
        assert r.headers["content-type"].startswith(_DOCX_MEDIA)
        cd = r.headers["content-disposition"]
        assert "attachment" in cd
        assert "filename*=" in cd  # RFC 5987 中文文件名
        assert "%E5%AE%9E%E9%AA%8C%E8%AE%B0%E5%BD%95" in cd  # 「实验记录」
        assert r.content[:2] == b"PK"  # docx 是 zip 包

    def test_export_content_key_fields(self, isolated_data):
        rec = _create_record(_create_favorite())
        r = client.get(f"/api/records/{rec['record_id']}/export")
        text = _docx_text(r.content)
        for kw in ("实验记录", "A5", rec["record_id"], "正式",
                   "苯甲醛", "O=Cc1ccccc1", "100-52-7",  # 醛单体名称/SMILES/CAS
                   "苯胺", "62-53-3",                     # 胺单体
                   "0.820", "树模型分量", "GNN 分量",      # 模型打分（含分量）
                   "溶剂一", "甲苯", "洗脱剂", "石油醚",   # 实验条件
                   "投料→陈化→干燥",                       # 完整流程
                   "投料开始", "第二天观察",                # 时间线
                   "总体顺利", "加料太快",                  # 总结 / 失误
                   "由 COF 科研助手导出", "版本 v"):        # 页脚
            assert kw in text, f"导出文档缺少关键内容: {kw}"

    def test_export_timeline_sorted(self, isolated_data):
        rec = _create_record(_create_favorite())
        text = _docx_text(client.get(
            f"/api/records/{rec['record_id']}/export").content)
        # 乱序写入的两条时间点应按 time_label 升序输出
        assert text.index("投料开始") < text.index("第二天观察")

    def test_export_missing_record_404(self, isolated_data):
        r = client.get("/api/records/rec_20990101_001/export")
        assert r.status_code == 404


class TestExportDegradedPaths:
    def test_no_prediction_snapshot(self, isolated_data):
        """关联收藏无打分快照 → 模型打分小节注明「未打分」，其余正常。"""
        rec = _create_record(_create_favorite(with_snapshot=False))
        r = client.get(f"/api/records/{rec['record_id']}/export")
        assert r.status_code == 200
        text = _docx_text(r.content)
        assert "未打分" in text
        assert "甲苯" in text  # 其余小节不受影响

    def test_llm_unavailable_keeps_placeholder(self, isolated_data, monkeypatch):
        """LLM 未配置/性质卡生成失败 → 单体性质小节保留占位，导出仍 200。"""
        import recommend.monomer_props as monomer_props

        def _boom(smiles, name=""):
            raise RuntimeError("LLM 未配置")

        monkeypatch.setattr(monomer_props, "get_monomer_properties", _boom)
        rec = _create_record(_create_favorite())
        r = client.get(f"/api/records/{rec['record_id']}/export")
        assert r.status_code == 200
        text = _docx_text(r.content)
        assert "单体性质" in text          # 小节占位保留
        assert "未配置 LLM" in text

    def test_llm_narrative_rendered(self, isolated_data, monkeypatch):
        """性质卡可用时写入 RDKit 事实与 LLM 解读正文。"""
        import recommend.monomer_props as monomer_props

        def _fake(smiles, name=""):
            return {"facts": {"mw": 106.12},
                    "narrative": "该单体溶解性良好。",
                    "narrative_source": "llm"}

        monkeypatch.setattr(monomer_props, "get_monomer_properties", _fake)
        rec = _create_record(_create_favorite())
        text = _docx_text(client.get(
            f"/api/records/{rec['record_id']}/export").content)
        assert "该单体溶解性良好。" in text
        assert "分子量" in text
