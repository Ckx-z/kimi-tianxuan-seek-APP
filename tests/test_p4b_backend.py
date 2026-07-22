"""P4b 后端测试：单体性质卡 + 方案卡模板系统（任务 C）。"""

from __future__ import annotations

import json
import sys
import types

import pytest

from src.recommend import monomer_props, plan_card, plan_templates


# ---------------------------------------------------------------- facts

class TestFacts:
    def test_benzene_facts(self):
        props = monomer_props.get_monomer_properties("c1ccccc1")
        facts = props["facts"]
        assert set(facts) == set(monomer_props._FACT_KEYS)
        assert facts["aromatic_rings"] == 1
        assert facts["f_count"] == 0
        assert facts["mw"] > 70

    def test_fluorinated(self):
        facts = monomer_props.compute_facts("Fc1ccccc1F")
        assert facts["f_count"] == 2
        assert facts["aromatic_rings"] == 1

    def test_invalid_smiles_empty_facts(self):
        assert monomer_props.compute_facts("not_a_smiles!!") == {}
        assert monomer_props.compute_facts("") == {}
        assert monomer_props.compute_facts(None) == {}

    def test_get_properties_structure(self):
        props = monomer_props.get_monomer_properties("c1ccccc1", name="苯")
        assert set(props) == {"facts", "narrative", "narrative_source"}
        assert props["narrative_source"] in ("llm", "none")


# ---------------------------------------------------------------- narrative 降级

class TestNarrativeFallback:
    def test_no_llm_module_returns_none(self, monkeypatch, tmp_path):
        """src.llm 不存在/未配置时 narrative=None，不抛异常。"""
        monkeypatch.setattr(monomer_props, "CACHE_DIR", tmp_path)
        monkeypatch.setitem(sys.modules, "src.llm", None)
        monkeypatch.setitem(sys.modules, "src.llm.client", None)
        assert (
            monomer_props.get_monomer_properties("c1ccccc1")["narrative"] is None
        )
        assert (
            monomer_props.get_monomer_properties("c1ccccc1")["narrative_source"]
            == "none"
        )

    def test_llm_success_and_cache(self, monkeypatch, tmp_path):
        monkeypatch.setattr(monomer_props, "CACHE_DIR", tmp_path)
        calls = []

        fake = types.SimpleNamespace(
            is_configured=lambda: True,
            chat_completion=lambda messages, max_tokens, temperature: (
                calls.append(messages) or "这是 LLM 生成的解读。"
            ),
        )
        monkeypatch.setattr(monomer_props, "_llm_client", lambda: fake)

        p1 = monomer_props.get_monomer_properties("c1ccccc1")
        assert p1["narrative"] == "这是 LLM 生成的解读。"
        assert p1["narrative_source"] == "llm"
        assert len(calls) == 1

        # 同 SMILES（等价写法）命中缓存，不再调 LLM
        p2 = monomer_props.get_monomer_properties("C1=CC=CC=C1")
        assert p2["narrative"] == "这是 LLM 生成的解读。"
        assert len(calls) == 1

    def test_llm_none_response(self, monkeypatch, tmp_path):
        monkeypatch.setattr(monomer_props, "CACHE_DIR", tmp_path)
        fake = types.SimpleNamespace(
            is_configured=lambda: True,
            chat_completion=lambda **kw: None,
        )
        fake.chat_completion = lambda messages, max_tokens, temperature: None
        monkeypatch.setattr(monomer_props, "_llm_client", lambda: fake)
        assert monomer_props.get_monomer_properties("c1ccccc1")["narrative"] is None


# ---------------------------------------------------------------- 模板 CRUD

@pytest.fixture()
def tpl_dir(tmp_path, monkeypatch):
    """隔离模板目录，内置模板拷入。"""
    d = tmp_path / "plan_templates"
    d.mkdir()
    src = plan_templates.TEMPLATES_DIR / "builtin_hou_v3_9.json"
    (d / "builtin_hou_v3_9.json").write_text(src.read_text(encoding="utf-8"), encoding="utf-8")
    monkeypatch.setattr(plan_templates, "TEMPLATES_DIR", d)
    return d


class TestTemplateCRUD:
    def test_builtin_template_valid(self):
        tpl = json.loads(
            (plan_templates.PROJECT_ROOT / "data" / "plan_templates" / "builtin_hou_v3_9.json")
            .read_text(encoding="utf-8")
        )
        out = plan_templates.validate_template(tpl)
        assert out["id"] == plan_templates.BUILTIN_ID
        assert out["steps"] and out["checklist"] and out["hints_rules"]

    def test_list_and_get(self, tpl_dir):
        templates = plan_templates.list_templates()
        assert len(templates) == 1
        assert templates[0]["builtin"] is True
        tpl = plan_templates.get_template(plan_templates.BUILTIN_ID)
        assert tpl["name"].startswith("侯老师")

    def test_save_and_list_user_template(self, tpl_dir):
        user_tpl = {
            "id": "user_test",
            "name": "测试模板",
            "source": "测试",
            "conditions": {"temperature_c": 80},
            "steps": ["步骤一"],
            "checklist": [{"item": "核对", "detail": "细节"}],
            "hints_rules": [{"rule": "r", "match": "m", "hint": "h"}],
        }
        saved = plan_templates.save_template(user_tpl)
        assert (tpl_dir / "user_test.json").exists()
        ids = [t["id"] for t in plan_templates.list_templates()]
        assert "user_test" in ids and plan_templates.BUILTIN_ID in ids
        user_entry = [t for t in plan_templates.list_templates() if t["id"] == "user_test"][0]
        assert user_entry["builtin"] is False

    def test_get_missing_raises(self, tpl_dir):
        with pytest.raises(plan_templates.TemplateError):
            plan_templates.get_template("no_such_template")

    def test_validate_rejects_bad_schema(self):
        with pytest.raises(plan_templates.TemplateError):
            plan_templates.validate_template({"id": "x"})
        with pytest.raises(plan_templates.TemplateError):
            plan_templates.validate_template("not a dict")

    def test_save_template_id_sanitized(self, tpl_dir):
        tpl = {
            "id": "user ../evil",
            "name": "n",
            "source": "s",
            "conditions": {},
            "steps": [],
            "checklist": [],
            "hints_rules": [],
        }
        plan_templates.save_template(tpl)
        # 文件名被净化，不逃逸目录
        assert list(tpl_dir.glob("user*evil*.json"))
        assert not (tpl_dir.parent / "evil.json").exists()


# ---------------------------------------------------------------- docx 提取

class TestDocxExtract:
    def _make_docx(self, path):
        import docx

        doc = docx.Document()
        doc.add_paragraph("界面法合成 COF：甲苯/氯仿，120 °C，苯胺调制剂。")
        doc.add_paragraph("步骤：先醛后胺，最后 6M 乙酸。")
        doc.save(str(path))

    def test_extract_success_mock_llm(self, tmp_path, monkeypatch):
        docx_path = tmp_path / "scheme.docx"
        self._make_docx(docx_path)

        llm_json = json.dumps(
            {
                "name": "文献法",
                "source": "测试来源",
                "conditions": {"solvent": "甲苯", "temperature_c": 120},
                "steps": ["先醛后胺", "最后加乙酸"],
                "checklist": [{"item": "浓度", "detail": "6M"}],
                "hints_rules": [{"rule": "fluorinated", "match": "含F", "hint": "预溶"}],
            },
            ensure_ascii=False,
        )
        fake_client = types.SimpleNamespace(
            is_configured=lambda: True,
            chat_completion=lambda messages, max_tokens, temperature: llm_json,
        )
        fake_llm = types.ModuleType("src.llm")
        fake_llm.client = fake_client
        monkeypatch.setitem(sys.modules, "src.llm", fake_llm)
        monkeypatch.setitem(sys.modules, "src.llm.client", fake_client)

        tpl = plan_templates.extract_template_from_docx(docx_path, name="文献法")
        assert tpl["name"] == "文献法"
        assert tpl["id"].startswith("user_")
        assert tpl["conditions"]["temperature_c"] == 120
        # 不落盘
        assert not (plan_templates.TEMPLATES_DIR / f"{tpl['id']}.json").exists()

    def test_extract_llm_not_configured(self, tmp_path, monkeypatch):
        docx_path = tmp_path / "scheme.docx"
        self._make_docx(docx_path)
        fake_client = types.SimpleNamespace(is_configured=lambda: False)
        fake_llm = types.ModuleType("src.llm")
        fake_llm.client = fake_client
        monkeypatch.setitem(sys.modules, "src.llm", fake_llm)
        monkeypatch.setitem(sys.modules, "src.llm.client", fake_client)
        with pytest.raises(plan_templates.TemplateError, match="未配置 LLM"):
            plan_templates.extract_template_from_docx(docx_path)

    def test_extract_bad_json(self, tmp_path, monkeypatch):
        docx_path = tmp_path / "scheme.docx"
        self._make_docx(docx_path)
        fake_client = types.SimpleNamespace(
            is_configured=lambda: True,
            chat_completion=lambda messages, max_tokens, temperature: "这不是JSON",
        )
        fake_llm = types.ModuleType("src.llm")
        fake_llm.client = fake_client
        monkeypatch.setitem(sys.modules, "src.llm", fake_llm)
        monkeypatch.setitem(sys.modules, "src.llm.client", fake_client)
        with pytest.raises(plan_templates.TemplateError):
            plan_templates.extract_template_from_docx(docx_path)


# ---------------------------------------------------------------- plan_card 模板注入

class TestPlanCardTemplate:
    ALD = "O=Cc1ccccc1"
    AMI = "Nc1ccccc1"

    def test_default_backward_compatible(self):
        card = plan_card.generate_plan_card(self.ALD, self.AMI)
        assert card["template"] == plan_card.TEMPLATE_NAME
        assert card["conditions"]["temperature_c"] == 120
        assert card["defaults_note"] == plan_card.DEFAULTS_NOTE
        assert len(card["steps"]) == 7

    def test_template_injection(self):
        tpl = plan_templates.get_template(plan_templates.BUILTIN_ID)
        tpl["name"] = "自定义法"
        tpl["conditions"] = {"solvent": "二氧六环", "temperature_c": 80}
        tpl["steps"] = ["自定义步骤1", "自定义步骤2"]
        card = plan_card.generate_plan_card(self.ALD, self.AMI, template=tpl)
        assert card["template"] == "自定义法"
        assert card["conditions"]["solvent"] == "二氧六环"
        assert card["steps"] == ["自定义步骤1", "自定义步骤2"]
        # monomer_hints 仍由结构检测生成
        assert "monomer_hints" in card

    def test_template_fluorine_hint_still_works(self):
        tpl = {"name": "x", "conditions": {}, "steps": [], "checklist": []}
        card = plan_card.generate_plan_card("O=Cc1ccc(F)cc1", self.AMI, template=tpl)
        assert any("含氟" in h for h in card["monomer_hints"])
