"""
bridge/generate_v7.py  (重写版)
==============================
v7 基于侯老师原文机理 (固-液/气三相界面 + 扩散/调制剂双介导)
- 中文术语
- DOI + 本地路径
- 美化排版 (docx 样式系统)
- 每步说明"为什么这么做" (沸点 + 夹带机理)
- 单体量的化学计量解释
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.styles.style import _ParagraphStyle

PROJ = r'C:\Users\ckx\Desktop\minimax'
OUT = os.path.join(PROJ, 'experiment', 'proposals',
                   'COF-TFPT-TAPT-2026-07-13-D-A1-v7.docx')

# ==================== 样式工具 ====================

def setup_styles(doc):
    """统一设置文档样式"""
    # Normal 样式
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')

    # Heading 1 - 章节标题
    h1 = doc.styles['Heading 1']
    h1.font.name = '黑体'
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.line_spacing = 1.5
    rPr1 = h1.element.get_or_add_rPr()
    rFonts1 = rPr1.find(qn('w:rFonts'))
    if rFonts1 is None:
        rFonts1 = OxmlElement('w:rFonts')
        rPr1.append(rFonts1)
    rFonts1.set(qn('w:eastAsia'), '黑体')
    rFonts1.set(qn('w:ascii'), 'Times New Roman')
    rFonts1.set(qn('w:hAnsi'), 'Times New Roman')

    # Heading 2 - 子章节
    h2 = doc.styles['Heading 2']
    h2.font.name = '黑体'
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x2E, 0x5C, 0x8A)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.line_spacing = 1.5

    # Heading 3 - 重点步骤标题
    h3 = doc.styles['Heading 3']
    h3.font.name = '黑体'
    h3.font.size = Pt(11.5)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.line_spacing = 1.5


def set_cell_bg(cell, color_hex):
    """设置单元格背景色"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def set_para_bg(para, color_hex):
    """段落背景"""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    pPr.append(shd)


def add_run(p, text, size=10.5, bold=False, italic=False, color=None, font='宋体'):
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    return run


def add_h1(doc, text):
    p = doc.add_heading('', level=1)
    add_run(p, text, size=16, bold=True, color=RGBColor(0x1F, 0x3A, 0x5F), font='黑体')
    return p


def add_h2(doc, text):
    p = doc.add_heading('', level=2)
    add_run(p, text, size=13, bold=True, color=RGBColor(0x2E, 0x5C, 0x8A), font='黑体')
    return p


def add_h3(doc, text):
    p = doc.add_heading('', level=3)
    add_run(p, text, size=11.5, bold=True, color=RGBColor(0xC0, 0x39, 0x2B), font='黑体')
    return p


def add_p(doc, text, indent=True, size=10.5, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    add_run(p, text, size=size, bold=bold, italic=italic, color=color)
    return p


def add_bullet(doc, text, size=10.5):
    """蓝色项目符号"""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.first_line_indent = Cm(0.5)
    add_run(p, text, size=size)
    return p


def add_callout(doc, label, content, color_hex='F0F7FF'):
    """高亮信息框 (灰色/淡蓝背景)"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    set_para_bg(p, color_hex)
    add_run(p, f'【{label}】', size=10.5, bold=True, color=RGBColor(0xC0, 0x39, 0x2B))
    add_run(p, content, size=10.5)
    return p


def add_table(doc, header, rows, header_bg='1F3A5F', col_widths=None):
    """美化表格 (深蓝表头 + 浅色边框)"""
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'

    # 设置列宽
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = w

    # 表头
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_bg(hdr_cells[i], header_bg)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 数据行
    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        # 斑马纹
        bg = 'F8F9FA' if r_idx % 2 == 0 else 'FFFFFF'
        for i, v in enumerate(row_data):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            add_run(p, str(v), size=9.5)
            set_cell_bg(cells[i], bg)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return table


def add_page_break(doc):
    doc.add_page_break()


# ==================== 主文档 ====================

def main():
    doc = Document()
    setup_styles(doc)

    # 设置页边距
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # ===== 封面 =====
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(36)
    title_p.paragraph_format.space_after = Pt(12)
    add_run(title_p, '含氟亚胺键 COF 膜合成实验方案',
            size=22, bold=True, color=RGBColor(0x1F, 0x3A, 0x5F), font='黑体')

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(6)
    add_run(subtitle_p, 'D 系 + A1 系 迭代方案 v7',
            size=16, bold=True, color=RGBColor(0x2E, 0x5C, 0x8A), font='黑体')

    # 副标题
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2.paragraph_format.space_after = Pt(36)
    add_run(sub2, '基于侯盛怀等固-液/气三相界面法 + 扩散/调制剂双介导机理',
            size=11, italic=True, color=RGBColor(0x6C, 0x75, 0x7D))

    # 元信息表
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(info_p, '方案编号: ', size=10.5, bold=True)
    add_run(info_p, 'COF-TFPT-TAPT-2026-07-13-D-A1-v7', size=10.5)
    info_p.add_run().add_break()
    add_run(info_p, '生成日期: ', size=10.5, bold=True)
    add_run(info_p, '2026-07-13', size=10.5)
    info_p.add_run().add_break()
    add_run(info_p, '迭代依据: ', size=10.5, bold=True)
    add_run(info_p, '14 条 ABCDEF + A1-D9 历史失败', size=10.5)
    info_p.add_run().add_break()
    add_run(info_p, '核心参考文献: ', size=10.5, bold=True)
    add_run(info_p, 'Hou et al., Angew. Chem. Int. Ed. 2025, 64, e202421555 (DOI: 10.1002/anie.202421555)',
            size=10.5, italic=True)
    info_p.add_run().add_break()
    add_run(info_p, '   本地路径: ', size=10.5)
    add_run(info_p, r'实验\文章\侯老师实验\侯盛怀德国应化.pdf', size=10, italic=True)
    info_p.add_run().add_break()
    add_run(info_p, '知识库检索: ', size=10.5, bold=True)
    add_run(info_p, 'GraphRAG v2 (9066 节点 / 23920 边 / 954 文献)', size=10.5)

    add_page_break(doc)

    # ===== 摘要 =====
    add_h1(doc, '摘要')
    add_p(doc, '本方案针对实验 ABCDEF 与 A1-D9 共 14 条历史失败反馈, 基于侯盛怀等于德国应用化学发表的固-液/气三相界面合成策略, 重新设计 D 系 (TFPT + TFMB) 与 A1 系 (TAPT + A6) 体系的实验方案。核心要点如下:')

    add_callout(doc, '关键发现',
                '侯老师文献的核心是【扩散/调制剂双介导】机理——低沸点溶剂 (氯仿 bp 61.2°C / 甲醇 bp 64.7°C) 在 120°C 油浴中快速蒸发并【夹带】单体向上输运, 在玻璃壁面形成超饱和铺展液层作为限域微反应器, 实现 COF 膜的有序生长。苯胺作为调制剂优先与醛基反应形成希夫碱中间体 (MS 检出 m/z 469.11), 随后与二胺缓慢交换完成聚合。',
                color_hex='FFF8E1')

    add_bullet(doc, '【D 系失败根因】用户 D 实验用了 12 M 乙酸 (过量 2 倍) + 量放大 2 倍, 偏离侯老师原文 6 M 乙酸 + 标准量级 → 反应过快, 无法形成超饱和液层, 导致沉淀。本方案严格回归侯老师原文参数。')
    add_bullet(doc, '【A1 系失败根因】自反应法 (A1/A2/A5/D3) 与新加料序 (A8/D9) 都偏离了【固-液/气三相界面】设计理念。本方案回归界面法, 但基于 A6 大位阻特性调整溶剂比例与反应时间。')
    add_bullet(doc, '【方法论】v3.9 SOP 是基准, 本方案对每步操作给出【为什么这么做】的化学解释, 避免下次再操作错。')

    # ===== 第一章: 侯老师文献机理 (核心) =====
    add_h1(doc, '第一章 侯老师文献机理 (本方案的科学基础)')

    add_p(doc, '本方案的所有设计依据来自侯盛怀等于 2025 年发表于《德国应用化学》的论文, 题: "扩散/调制剂双介导固-液/气界面合成结晶型共价有机框架膜"。下文将详细梳理本方案需要遵循的核心理论框架。')

    add_h2(doc, '1.1 侯老师方法的核心创新: 固-液/气三相界面')
    add_p(doc, '传统界面聚合法 (IP) 主要分为气-液、液-液、固-气、液-固四类, 各有局限:')
    add_bullet(doc, '气-液 / 液-液 IP: 形成的膜太薄, 必须转移到多孔基底才能测试')
    add_bullet(doc, '固-气 IP: 需要复杂基底刻蚀工艺')
    add_bullet(doc, '液-固 IP: 膜取向与结晶度随厚度增加而恶化 (权衡效应)')

    add_p(doc, '侯老师方法的核心创新: 利用玻璃管内壁这一【常被忽视的坚固表面】作为基底, 通过溶剂蒸发与调制剂调控, 在玻璃壁上原位生长自支撑、定向、高结晶度的 COF 膜。')

    add_h2(doc, '1.2 扩散介导机理 (本方案最重要的核心)')

    add_callout(doc, '夹带作用机理',
                '低沸点溶剂在 120°C 油浴中快速蒸发, 其蒸汽流【夹带】单体分子一起向上输运 (entrainment), 在玻璃管内壁 (极性玻璃表面) 形成一层极薄 (几微米厚) 的超饱和铺展液层, 这层液层作为限域微反应器, 大幅降低成核势垒, 促进 COF 膜在玻璃壁上定向生长。【此即"扩散介导"的核心】',
                color_hex='FFF8E1')

    add_p(doc, '夹带作用的关键条件:')

    # 溶剂沸点表
    add_table(doc,
              ['溶剂', '沸点 (°C)', '120°C 下蒸发速率', '夹带作用', '是否可用'],
              [
                  ['甲醇', '64.7', '极快', '强', '可用'],
                  ['二氯甲烷', '39.6', '极快', '强 (易爆沸, 慎用)', '可用'],
                  ['氯仿', '61.2', '快', '强 (平衡润湿性)', '推荐'],
                  ['甲苯', '110.6', '慢', '弱 (但润湿性好)', '推荐 (作润湿剂)'],
                  ['1,4-二氧六环', '101.3', '慢', '弱', '不可用'],
                  ['均三甲苯', '164.7', '极慢', '几乎无', '不可用'],
                  ['正丁醇', '117.7', '慢', '弱', '不可用'],
                  ['1,2-二氯苯', '180.4', '极慢', '无', '不可用'],
              ],
              col_widths=[Cm(2.5), Cm(2.0), Cm(3.5), Cm(4.0), Cm(3.5)])

    add_p(doc, '表 1 溶剂沸点与夹带作用关系 (引自侯老师 Table S5)')

    add_p(doc, '结论: 溶剂的沸点必须【显著低于】反应温度 (120°C), 以提供足够强的夹带驱动力。本方案优选【氯仿 (61.2°C)】作为夹带剂, 配以【甲苯 (110.6°C)】作为润湿剂 (玻璃接触角小), 二者比例 3:7 (v/v) 是侯老师原文最优条件。')

    add_h2(doc, '1.3 调制剂介导机理 (苯胺的真实作用)')

    add_p(doc, '苯胺 (一取代单胺) 在 COF 膜合成中的真实作用, 远超"封端醛基"的简单理解。侯老师用 HPLC-MS/MS 直接监测反应过程, 发现:')

    add_bullet(doc, '【步骤 1: 苯胺优先反应】反应初始阶段, 苯胺与 TFPT 反应生成苯胺-TFPT 希夫碱中间体 (MS 检出 m/z 469.11), 速率远快于 TFPT 直接与二胺反应。')
    add_bullet(doc, '【步骤 2: 缓慢亚胺交换】苯胺-TFPT 中间体随后与二胺 (BD-CF3 等) 发生缓慢交换反应 (中间体 IV, m/z 846.29), 生成目标 COF。')
    add_bullet(doc, '【消耗速率证据】苯胺消耗速率 : 二胺消耗速率 = 6.25 : 1 (苯胺优先消耗); 苯胺存在时平衡速率 0.81 mmol/L/h, 无苯胺时 1.94 mmol/L/h (苯胺降低反应速率, 提供动力学调控)。')
    add_bullet(doc, '【苯胺量优化】3 eq (相对 TFPT 的醛基数) 是结晶度最佳 (PXRD fwhm = 0.16°); 0/1/2 eq 结晶度差 (fwhm = 0.45°); 4/5 eq 过量, 反而破坏可逆性 (希夫碱交换被抑制), 结晶度下降。')

    add_callout(doc, '本方案结论',
                '苯胺必须用 3 eq (相对醛基数), 不能多不能少; 必须先与醛基混合 (形成希夫碱中间体), 再加二胺。本方案所有步骤严格遵循此顺序。',
                color_hex='E8F5E9')

    add_h2(doc, '1.4 玻璃壁预排列机理 (DFT 计算支持)')

    add_p(doc, 'DFT 计算表明: 玻璃壁的 -OH 基团与单体 (醛基/氨基) 形成强氢键 → 单体在玻璃壁上【预排列】为有序结构 → 聚合时直接形成定向晶域。')

    add_bullet(doc, 'PTFE 壁 (低表面能) 无膜 → 验证玻璃壁的氢键是关键')
    add_bullet(doc, '均相沉淀不在体相发生 → 验证玻璃壁的限域作用是关键')
    add_bullet(doc, 'DFT 模拟证实玻璃壁的强氢键诱导单体自组装')

    add_h2(doc, '1.5 侯老师 13 个 AMCOF 体系小结')

    add_p(doc, '侯老师用此方法成功制备了 13 个 AMCOF 膜, 体系覆盖: 不同节点 (TFPT, Tp, TPA, TPT, TFFT, TTFT)、不同连接臂 (BD-CF3, BD-CH3, PDA-CF3 等)、不同官能团 (氟/甲基/腈基/苯基取代)。其中:')

    add_table(doc,
              ['编号', '醛节点', '胺连接臂', '溶剂', '乙酸', '时间', '产率'],
              [
                  ['AMCOF-1', 'TFPT', 'BD-CF3 (TFMB)', '甲苯/氯仿 3/7', '6 M', '2 天', '88.36%'],
                  ['AMCOF-2', 'TFPT', 'BD-CH3', '甲苯/甲醇 4/6', '3 M', '3 天', '65.90%'],
                  ['AMCOF-3', 'TFPT', 'PDA-CF3 (B6)', '甲苯/氯仿 3/7', '6 M', '2 天', '71.20%'],
                  ['AMCOF-4 ~ 13', '(原文 Figure 1B)', '—', '—', '—', '—', '—'],
              ],
              col_widths=[Cm(2.5), Cm(2.5), Cm(3.5), Cm(3.5), Cm(2.0), Cm(2.0), Cm(2.5)])

    add_p(doc, '表 2 侯老师 AMCOF-1 ~ AMCOF-3 详细参数 (引自 Scheme S1-S3)')

    add_callout(doc, '本方案核心参考',
                '【AMCOF-1】与用户的【D 系】单体对完全一致 (TFPT + BD-CF3 = TFMB), 是直接可参考的合成方案。本方案 D-1 严格按 AMCOF-1 原文参数执行, 预期重现 ~88% 产率。',
                color_hex='FFF8E1')

    # ===== 第二章: 单体物性 =====
    add_h1(doc, '第二章 单体物性 (本方案使用的单体)')

    add_h2(doc, '2.1 单体结构与基本物性')

    add_table(doc,
              ['代号', '中文名', '英文名', 'CAS', '分子量', '官能团', '物性'],
              [
                  ['TFPT', '1,3,5-三(4-甲酰基苯基)-1,3,5-三嗪', '4,4\',4\'\'-(1,3,5-Triazine-2,4,6-triyl)tribenzaldehyde',
                   '443922-06-3', '393.4', '-CHO × 3 (醛基)', '黄色固体'],
                  ['TAPT', '1,3,5-三(4-氨基苯基)-1,3,5-三嗪', '2,4,6-Tris(4-aminophenyl)-1,3,5-triazine',
                   '14544-47-9', '354.4', '-NH₂ × 3 (氨基)', '黄色固体'],
                  ['TFMB / B5 / BD-CF3', '2,2\'-双(三氟甲基)-4,4\'-联苯二胺', '2,2\'-Bis(trifluoromethyl)benzidine',
                   '341-58-2', '320.2', '-NH₂ × 2 (氨基)', '白色固体'],
                  ['A6', '4,4\'\'-双(三氟甲基)-2\',5\'-二甲酰基对三联苯', '4,4\'\'-Bis(trifluoromethyl)-2\',5\'-diformyl-p-terphenyl',
                   '1300701-03-4', '422.4', '-CHO × 2 (醛基)', '淡黄色固体'],
              ],
              col_widths=[Cm(1.8), Cm(4.0), Cm(4.5), Cm(2.5), Cm(1.5), Cm(2.5), Cm(2.0)])

    add_p(doc, '表 3 本方案使用的四个单体 (TFPT/TAPT 为节点, TFMB/A6 为线性连接臂)')

    add_h2(doc, '2.2 单体溶解度 (决定溶剂选择)')

    add_p(doc, '下表汇总本方案使用单体的溶解度, 是选择溶剂体系的依据。标度: ++ 易溶 (>20 mg/mL), + 可溶 (5-20 mg/mL), +/- 微溶 (1-5 mg/mL), - 几乎不溶 (<1 mg/mL)。')

    add_table(doc,
              ['单体', '甲醇', '乙醇', '甲苯', '氯仿', '三氟甲苯', '二氧六环', 'DMF', 'DMSO'],
              [
                  ['TAPT', '-', '-', '-', '+/-', '+/-', '+', '++', '++'],
                  ['TFPT', '-', '-', '-', '+/-', '+', '+', '++', '++'],
                  ['A6', '-', '-', '+', '++', '++', '+', '+', '+'],
                  ['TFMB', '-', '-', '+/-', '+', '+', '+', '++', '++'],
              ],
              col_widths=[Cm(1.5), Cm(1.5), Cm(1.5), Cm(1.5), Cm(1.5), Cm(2.0), Cm(1.8), Cm(1.5), Cm(1.5)])

    add_p(doc, '表 4 单体溶解度 (引自实验方案 v3.9 §2.5)')

    add_p(doc, '关键观察:')
    add_bullet(doc, 'TAPT 和 TFPT 在 DMF/DMSO 中溶解度极佳 (++), 但 DMF/DMSO 与水相 (含乙酸) 互溶, 无法形成稳定液-液界面, 且与本方案的【固-液/气三相界面】设计冲突 → 排除 DMF/DMSO')
    add_bullet(doc, 'A6 和 TFMB 在甲苯/氯仿中溶解度好 (+ ~ ++), 是合适的有机相溶剂')
    add_bullet(doc, 'TAPT/TFPT 微溶于氯仿 (+/-), 必须用甲苯/氯仿混合溶剂 (3:7) 才能完全溶解')

    add_h2(doc, '2.3 溶剂极性与选择依据')

    add_table(doc,
              ['溶剂', '沸点 (°C)', '介电常数', '偶极矩 (D)', '极性', '本方案角色'],
              [
                  ['DMSO', '189', '46.7', '3.96', '强极性', '排除 (与水互溶)'],
                  ['DMF', '153', '36.7', '3.86', '强极性', '排除 (与水互溶)'],
                  ['甲醇', '64.7', '32.7', '1.70', '强极性', '可选 (强夹带)'],
                  ['乙醇', '78.4', '24.5', '1.69', '较强极性', '可选'],
                  ['1,4-二氧六环', '101.3', '2.2', '0.45', '中等极性', '排除 (沸点过高)'],
                  ['氯仿', '61.2', '4.8', '1.15', '弱极性', '★ 关键 (夹带剂)'],
                  ['三氟甲苯 (BTF)', '102', '~9.2', '~2.6', '弱-中等极性', '可选 (含氟膜亲氟)'],
                  ['甲苯', '110.6', '2.4', '0.36', '非极性', '★ 关键 (润湿剂)'],
              ],
              col_widths=[Cm(2.5), Cm(1.8), Cm(2.0), Cm(2.5), Cm(2.0), Cm(4.5)])

    add_p(doc, '表 5 溶剂极性与本方案角色 (介电常数引自 depts.washington.edu/eooptic)')

    add_callout(doc, '溶剂设计原理',
                '本方案优选【甲苯/氯仿 = 3/7 (v/v)】混合溶剂: (a) 甲苯润湿性好 (玻璃接触角小) → 液层能铺展到玻璃壁; (b) 氯仿沸点低 (61.2°C) → 120°C 下蒸发快, 夹带单体能力强; (c) 二者比例 3:7 是侯老师原文最优 (Table S5); (d) 甲苯极低极性不会破坏亚胺键可逆性。',
                color_hex='FFF8E1')

    # ===== 第三章: 历史失败分析 =====
    add_h1(doc, '第三章 历史失败深度分析 (14 条反馈)')

    add_p(doc, '下表汇总用户 14 条实验反馈, 并基于侯老师机理给出根因诊断。')

    add_table(doc,
              ['ID', '醛', '胺', '概率', 'Class', '用户描述', '侯老师机理诊断', '本方案修正'],
              [
                  ['D', 'TFPT', 'TFMB', '0.596', 'E',
                   '白色固体 0.2 cm 粗糙, 乙酸 0.4 mL × 12 M',
                   '【乙酸 12 M 过量】破坏可逆性 → 沉淀而非膜',
                   '改为 6 M 乙酸, 量不放大'],
                  ['D3', 'TFPT', 'TFMB', '0.596', 'E',
                   '12 h 自反应膜反复溶解, 加 TFMB 操作失误爆沸',
                   '【自反应法】偏离固-液/气界面 → 玻璃壁上没超饱和液层',
                   '回归侯老师原文 4 步加料序'],
                  ['D4', 'TFPB', 'TFMB', '0.596', 'A',
                   '#aborted 未冷却开盖 → 氯仿爆沸',
                   '【操作错】违反温度安全规程',
                   '【废除】改用 TFPT 节点'],
                  ['D7', 'TFPT', 'TFMB', '0.596', 'E',
                   '标准序 72 h 后不连续粗糙浅黄色固体',
                   '与 D 类似, 乙酸 + 量错误',
                   '改为 6 M 乙酸, 量不放大'],
                  ['D9', 'TFPT', 'TFMB', '0.596', 'E',
                   '放大 3 倍 + 新序, 膜部分碳化',
                   '【量放大破坏夹带动力学】夹带速率跟不上反应速率',
                   '【严禁放大】严格按侯老师原文量级'],
                  ['A', 'TAPB', 'A6', '0.819', 'E',
                   '黄褐色固体, 不连续粗糙',
                   '【TAPB 节点不适合】+ 乙酸过量',
                   '改用 TAPT 节点 (更对称)'],
                  ['A1', 'TAPT', 'A6', '0.819', 'E',
                   '12 h 自反应膜, 加醛后一碰就碎',
                   '【自反应法】+ A6 大位阻影响苯胺-醛中间体形成',
                   '回归固-液/气界面, 用 BTF 增强溶解'],
                  ['A2', 'TAPB', 'A6', '0.819', 'E',
                   'TAPB 自反应更差',
                   'TAPB 比 TAPT 反应性弱',
                   '改用 TAPT 节点'],
                  ['A5', 'TAPT', 'A6', '0.819', 'E',
                   '氯仿 1.08 vs 0.88 mL 微调, 无效',
                   '微调不足以解决根本问题',
                   '回归侯老师原文比例'],
                  ['A8', 'TAPT', 'A6', '0.819', 'E',
                   '新加料序仍粗糙',
                   '新序破坏了苯胺-TFPT 中间体的优先形成',
                   '回归"醛+苯胺 → 立即胺 → 最后酸"'],
              ],
              col_widths=[Cm(1.0), Cm(1.3), Cm(1.5), Cm(1.2), Cm(1.0), Cm(2.5), Cm(3.5), Cm(2.5)])

    add_p(doc, '表 6 历史失败 10 条 + 侯老师机理诊断 + 本方案修正')

    add_callout(doc, '核心诊断',
                '【14 个失败中 11 个的根本原因都是操作错 (乙酸过量 12 M vs 应 6 M / 量放大 / 偏离加料序)】, 而非化学体系错。tianxuan-seek 模型预测 TFPT+TFMB 概率 0.596, 是中等概率, 侯老师 AMCOF-1 实际产率 88.36% 已经验证——只要严格执行侯老师 SOP, 就能成功。',
                color_hex='E8F5E9')

    # ===== 第四章: 化学计量与单体量 =====
    add_h1(doc, '第四章 单体量与化学计量 (本方案的化学基础)')

    add_h2(doc, '4.1 化学计量比 (TFPT : TFMB = 1 : 1.5)')

    add_p(doc, 'TFPT 含 3 个醛基 (-CHO), TFMB 含 2 个氨基 (-NH₂)。希夫碱反应 1 个 -CHO 消耗 1 个 -NH₂, 形成 -CH=N- 键。化学计量比推导:')

    add_p(doc, '设 TFPT 取 n₁ mmol (含 3n₁ 个 -CHO), TFMB 取 n₂ mmol (含 2n₂ 个 -NH₂)。')
    add_p(doc, '完全反应条件: 3n₁ = 2n₂, 即 n₂ = 1.5 n₁')
    add_p(doc, '侯老师 AMCOF-1 实际取 n₁ = 0.03 mmol, n₂ = 0.045 mmol = 1.5 × 0.03 ✓')

    add_table(doc,
              ['项目', '侯老师 AMCOF-1', '本方案 D-1 (严格回归)', '比例'],
              [
                  ['TFPT', '11.8 mg (0.03 mmol)', '11.8 mg (0.03 mmol)', '1.0'],
                  ['TFMB / BD-CF3', '14.4 mg (0.045 mmol)', '14.4 mg (0.045 mmol)', '1.5'],
                  ['TFPT : TFMB 摩尔比', '1 : 1.5', '1 : 1.5', '—'],
                  ['醛基 : 氨基 当量比', '3 × 0.03 : 2 × 0.045 = 1 : 1', '1 : 1', '化学计量'],
              ],
              col_widths=[Cm(4.0), Cm(5.0), Cm(5.5), Cm(3.0)])

    add_p(doc, '表 7 TFPT + TFMB 化学计量对照')

    add_h2(doc, '4.2 苯胺量 (3 eq 是关键)')

    add_p(doc, '苯胺用量按醛基数计算: 苯胺 : TFPT 醛基 = 3 : 1, 即苯胺 = 3 × 0.03 = 0.09 mmol = 8.2 μL (苯胺密度 1.022 g/mL, 摩尔质量 93.13 g/mol)。')

    add_p(doc, '侯老师用 HPLC 直接测了苯胺消耗动力学, 证明 3 eq 最佳:')
    add_bullet(doc, '0 eq: 膜为多晶微球堆叠 (fwhm = 0.45°), 无连续膜')
    add_bullet(doc, '1 eq: 结晶度差, 颗粒不连续')
    add_bullet(doc, '2 eq: 接近最佳, 但仍有缺陷')
    add_bullet(doc, '3 eq: ★ 最佳 (fwhm = 0.16°), 连续完整膜, 棱面晶')
    add_bullet(doc, '4 eq: 结晶度下降 (过量苯胺抑制亚胺交换)')
    add_bullet(doc, '5 eq: 更差')

    add_h2(doc, '4.3 乙酸浓度 (6 M 是关键)')

    add_p(doc, '乙酸 (CH₃COOH / HOAc) 作为催化剂, 浓度选择有严格范围:')
    add_bullet(doc, '3 M: 催化不足, 结晶度差 (AMCOF-2 用 3 M 因为 BD-CH3 反应性较强)')
    add_bullet(doc, '6 M: ★ 最佳 (AMCOF-1, AMCOF-3) — 与苯胺 pKa 匹配, 既能催化又能维持可逆性')
    add_bullet(doc, '12 M: 用户 D 系列错误, 过量乙酸把平衡推到不可逆, 导致沉淀')
    add_bullet(doc, '18 M: 用户 ABCDEF 错误, 严重过量')

    add_callout(doc, '6 M 乙酸的最佳性证据',
                '侯老师 Table S8 比较了 0.6 / 1.5 / 3 / 6 / 9 / 12 M 乙酸, 发现 6 M 时 fwhm 最小 (0.16°), 结晶度最高。其机制: 6 M 乙酸的 pH (~1.7) 与苯胺的 pKa (4.6) 形成最佳酸碱缓冲, 既能催化希夫碱形成, 又能让亚胺交换维持可逆性。',
                color_hex='E8F5E9')

    add_h2(doc, '4.4 溶剂比例 (甲苯/氯仿 = 3/7)')

    add_p(doc, '侯老师测试了多种溶剂组合, 甲苯/氯仿 3/7 是结晶度最优的 (Table S5):')
    add_bullet(doc, '纯氯仿: 反应过快, 膜厚但不均匀')
    add_bullet(doc, '纯甲苯: 沸点 110.6°C, 120°C 下夹带作用弱, 无膜')
    add_bullet(doc, '甲苯/氯仿 = 7/3: 氯仿少, 夹带不足')
    add_bullet(doc, '甲苯/氯仿 = 3/7: ★ 最佳 (甲苯提供润湿性, 氯仿提供夹带)')
    add_bullet(doc, '甲苯/氯仿 = 1/9: 膜不连续 (润湿性不足)')
    add_bullet(doc, '纯 1,4-二氧六环 / 均三甲苯 / 正丁醇: 无膜 (沸点过高, 夹带失效)')

    add_h2(doc, '4.5 反应温度 (120°C 是关键)')

    add_p(doc, '侯老师测试了 25 / 60 / 90 / 120 / 150°C:')
    add_bullet(doc, '25°C: 体相中只得到无定形粉末 (溶剂蒸发太慢, 玻璃壁没超饱和液层)')
    add_bullet(doc, '60°C: 开始在玻璃壁上成膜, 结晶度低')
    add_bullet(doc, '90-120°C: 结晶度递增, 120°C 最佳 (fwhm = 0.16°)')
    add_bullet(doc, '150°C: 结晶度反而下降 (蒸发过快, 反应物来不及反应就被带走)')

    add_p(doc, '注意: 即使在室温得到的无定形粉末, 再用 120°C 加热 48 h 可逆转为结晶膜 — 验证 120°C 是结晶的热力学最优温度。')

    add_h2(doc, '4.6 反应时间 (2 天 = 48 h)')

    add_p(doc, '侯老师监测了 1 / 6 / 12 / 24 / 48 / 72 / 96 h 的膜生长过程:')
    add_bullet(doc, '1 h: 形成大量寡聚物 (MS 检出 V m/z 1073.30, VI m/z 1300.25), 产率 22.28%')
    add_bullet(doc, '6 h: 苯胺浓度降到最低 (消耗最快), 之后缓慢回升')
    add_bullet(doc, '12 h: BD-CF3 浓度达平衡')
    add_bullet(doc, '24 h: 寡聚物逐渐消失, COF 产率 87.55%')
    add_bullet(doc, '48 h: PXRD 信号最强, 结晶度最优, 产率 88.36% — ★ 反应终点')
    add_bullet(doc, '72 h 之后: PXRD/FT-IR/SEM 几乎不变 (成核已完成, 结晶度已稳定)')

    # ===== 第五章: D 系新方案 =====
    add_h1(doc, '第五章 D 系新方案 (TFPT + TFMB)')

    add_p(doc, 'D 系 (TFPT + TFMB) 与侯老师 AMCOF-1 体系完全一致 (TFPT + BD-CF3, BD-CF3 = TFMB)。本方案设计 3 个对照, 严格回归侯老师原文参数。')

    add_h2(doc, '5.1 D-1: 严格回归 AMCOF-1 原文 (推荐先跑这个)')

    add_h3(doc, '设计依据')
    add_p(doc, '侯老师 AMCOF-1 (Scheme S1) 用 TFPT + BD-CF3 在标准参数下达到 88.36% 产率。本方案 D-1 完全照搬此参数, 验证 TFMB (= BD-CF3) 在用户实验条件下也能重现。')

    add_h3(doc, '化学计量核算')
    add_bullet(doc, 'TFPT: 11.8 mg (0.03 mmol), 含醛基 0.09 mmol')
    add_bullet(doc, 'TFMB: 14.4 mg (0.045 mmol), 含氨基 0.09 mmol (与醛基等量)')
    add_bullet(doc, '苯胺: 8.2 μL (0.09 mmol, 3 eq 相对醛基)')
    add_bullet(doc, '乙酸: 0.20 mL × 6 M = 1.2 mmol')
    add_bullet(doc, '溶剂: 甲苯 0.6 mL + 氯仿 0.4 mL + 氯仿 1.0 mL (溶解 TFMB 用) = 总 2.0 mL')

    add_h3(doc, '详细操作步骤')

    add_p(doc, '步骤 1 【Pyrex 管准备】取 35 mL 派热克斯玻璃管, 依次用丙酮、乙醇、去离子水各超声清洗 10 分钟, 120°C 烘箱干燥 30 分钟备用。【目的: 去除表面有机污染物, 暴露极性 -OH 基团, 为后续氢键预排列做准备】')

    add_p(doc, '步骤 2 【加醛基单体】向 Pyrex 管中依次加入:')
    add_p(doc, '   • TFPT (1,3,5-三(4-甲酰基苯基)-1,3,5-三嗪): 11.8 mg (0.03 mmol)')
    add_p(doc, '   • 甲苯: 0.6 mL')
    add_p(doc, '   • 氯仿: 0.4 mL')
    add_p(doc, '   • 苯胺 (C₆H₅NH₂): 8.2 μL')
    add_p(doc, '【目的: TFPT 是三醛节点; 甲苯润湿玻璃壁 (铺展液层), 氯仿 120°C 下蒸发夹带单体向上; 苯胺优先与 TFPT 醛基反应形成希夫碱中间体 (m/z 469.11, MS 验证)】')

    add_p(doc, '步骤 3 【加催化剂和第二单体】沿管壁缓慢加入:')
    add_p(doc, '   • 6.0 M 乙酸水溶液: 0.20 mL')
    add_p(doc, '   • TFMB (2,2\'-双(三氟甲基)-4,4\'-联苯二胺): 14.4 mg (0.045 mmol), 预先溶于氯仿 1.0 mL')
    add_p(doc, '【目的: 乙酸 6 M 是最佳催化浓度 (与苯胺 pKa 匹配); TFMB 是二胺连接臂, 溶于氯仿保证加入时立即分散; 加料顺序——先乙酸后 TFMB——确保催化剂先与希夫碱中间体接触, 调控交换动力学】')

    add_p(doc, '步骤 4 【超声混合】将混合液超声处理 10 分钟, 至完全溶解得到澄清 (或微浊) 溶液。【目的: 超声促进单体在有机相中均匀分散, 避免局部浓度过高引起均相沉淀】')

    add_p(doc, '步骤 5 【加热反应】将 Pyrex 管用聚四氟乙烯内衬螺旋盖密封, 置于 120°C 油浴中静置反应 2 天 (48 h)。【目的: 120°C 高于所有溶剂沸点, 持续蒸发 + 夹带 → 玻璃壁超饱和液层 → 限域微反应器; 2 天是侯老师监测到的最优反应终点 (48 h 后 PXRD/FT-IR/SEM 几乎不变)】')

    add_p(doc, '步骤 6 【膜收集】反应结束后, 冷却至室温, 用金属刮刀小心将玻璃壁上的黄色膜剥离。【目的: 自然冷却避免热应力破坏膜; 玻璃壁上的膜容易剥离, 因为玻璃与膜之间是物理吸附而非化学键】')

    add_p(doc, '步骤 7 【洗涤】依次用丙酮和四氢呋喃 (THF) 各洗涤 3 次, 去除未反应单体和低聚物。【目的: 丙酮溶解有机杂质, THF 溶解 COF 低聚物; 三次洗涤保证纯度】')

    add_p(doc, '步骤 8 【干燥与称量】室温自然干燥, 称量膜质量, 计算分离产率。预期产率: ~88% (对照 AMCOF-1)')

    add_h3(doc, '预期结果与失败排查')

    add_table(doc,
              ['预测现象', '原因分析', '修正方法'],
              [
                  ['预期: 黄色连续膜, 产率 ~88%',
                   '严格按侯老师参数 → 重现 AMCOF-1',
                   '—'],
                  ['膜不连续, 玻璃上散在颗粒',
                   '可能: TFMB 未完全溶解 → 加料前超声 TFMB+氯仿',
                   '增加超声时间至 20 min'],
                  ['玻璃上有粉末, 体相浑浊',
                   '【乙酸过量】或【量放大】破坏可逆性',
                   '【严格 6 M 乙酸】+【严禁放量】'],
                  ['膜黄变深 / 碳化',
                   '【温度过高】或【反应时间过长】',
                   '温度 120°C 严格, 时间 48 h 不超'],
                  ['未冷却开盖 → 爆沸 (D4 类操作错)',
                   '违反温度安全规程',
                   '冷却至 30-40°C 再开盖, 用 PTFE 盖而非螺纹盖'],
              ],
              col_widths=[Cm(4.5), Cm(5.5), Cm(5.0)])

    add_p(doc, '表 8 D-1 预期结果与失败排查')

    add_h2(doc, '5.2 D-2: 借鉴 AMCOF-3 (PDA-CF3 路线)')

    add_h3(doc, '设计依据')
    add_p(doc, '侯老师 AMCOF-3 (Scheme S3) 用了不同的二胺 PDA-CF3 (1.0 eq 相对 TFPT 醛基), 用 2.0 mL 氯仿溶解二胺。本方案 D-2 用同样的二胺 + TFMB (结构相似, 苯环 vs 联苯), 但溶剂体系不同: 氯仿 2.0 mL 一次性加, 不分两批。')

    add_h3(doc, '关键差异 (与 D-1 对比)')
    add_bullet(doc, 'D-1: TFMB 用 1.0 mL 氯仿溶解, 加到 1.0 mL 体系中')
    add_bullet(doc, 'D-2: TFMB 用 2.0 mL 氯仿溶解, 直接形成 3.0 mL 体系 (更多氯仿 → 更快夹带 → 更快结晶)')

    add_h3(doc, '详细操作')
    add_p(doc, '步骤 1-1 同 D-1。')
    add_p(doc, '步骤 2 同 D-1, 但 TFPT+甲苯+氯仿+苯胺总体积 1.0 mL。')
    add_p(doc, '步骤 3: 沿管壁加入 6.0 M 乙酸 0.20 mL, 然后加入 TFMB (14.4 mg, 0.045 mmol) 溶于氯仿 2.0 mL。【注意: 此处氯仿总体积 3.0 mL, 比 D-1 多 1.0 mL, 强化夹带作用】')
    add_p(doc, '步骤 4-8 同 D-1。')

    add_h2(doc, '5.3 D-3: 改良 (用三氟甲苯 BTF 替代部分甲苯)')

    add_h3(doc, '设计依据')
    add_p(doc, 'A6 是含氟二醛 (含 -CF₃ 基团), 与含氟 BTF (沸点 102°C, 含 -CF₃) 有氟-氟亲和性。本方案 D-3 用 BTF 替代部分甲苯, 增强含氟单体的溶解性。但 BTF 沸点 102°C 接近反应温度 120°C, 夹带作用比氯仿弱 → 必须保留氯仿作为主要夹带剂。')

    add_h3(doc, '关键差异')
    add_bullet(doc, 'D-1/D-2: 甲苯/氯仿 = 3/7 (v/v), 0.6 mL 甲苯 + 0.4 mL 氯仿')
    add_bullet(doc, 'D-3: 甲苯/BTF/氯仿 = 1/1/2 (v/v/v), 0.2 mL 甲苯 + 0.2 mL BTF + 0.6 mL 氯仿')

    add_h3(doc, '详细操作')
    add_p(doc, '步骤 1: Pyrex 管清洗同 D-1。')
    add_p(doc, '步骤 2: 向管中加入 TFPT (11.8 mg, 0.03 mmol) + 甲苯 (0.2 mL) + BTF (0.2 mL) + 氯仿 (0.6 mL) + 苯胺 (8.2 μL)。【目的: BTF 含 -CF₃ 与 A6/TFMB 亲氟, 增强溶解】')
    add_p(doc, '步骤 3: 加 6.0 M 乙酸 0.20 mL + TFMB 14.4 mg 溶于氯仿 1.0 mL。')
    add_p(doc, '步骤 4-8 同 D-1。')

    add_h2(doc, '5.4 D 系对照实验汇总')

    add_table(doc,
              ['编号', '醛', '胺', '溶剂', '苯胺', '乙酸', '时间', '核心差异'],
              [
                  ['D-1', 'TFPT', 'TFMB', '甲苯/氯仿 3/7', '8.2 μL (3 eq)', '6 M × 0.20 mL', '48 h', '★ 严格 AMCOF-1 原文'],
                  ['D-2', 'TFPT', 'TFMB', '甲苯/氯仿 1/3 (加氯仿)', '8.2 μL (3 eq)', '6 M × 0.20 mL', '48 h', '借鉴 AMCOF-3, 氯仿多'],
                  ['D-3', 'TFPT', 'TFMB', '甲苯/BTF/氯仿 1/1/2', '8.2 μL (3 eq)', '6 M × 0.20 mL', '48 h', 'BTF 强化含氟单体溶解'],
              ],
              col_widths=[Cm(1.0), Cm(1.5), Cm(1.5), Cm(3.5), Cm(2.0), Cm(2.5), Cm(1.5), Cm(4.0)])

    add_p(doc, '表 9 D 系 3 个对照方案汇总')

    # ===== 第六章: A1 系新方案 =====
    add_h1(doc, '第六章 A1 系新方案 (TAPT + A6)')

    add_p(doc, 'A1 系 (TAPT + A6) 不在侯老师 13 个 AMCOF 体系中 (侯老师全用 TFPT 或其他三醛做节点, 没用 TAPT 三胺节点), 但本方案推断其机理应一致: 固-液/气三相界面 + 扩散/调制剂双介导。基于以下 3 个依据:')
    add_bullet(doc, 'AMCOF-1 ~ AMCOF-13 涵盖不同节点 (TFPT, Tp, TPA, TPT, TFFT, TTFT) 和不同连接臂, 证明机理通用')
    add_bullet(doc, 'TAPT 与 TFPT 结构相似 (都是 1,3,5-三嗪核 + 三个苯环臂), 物性相近')
    add_bullet(doc, 'A6 是二醛 (不是三醛), 与 TAPT 三胺形成 hcb 拓扑 COF — 节点对换不改变机理')

    add_h2(doc, '6.1 A1-1: 借鉴 AMCOF-1 但用 TAPT+A6')

    add_h3(doc, '化学计量核算')
    add_bullet(doc, 'TAPT: 10.6 mg (0.03 mmol), 含氨基 0.09 mmol')
    add_bullet(doc, 'A6: 19.0 mg (0.045 mmol), 含醛基 0.09 mmol (与氨基等量)')
    add_bullet(doc, '苯胺: 8.2 μL (0.09 mmol, 3 eq 相对氨基)')
    add_bullet(doc, '乙酸: 0.20 mL × 6 M')
    add_bullet(doc, '溶剂: 甲苯 0.6 mL + 氯仿 0.4 mL (调节点用) + 氯仿 1.5 mL (溶解 A6 用)')

    add_callout(doc, '关键差异 (与 D-1 对比)',
                'A1-1 把 D-1 中的【醛+苯胺】换成【胺+苯胺】——这是因为 TAPT 是胺, 苯胺必须优先与 A6 (醛) 反应形成希夫碱中间体。所以加料序调整为: (1) A6+甲苯+氯仿+苯胺 → (2) 乙酸 → (3) TAPT+氯仿。',
                color_hex='FFF8E1')

    add_h3(doc, '详细操作')
    add_p(doc, '步骤 1-1 同 D-1。')
    add_p(doc, '步骤 2: 加 A6 (19.0 mg, 0.045 mmol) + 甲苯 (0.6 mL) + 氯仿 (0.4 mL) + 苯胺 (8.2 μL)。')
    add_p(doc, '步骤 3: 加 6.0 M 乙酸 0.20 mL + TAPT (10.6 mg, 0.03 mmol) 溶于氯仿 1.5 mL。')
    add_p(doc, '步骤 4-8 同 D-1, 反应 48 h。')

    add_h3(doc, '预期结果与失败排查')

    add_table(doc,
              ['预测现象', '原因分析', '修正方法'],
              [
                  ['预期: 黄色连续膜 (TAPT+A6 概率 0.819)',
                   '严格按侯老师参数 + TAPT/A6 化学计量',
                   '—'],
                  ['膜不连续, 粗糙',
                   'A6 三联苯骨架大位阻 → 苯胺-醛中间体形成慢',
                   '改用 D-2 (更多氯仿) 或延长至 72 h'],
                  ['膜黄变深',
                   '反应时间过长或温度过高',
                   '严格 48 h, 120°C'],
                  ['底部有粉末, 壁上无膜',
                   '体相均相沉淀 (夹带不足)',
                   '增加氯仿比例 (A1-2 用甲苯/甲醇)'],
              ],
              col_widths=[Cm(4.5), Cm(5.5), Cm(5.0)])

    add_p(doc, '表 10 A1-1 预期结果与失败排查')

    add_h2(doc, '6.2 A1-2: 借鉴 AMCOF-2 (甲苯/甲醇体系, 3 M 乙酸, 3 天)')

    add_h3(doc, '设计依据')
    add_p(doc, 'AMCOF-2 (TFPT + BD-CH3) 用甲苯/甲醇 = 4/6 + 3 M 乙酸 + 3 天, 产率 65.90%。A6 比 BD-CH3 位阻大, 用更慢的反应条件 (3 M 乙酸 + 3 天) 可能更适合 A6 的缓慢有序聚合。')

    add_h3(doc, '关键差异 (与 A1-1 对比)')
    add_bullet(doc, 'A1-1: 甲苯/氯仿 = 3/7, 6 M 乙酸, 2 天')
    add_bullet(doc, 'A1-2: 甲苯/甲醇 = 4/6, 3 M 乙酸, 3 天')

    add_h3(doc, '详细操作')
    add_p(doc, '步骤 1-1 同 D-1。')
    add_p(doc, '步骤 2: 加 A6 (19.0 mg, 0.045 mmol) + 甲苯 (0.4 mL) + 甲醇 (0.6 mL) + 苯胺 (8.2 μL)。【甲醇极性强, 含 -OH, 与苯胺竞争反应, 因此乙酸降至 3 M 平衡】')
    add_p(doc, '步骤 3: 加 3.0 M 乙酸 0.20 mL + TAPT (10.6 mg, 0.03 mmol) 溶于甲醇 2.0 mL。')
    add_p(doc, '步骤 4-5 同 D-1。')
    add_p(doc, '步骤 6: 120°C 油浴 3 天 (72 h) — 延长反应时间, 因为 3 M 乙酸催化更慢。')
    add_p(doc, '步骤 7-8 同 D-1。')

    add_h2(doc, '6.3 A1-3: 用三氟甲苯强化溶解 (TAPT + A6 + BTF)')

    add_h3(doc, '设计依据')
    add_p(doc, 'A6 含 -CF₃, 与 BTF (含 -CF₃) 有强亲氟性。但 BTF 沸点 102°C, 夹带作用弱 → 仍需氯仿辅助。本方案 A1-3 用甲苯/BTF/氯仿三元混合, 兼顾溶解性与夹带。')

    add_h3(doc, '详细操作')
    add_p(doc, '步骤 2: 加 A6 (19.0 mg, 0.045 mmol) + 甲苯 (0.2 mL) + BTF (0.3 mL) + 氯仿 (0.5 mL) + 苯胺 (8.2 μL)。')
    add_p(doc, '步骤 3: 加 6.0 M 乙酸 0.20 mL + TAPT (10.6 mg, 0.03 mmol) 溶于氯仿 1.5 mL。')
    add_p(doc, '步骤 4-6: 120°C 油浴 3 天 (因 A6 位阻大)。')

    add_h2(doc, '6.4 A1 系对照实验汇总')

    add_table(doc,
              ['编号', '醛', '胺', '溶剂', '苯胺', '乙酸', '时间', '核心差异'],
              [
                  ['A1-1', 'A6', 'TAPT', '甲苯/氯仿 3/7', '8.2 μL (3 eq)', '6 M × 0.20 mL', '48 h', '★ 借鉴 AMCOF-1'],
                  ['A1-2', 'A6', 'TAPT', '甲苯/甲醇 4/6', '8.2 μL (3 eq)', '3 M × 0.20 mL', '72 h', '借鉴 AMCOF-2 (慢反应)'],
                  ['A1-3', 'A6', 'TAPT', '甲苯/BTF/氯仿 2/3/5', '8.2 μL (3 eq)', '6 M × 0.20 mL', '72 h', 'BTF 强化含氟溶解'],
              ],
              col_widths=[Cm(1.0), Cm(1.5), Cm(1.5), Cm(3.5), Cm(2.0), Cm(2.5), Cm(1.5), Cm(4.0)])

    add_p(doc, '表 11 A1 系 3 个对照方案汇总')

    # ===== 第七章: 加料顺序的化学解释 =====
    add_h1(doc, '第七章 加料顺序的化学逻辑 (每步为什么)')

    add_p(doc, '本方案的加料顺序严格按侯老师 Scheme S1-S3 原文, 每步都有明确的化学目的。下表解释每一步的"为什么"。')

    add_table(doc,
              ['步骤', '操作', '化学目的'],
              [
                  ['1', 'Pyrex 管清洗',
                   '去除表面有机污染物; 暴露极性 -OH 基团, 为后续氢键预排列做准备 (DFT 证实)'],
                  ['2a', '加醛/胺单体',
                   '单体先进入反应体系'],
                  ['2b', '加溶剂 (甲苯/氯仿 3/7)',
                   '甲苯润湿玻璃壁 (铺展液层); 氯仿 120°C 下蒸发夹带单体向上'],
                  ['2c', '加苯胺 (3 eq)',
                   '★ 苯胺优先与醛基反应, 形成苯胺-醛希夫碱中间体 (m/z 469.11) — MS 验证'],
                  ['3a', '加乙酸 (6 M)',
                   '催化亚胺形成与交换; 6 M 与苯胺 pKa (4.6) 匹配, 维持可逆性'],
                  ['3b', '加第二单体',
                   '★ 与苯胺-醛中间体缓慢交换 (m/z 846.29 中间体), 形成有序 COF'],
                  ['4', '超声 10 min',
                   '均匀分散, 避免局部浓度过高'],
                  ['5', '120°C × 48 h',
                   '★ 持续蒸发 + 夹带 → 玻璃壁超饱和液层 → 限域微反应器'],
                  ['6', '刮刀剥离',
                   '膜与玻璃壁是物理吸附 (非化学键), 容易剥离'],
                  ['7', '丙酮 + THF 各洗 3 次',
                   '丙酮溶解有机杂质; THF 溶解低聚物; 三次保证纯度'],
                  ['8', '室温干燥',
                   '低温避免膜氧化或结构变化'],
              ],
              col_widths=[Cm(1.5), Cm(5.5), Cm(8.5)])

    add_p(doc, '表 12 通用 11 步流程的化学解释')

    # ===== 第八章: 失败排查 =====
    add_h1(doc, '第八章 失败排查与修正 (基于侯老师数据)')

    add_h2(doc, '8.1 通用排查')

    add_h3(doc, '症状 1: 未见膜形成, 仅浑浊溶液或沉淀')
    add_p(doc, '【原因】(a) 单体未完全溶解即加催化剂, 析出无定形聚合物; (b) 乙酸浓度过低或温度不达标; (c) 加料顺序错——先加第二单体再加苯胺, 破坏苯胺优先反应机制。')
    add_p(doc, '【修正】(a) 加料前先在 60-80°C 超声 10 min 完全溶解; (b) 严格 6 M 乙酸 + 120°C; (c) 严格按"醛+苯胺 → 乙酸 → 第二单体"顺序。')

    add_h3(doc, '症状 2: 膜易碎、不连续或局部卷曲')
    add_p(doc, '【原因】(a) 催化剂量过大导致反应过快, 缺陷累积; (b) 反应时间不足 (< 48 h); (c) 苯胺量错 (0/1/2 eq 或 4/5 eq 都差, 必须 3 eq)。')
    add_p(doc, '【修正】(a) 严格 6 M 乙酸, 不加过量; (b) 延长至 48-72 h; (c) 严格 8.2 μL 苯胺 (3 eq)。')

    add_h3(doc, '症状 3: 膜过厚、不透明或深黄/棕色')
    add_p(doc, '【原因】(a) 反应温度过高 (>150°C) 或时间过长 (>72 h); (b) 单体氧化 (含 CF₃/CF₂ 链的 A6 易氧化变色)。')
    add_p(doc, '【修正】(a) 严格 120°C, 48 h; (b) 反应前对有机相鼓氮气 15 min, 全程氮气保护。')

    add_h3(doc, '症状 4: PXRD 无衍射峰 (非晶)')
    add_p(doc, '【原因】(a) 反应速率过快或过慢均不利结晶; (b) 苯胺量错 (0/1/2/4/5 eq 都差); (c) 溶剂体系错 (用了二氧六环/均三甲苯等高沸点, 夹带失效)。')
    add_p(doc, '【修正】(a) 严格 6 M 乙酸 + 120°C + 48 h; (b) 严格 8.2 μL 苯胺; (c) 严格甲苯/氯仿 3/7。')

    add_h3(doc, '症状 5: 未冷却开盖 → 氯仿爆沸 (D4 类)')
    add_p(doc, '【原因】违反温度安全规程, 120°C 时管内压力大, 突然开盖氯仿 (bp 61.2°C) 瞬间沸腾。')
    add_p(doc, '【修正】冷却至 30-40°C 再开盖。用 PTFE 内衬盖而非普通玻璃塞。绝不要直接开盖!')

    add_h2(doc, '8.2 D 系特异排查')

    add_h3(doc, '症状: TFMB 加料后立即浑浊')
    add_p(doc, '【原因】TFMB 在氯仿中溶解度 + (5-20 mg/mL), 14.4 mg / 1.0 mL 已接近饱和; 若溶解不彻底会析出。')
    add_p(doc, '【修正】加料前超声 TFMB + 氯仿 10 min 完全溶解后再加入。')

    add_h2(doc, '8.3 A1 系特异排查')

    add_h3(doc, '症状: A6 在体系中析出 (颗粒物)')
    add_p(doc, '【原因】A6 在甲苯/氯仿中溶解度好 (+ ~ ++), 但 A6 三联苯骨架大, 19 mg / 1.0 mL 接近溶解度上限。')
    add_p(doc, '【修正】改用 A1-3 (甲苯/BTF/氯仿三元), BTF 含氟与 A6 亲氟, 显著提升溶解度。')

    add_h3(doc, '症状: TAPT 在甲醇中不溶解 (A1-2)')
    add_p(doc, '【原因】TAPT 在甲醇中溶解度 - (几乎不溶)。')
    add_p(doc, '【修正】A1-2 把 TAPT 溶于【甲醇 2.0 mL + 0.20 mL 6 M 乙酸水溶液】, 酸化的甲醇微溶 TAPT; 或改用 A1-1 / A1-3 (用氯仿溶解 TAPT)。')

    # ===== 第九章: 表征方案 =====
    add_h1(doc, '第九章 表征方案')

    add_table(doc,
              ['表征', '目的', '关键判定', '数据对比'],
              [
                  ['PXRD (粉末 X 射线衍射)',
                   '评估结晶度',
                   'fwhm(100) ≤ 0.16° 为高结晶 (对照 AMCOF-1)',
                   'AMCOF-1: 2θ = 2.38°, 4.14°, ...'],
                  ['FT-IR (红外光谱)',
                   '确认亚胺键形成',
                   'C=N 峰 ~1631 cm⁻¹ (出现); C=O 峰 ~1708 cm⁻¹ (消失)',
                   'AMCOF-1: 1631 cm⁻¹ 亚胺特征峰'],
                  ['固体 ¹³C NMR',
                   '确认结构',
                   '亚胺碳信号 ~159 ppm',
                   'AMCOF-1: 159 ppm'],
                  ['SEM (扫描电镜)',
                   '膜表面形貌',
                   '气相侧: 立方/六棱柱晶体; 玻璃侧: 光滑',
                   'AMCOF-1: 棱面晶, 厚 0.96-21.7 μm'],
                  ['AFM (原子力显微镜)',
                   '膜厚与粗糙度',
                   '膜厚均匀, 表面 Ra < 5 nm',
                   'AMCOF-1: 4.60 μm 厚, Ra 较低'],
                  ['BET (N₂ 吸附, 77 K)',
                   '比表面积 + 孔径',
                   'BET > 500 m²/g, 孔径 ~3.9 nm (对照 AA 堆积)',
                   'AMCOF-1: 594.87 m²/g, 3.93 nm 孔径'],
                  ['XPS',
                   '元素组成 + 化学态',
                   'C, N, F 峰对应 (含氟膜应见 F 1s)',
                   'AMCOF-1: 含 F, C, N'],
                  ['接触角',
                   '膜疏水性',
                   '含氟膜 > 100°',
                   'A6 / TFMB 应 > 110°'],
              ],
              col_widths=[Cm(3.0), Cm(3.0), Cm(5.0), Cm(5.0)])

    add_p(doc, '表 13 本方案推荐表征项目 (引自侯老师 AMCOF-1 表征方法)')

    # ===== 第十章: 参考文献 =====
    add_h1(doc, '第十章 参考文献')

    add_h2(doc, '10.1 核心文献 (本方案科学基础)')

    add_p(doc, '【1】侯盛怀, 张干兵, 乔钊宇, 白玉轩, 邸昊昕, 华烨同, 郝甜甜, 徐慧* "扩散/调制剂双介导固-液/气界面合成结晶型共价有机框架膜"《德国应用化学》 2025, 64, e202421555.')
    add_p(doc, '   DOI: 10.1002/anie.202421555', indent=False)
    add_p(doc, '   本地路径: 实验\\文章\\侯老师实验\\侯盛怀德国应化.pdf', indent=False)
    add_p(doc, '   补充信息: 实验\\文章\\侯老师实验\\sl-884-anie202421555-sup-0001-misc_information.pdf', indent=False)
    add_p(doc, '   提取文本: 实验\\文章\\侯老师实验\\_main_full.txt + _si.txt', indent=False)

    add_h2(doc, '10.2 GraphRAG v2 检索到的相关文献')

    add_p(doc, '【2】L-0ca3cbf692f8 (Science Advances). 界面合成制备高结晶度 2D COF 膜 + 溶剂响应. DOI 待补充.')

    add_p(doc, '【3】L-c8cc6ff417dd (J. Am. Chem. Soc.). SI-SBMAP 方法, sp²c-COF 薄膜. DOI 待补充.')

    add_p(doc, '【4】L-73ca3ac9b632 (ACS Appl. Mater. Interfaces). TpTFMB 2D COF 体系, 原位生长在石英毛细管内壁. DOI 待补充.')

    add_p(doc, '【5】L-a35edf0e281d (J. Am. Chem. Soc.). Boc 保护 + TFA 脱保护策略, 编织结构 COF-112. DOI 待补充.')

    add_p(doc, '【6】R-101-1 (GraphRAG). Tp + Pa, water/CH₂Cl₂, 室温 25°C, film. (仅作为液-液界面法对照参考, 不属于本方案主线)')

    add_p(doc, '【7】R-563-6 (GraphRAG). Tp + Tta, DCM/water 液-液界面, 室温, film. (同上)')

    add_p(doc, '【8】R-161-1 (GraphRAG). TFTA + TAPT, mesitylene:1,4-dioxane (1:1), 120°C, film, 异相 (Fe₃O₄ 磁核). DOI 待补充.')

    add_p(doc, '【9】R-60-6 (GraphRAG). FPDA + TAPT, 1,4-dioxane:mesitylene (1:1), 120°C, film. DOI 待补充.')

    add_p(doc, '【10】R-62 系列 (GraphRAG). Tp + Boc 保护胺, dioxane/water, 120°C, film (Boc 保护策略). DOI 待补充.')

    add_h2(doc, '10.3 本方案相关本地文献')

    add_p(doc, '【11】侯盛怀等, "固-液/气三相界面合成 COF 膜" 实验数据. 实验\\文章\\侯老师实验\\侯盛怀德国应化.pdf')

    add_p(doc, '【12】tianxuan-seek 训练数据集 (含 954 篇 COF 文献结构化抽取). C:\\Users\\ckx\\Desktop\\tianxuan seek\\data\\structured\\')

    add_p(doc, '【13】实验 ABCDEF 反馈记录 (14 条). C:\\Users\\ckx\\Desktop\\实验\\方案\\实验ABCDEF.docx')

    add_p(doc, '【14】v3.9 实验方案 (本方案的上游模板). C:\\Users\\ckx\\Desktop\\实验\\方案\\实验方案_含氟COF薄膜合成_v3.9_20260626.docx')

    add_h2(doc, '10.4 关键数据点 (本方案涉及的化学反应常数)')

    add_table(doc,
              ['参数', '数值', '来源'],
              [
                  ['TFPT 分子量', '393.40 g/mol', 'PubChem CID 102335589'],
                  ['TAPT 分子量', '354.41 g/mol', 'PubChem CID 1515256'],
                  ['TFMB / BD-CF3 分子量', '320.23 g/mol', 'D&B Biotechnology 货品规格'],
                  ['A6 分子量', '422.37 g/mol', 'v3.9 方案 §2.2'],
                  ['苯胺密度', '1.022 g/mL', 'Sigma-Aldrich'],
                  ['苯胺摩尔质量', '93.13 g/mol', '—'],
                  ['甲苯沸点', '110.6 °C', 'depts.washington.edu/eooptic'],
                  ['氯仿沸点', '61.2 °C', '同上'],
                  ['甲醇沸点', '64.7 °C', '同上'],
                  ['乙酸 pKa', '4.76', '—'],
                  ['苯胺 pKa', '4.60 (共轭酸)', '—'],
              ],
              col_widths=[Cm(5.0), Cm(5.0), Cm(5.5)])

    add_p(doc, '表 14 本方案涉及的关键物性数据')

    # ===== 第十一章: 实验准备清单 =====
    add_h1(doc, '第十一章 实验准备清单')

    add_h2(doc, '11.1 单体 (4 种, 全部已买)')

    add_table(doc,
              ['名称', 'CAS', '来源', '用量 (D-1)', '用量 (A1-1)'],
              [
                  ['TFPT', '443922-06-3', 'D&B Biotechnology', '11.8 mg', '—'],
                  ['TAPT', '14544-47-9', '—', '—', '10.6 mg'],
                  ['TFMB (BD-CF3)', '341-58-2', 'D&B Biotechnology', '14.4 mg', '—'],
                  ['A6', '1300701-03-4', '—', '—', '19.0 mg'],
                  ['苯胺 (C₆H₅NH₂)', '62-53-3', 'Sinopharm (国药)', '8.2 μL', '8.2 μL'],
                  ['甲苯', '108-88-3', 'Sinopharm', '0.6 mL', '0.6 mL'],
                  ['氯仿', '67-66-3', 'Kelong (科龙)', '0.4 + 1.0 = 1.4 mL', '0.4 + 1.5 = 1.9 mL'],
                  ['甲醇', '67-56-1', 'Sinopharm', '— (D-2/D-3)', '0.6 mL (A1-2)'],
                  ['三氟甲苯 (BTF)', '98-08-8', 'Aladdin (阿拉丁)', '— (D-3)', '— (A1-3)'],
                  ['乙酸 (冰醋酸)', '64-19-7', 'Sinopharm', '0.20 mL × 6 M', '0.20 mL × 6 M'],
                  ['去离子水', '—', 'Milli-Q', '配 6 M 乙酸', '配 6 M 乙酸'],
                  ['丙酮 (洗涤)', '67-64-1', '—', '5 mL × 3', '5 mL × 3'],
                  ['四氢呋喃 THF (洗涤)', '109-99-9', '—', '5 mL × 3', '5 mL × 3'],
              ],
              col_widths=[Cm(3.5), Cm(2.5), Cm(3.0), Cm(2.5), Cm(2.5)])

    add_p(doc, '表 15 实验试剂用量清单 (D-1 和 A1-1 为例)')

    add_h2(doc, '11.2 设备')

    add_bullet(doc, '35 mL 派热克斯玻璃管 (Pyrex tube) — 侯老师指定, 极性玻璃表面是关键')
    add_bullet(doc, '聚四氟乙烯 (PTFE) 内衬螺旋盖 — 防止氯仿爆沸')
    add_bullet(doc, '120°C 油浴 (带磁力搅拌或不搅拌均可) — 侯老师用沙浴或油浴')
    add_bullet(doc, '超声清洗仪 — 用于单体溶解和玻璃管清洗')
    add_bullet(doc, '金属刮刀 — 用于剥离膜')
    add_bullet(doc, '微量注射器 (10 μL / 100 μL / 1 mL) — 精确加苯胺和乙酸')
    add_bullet(doc, '分析天平 (0.1 mg 精度) — 称量单体')

    # ===== 第十二章: 附录 =====
    add_h1(doc, '第十二章 附录')

    add_h2(doc, '12.1 与 v3.9 通用步骤对照')

    add_table(doc,
              ['步骤', 'v3.9 通用', '本方案 v7', '改进理由'],
              [
                  ['1. 管准备', '丙酮/乙醇/水超声', '同', '—'],
                  ['2. 加醛/胺单体', '醛+苯胺→立即胺→最后酸', '严格按侯老师原序',
                   '侯老师 HPLC-MS 证明必须先醛+苯胺'],
                  ['3. 加乙酸', '6.0 M × 0.20 mL', '同',
                   '6 M 是侯老师最佳浓度, 12/18 M 都过量'],
                  ['4. 超声', '10-15 min', '同', '—'],
                  ['5. 加热', '120°C × 48-72 h', '同', '48 h 是结晶度最优终点'],
                  ['6. 监测', '观察膜形成', '同 + 关注玻璃壁超饱和液层',
                   '液层是侯老师机理的关键中间体'],
                  ['7. 收集', '刮刀剥离', '同', '—'],
                  ['8. 洗涤', '丙酮 + THF × 3', '同', '—'],
                  ['9. 干燥', '室温 / 60°C 真空', '同', '—'],
                  ['10. 称量', '分离产率', '同', '对照 88.36% (AMCOF-1)'],
              ],
              col_widths=[Cm(2.0), Cm(5.0), Cm(5.0), Cm(4.5)])

    add_p(doc, '表 16 v7 与 v3.9 通用步骤对照')

    add_h2(doc, '12.2 生成工具信息')

    add_p(doc, '本方案生成工具: bridge/generate_v7.py (基于侯老师原文机理 + GraphRAG v2 检索 + 人工精修)')
    add_p(doc, '生成时间: 2026-07-13')
    add_p(doc, '数据来源: tianxuan-seek/data/ (954 篇 yaml + 6197 反应 + 1059 单体)')
    add_p(doc, 'GraphRAG v2 模块: nl2graph + router + multimodal + importance + community + reasoning')
    add_p(doc, 'v1-v6 已删除 (过时方案, 偏离侯老师机理)')

    add_p(doc, '')
    add_p(doc, '--- 本方案结束 ---')

    doc.save(OUT)
    print(f'✓ 已生成: {OUT}')
    print(f'  大小: {os.path.getsize(OUT)/1024:.1f} KB')


if __name__ == '__main__':
    main()