"""
bridge/generate_v8.py
=====================
v8 基于【用户的两步法创新 + 侯老师调制剂双介导机理】

用户的两步法:
  步骤 1: 三嗪节点 (TAPT/TFPT) 在液面上方玻璃管壁上自组装成基底膜
  步骤 2: 冷却后加线性连接臂 + 苯胺 + 乙酸, 反应形成 COF

v7 的错误: 回归侯老师原始一锅法, 忽略了用户的两步法创新
v8 的修正: 保留两步法, 加入侯老师的调制剂双介导 (步骤 1 也要加苯胺 + 乙酸)

设计 6 个方案:
  D-1 (推荐): 两步法 + 调制剂双介导 (步骤 1 加苯胺 + 乙酸)
  D-2 (对照): 两步法无调制剂 (用户原方案, 验证失败原因)
  D-3 (基线): 一锅法 (侯老师 AMCOF-1 原文, 作为对照)
  A1-1 ~ A1-3: 同 D 系列但用 TAPT + A6
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJ = r'C:\Users\ckx\Desktop\minimax'
OUT = os.path.join(PROJ, 'experiment', 'proposals',
                   'COF-TFPT-TAPT-2026-07-15-D-A1-v8.docx')

# ==================== 样式工具 ====================

def setup_styles(doc):
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')

    h1 = doc.styles['Heading 1']
    h1.font.name = '黑体'
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.line_spacing = 1.5
    rPr1 = h1.element.get_or_add_rPr()
    rFonts1 = rPr1.find(qn('w:rFonts'))
    if rFonts1 is None:
        rFonts1 = OxmlElement('w:rFonts')
        rPr1.append(rFonts1)
    rFonts1.set(qn('w:eastAsia'), '黑体')

    h2 = doc.styles['Heading 2']
    h2.font.name = '黑体'
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x2E, 0x5C, 0x8A)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.line_spacing = 1.5

    h3 = doc.styles['Heading 3']
    h3.font.name = '黑体'
    h3.font.size = Pt(11.5)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.line_spacing = 1.5


def set_cell_bg(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def set_para_bg(para, color_hex):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    pPr.append(shd)


def add_run(p, text, size=10.5, bold=False, italic=False, color=None, font='宋体'):
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    return run


def add_h1(doc, text):
    p = doc.add_heading('', level=1)
    add_run(p, text, size=16, bold=True, color=RGBColor(0x1F, 0x3A, 0x5F), font='黑体')
    return p


def add_h2(doc, text):
    p = doc.add_heading('', level=2)
    add_run(p, text, size=13, bold=True, color=RGBColor(0x2E, 0x5C, 0x8A), font='黑体')
    return p


def add_h3(doc, text):
    p = doc.add_heading('', level=3)
    add_run(p, text, size=11.5, bold=True, color=RGBColor(0xC0, 0x39, 0x2B), font='黑体')
    return p


def add_p(doc, text, indent=True, size=10.5, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    add_run(p, text, size=size, bold=bold, italic=italic, color=color)
    return p


def add_bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.first_line_indent = Cm(0.5)
    add_run(p, text, size=size)
    return p


def add_callout(doc, label, content, color_hex='F0F7FF'):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    set_para_bg(p, color_hex)
    add_run(p, f'【{label}】', size=10.5, bold=True, color=RGBColor(0xC0, 0x39, 0x2B))
    add_run(p, content, size=10.5)
    return p


def add_table(doc, header, rows, header_bg='1F3A5F', col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = w

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_bg(hdr_cells[i], header_bg)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        bg = 'F8F9FA' if r_idx % 2 == 0 else 'FFFFFF'
        for i, v in enumerate(row_data):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            add_run(p, str(v), size=9.5)
            set_cell_bg(cells[i], bg)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return table


def add_page_break(doc):
    doc.add_page_break()


# ==================== 主文档 ====================

def main():
    doc = Document()
    setup_styles(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # ===== 封面 =====
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(36)
    title_p.paragraph_format.space_after = Pt(12)
    add_run(title_p, '含氟亚胺键 COF 膜合成实验方案',
            size=22, bold=True, color=RGBColor(0x1F, 0x3A, 0x5F), font='黑体')

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(6)
    add_run(subtitle_p, 'D 系 + A1 系 迭代方案 v8',
            size=16, bold=True, color=RGBColor(0x2E, 0x5C, 0x8A), font='黑体')

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2.paragraph_format.space_after = Pt(36)
    add_run(sub2, '基于【用户两步法创新】+【侯老师调制剂双介导机理】',
            size=11, italic=True, color=RGBColor(0x6C, 0x75, 0x7D))

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(info_p, '方案编号: ', size=10.5, bold=True)
    add_run(info_p, 'COF-TFPT-TAPT-2026-07-15-D-A1-v8', size=10.5)
    info_p.add_run().add_break()
    add_run(info_p, '生成日期: ', size=10.5, bold=True)
    add_run(info_p, '2026-07-15', size=10.5)
    info_p.add_run().add_break()
    add_run(info_p, '迭代依据: ', size=10.5, bold=True)
    add_run(info_p, '14 条 ABCDEF + A1-D9 历史失败 (用户两步法实验)', size=10.5)
    info_p.add_run().add_break()
    add_run(info_p, '核心参考文献: ', size=10.5, bold=True)
    add_run(info_p, 'Hou et al., Angew. Chem. Int. Ed. 2025, 64, e202421555 (DOI: 10.1002/anie.202421555)',
            size=10.5, italic=True)
    info_p.add_run().add_break()
    add_run(info_p, '   本地路径: ', size=10.5)
    add_run(info_p, r'实验\文章\侯老师实验\侯盛怀德国应化.pdf', size=10, italic=True)

    doc.add_page_break()

    # ===== 摘要 =====
    add_h1(doc, '摘要')

    add_callout(doc, '关键修正',
                'v7 错误地回归侯老师原始一锅法 (醛+苯胺+胺+乙酸一次加完), 忽略了用户的两步法创新。本 v8 方案重新基于用户的真实两步法 (步骤 1: 三嗪节点在玻璃壁上自组装成基底膜 → 步骤 2: 冷却后加线性连接臂反应), 并融入侯老师的调制剂双介导机理 (步骤 1 加入苯胺+乙酸, 让基底膜稳定可逆形成而非无序自聚)。',
                color_hex='FFF8E1')

    add_p(doc, '本方案核心洞察 (基于对用户 A1/D3 实验的真实理解):')
    add_bullet(doc, '【用户的两步法是创新】步骤 1 让三嗪节点在玻璃壁上预先自组装成基底膜, 步骤 2 加线性连接臂反应形成 COF。这与侯老师一锅法不同, 是用户试图形成"易脱落基底膜"的尝试。')
    add_bullet(doc, '【失败的真原因】步骤 1 没加苯胺/乙酸 → 纯三嗪单体自聚得到的是无序聚集体 → 重启加热后反复溶解-再形成 → 不能形成稳定基底膜。同时, D3 在步骤 2 用了 12 M 乙酸, 破坏了侯老师调制剂双介导的可逆性。')
    add_bullet(doc, '【本 v8 的修正】保留两步法的核心创新 (步骤 1 先形成基底膜), 但步骤 1 加入苯胺 (3 eq) + 乙酸 (6 M), 让基底膜以"苯胺-三嗪希夫碱"形式稳定可逆形成。这样基底膜既稳定又能与步骤 2 的连接臂发生亚胺交换反应。')
    add_bullet(doc, '【6 个对照方案】D-1/D-2/D-3 + A1-1/A1-2/A1-3, 每个都明确步骤 1 是否含调制剂, 步骤 2 是否严格冷却加料。')

    # ===== 第一章: 用户两步法 vs 侯老师一锅法 =====
    add_h1(doc, '第一章 用户两步法 vs 侯老师一锅法 (核心区别)')

    add_h2(doc, '1.1 侯老师一锅法 (AMCOF-1 原文)')

    add_p(doc, '侯老师方法是所有单体同时加入, 通过苯胺调制剂和限域微反应器在玻璃壁上原位生长 COF 膜:')
    add_p(doc, '  步骤 1: TFPT + 甲苯 + 氯仿 + 苯胺 + 乙酸 + BD-CF3(溶于氯仿) → 一次性加完')
    add_p(doc, '  步骤 2: 超声 10 min')
    add_p(doc, '  步骤 3: 120°C × 48 h 油浴')
    add_p(doc, '  步骤 4: 刮刀剥离 → 丙酮/THF 洗涤 → 室温干燥')
    add_p(doc, '   【结果】: 88.36% 产率, 黄色连续膜, PXRD fwhm = 0.16°')

    add_h2(doc, '1.2 用户两步法 (基于用户 A1 实验描述)')

    add_p(doc, '用户两步法是先把三嗪节点在玻璃壁上自组装成基底膜, 再加线性连接臂反应形成 COF:')
    add_p(doc, '  步骤 1: TAPT/TFPT + 溶剂 (甲苯 + 氯仿) → 超声 → 120°C × 12 h ~ 3 天')
    add_p(doc, '          [目的: 让三嗪节点在玻璃壁上预先自组装成基底膜]')
    add_p(doc, '  步骤 2: 取出 Pyrex 管, 冷却至 30-40°C')
    add_p(doc, '  步骤 3: 加第二种单体 + 苯胺 + 乙酸')
    add_p(doc, '  步骤 4: 重新 120°C × 3 天')
    add_p(doc, '          [目的: 让第二种单体与基底膜反应形成 COF]')
    add_p(doc, '   【用户原方案结果】: A1 形成连续但粗糙膜, 机械强度极低 (一碰就碎), 失败。D3 操作失误爆沸 + 12 M 乙酸, 失败。')

    add_h2(doc, '1.3 两步法的科学问题')

    add_p(doc, '用户的两步法本质是试图解决"易脱落基底膜"问题——如果步骤 1 形成的基底膜与玻璃壁是物理吸附, 步骤 2 加的连接臂反应生成的 COF 应该能与基底膜分离, 形成独立的 COF 膜。这比侯老师一锅法更有工业应用前景 (易转移)。')

    add_p(doc, '但用户的尝试存在 3 个核心问题:')
    add_bullet(doc, '【问题 1: 步骤 1 缺调制剂】纯三嗪单体在加热下自聚成无序聚集体, 没有苯胺的"封端"作用, 形成的是不稳定的寡聚物-沉淀动态平衡, 重启加热后反复溶解-再形成 (A1 现象描述)')
    add_bullet(doc, '【问题 2: 步骤 2 加新单体扰动步骤 1 形成的膜】A1 步骤 2 加入 A6+苯胺+乙酸时, 新加的溶剂 (1.0 mL 甲苯+氯仿) 会溶解步骤 1 形成的初始膜, 导致基底膜被冲刷掉')
    add_bullet(doc, '【问题 3: D3 的步骤 2 用了 12 M 乙酸】过量乙酸破坏可逆性 → 形成沉淀而非膜 (与用户 D 系列失败原因一致)')

    add_h2(doc, '1.4 本 v8 方案的核心修正')

    add_callout(doc, '修正策略',
                '保留两步法的创新 (步骤 1 先形成基底膜), 但融入侯老师调制剂双介导 (步骤 1 加入苯胺 + 乙酸):\n'
                '【修正后的步骤 1】: 三嗪节点 + 苯胺 (3 eq) + 乙酸 (6 M) + 溶剂 → 120°C × 12-24 h\n'
                '【关键变化】苯胺优先与三嗪的醛基反应形成苯胺-三嗪希夫碱 (m/z 469.11 类似物, 侯老师 MS 验证), 这层希夫碱是稳定的可逆聚合产物, 不是无序聚集体。它在玻璃壁上形成稳定的"基底膜", 重启加热不会反复溶解-再形成。\n'
                '【修正后的步骤 2】: 冷却至 30-40°C → 加线性连接臂 → 重新 120°C × 48 h',
                color_hex='E8F5E9')

    # ===== 第二章: 修正的两步法机理 =====
    add_h1(doc, '第二章 修正的两步法机理 (含调制剂双介导)')

    add_h2(doc, '2.1 步骤 1: 三嗪节点 + 调制剂 → 形成稳定基底膜')

    add_p(doc, '【反应方程式】(以 TFPT 为例):')
    add_p(doc, '  TFPT-CHO + C₆H₅NH₂ → TFPT-CH=N-C₆H₅ (苯胺-三嗪希夫碱, m/z ≈ 469)')
    add_p(doc, '  苯胺-三嗪希夫碱 + 玻璃壁 -OH → 氢键 + π-π 堆积 → 自组装成基底膜')

    add_p(doc, '【为什么基底膜稳定】')
    add_bullet(doc, '苯胺封端了一个醛基, 形成稳定的希夫碱 (不像纯三嗪自聚那样无序)')
    add_bullet(doc, '希夫碱有共轭结构, 与玻璃壁的 -OH 形成氢键, 在玻璃壁上预排列')
    add_bullet(doc, '6 M 乙酸催化亚胺形成 + 维持可逆性 (反应不会太快生成沉淀)')
    add_bullet(doc, '12-24 h 反应时间足够形成连续的基底膜, 但不会过厚影响步骤 2')

    add_h2(doc, '2.2 步骤 2: 加线性连接臂 → 亚胺交换反应')

    add_p(doc, '【反应方程式】')
    add_p(doc, '  TFPT-CH=N-C₆H₅ + TFMB-NH₂ → TFPT-CH=N-TFMB (COF) + C₆H₅NH₂ (苯胺释放)')
    add_p(doc, '  苯胺释放的苯胺重新进入催化循环, 调控反应动力学')

    add_p(doc, '【为什么步骤 2 要冷却加料】')
    add_bullet(doc, '冷却至 30-40°C 防止氯仿爆沸 (D3 类操作错)')
    add_bullet(doc, '冷却时基底膜不会溶解 (希夫碱在室温下稳定)')
    add_bullet(doc, '新加的线性连接臂与基底膜的希夫碱交换, 在基底膜上生长 COF')

    add_callout(doc, '机理优势',
                '与用户原两步法相比, 修正方案的关键优势:\n'
                '【基底膜稳定】苯胺-三嗪希夫碱是稳定有序结构, 不会反复溶解-再形成\n'
                '【易脱落设计保留】基底膜与玻璃壁是物理吸附 (氢键), 步骤 2 形成的 COF 与基底膜是化学键, 刮刀剥离时 COF 与基底膜一起脱落, 与玻璃壁分离 → "易脱落基底膜" 设想实现\n'
                '【反应动力学可控】苯胺作为调制剂持续调控反应速率, 避免快速沉淀',
                color_hex='E8F5E9')

    # ===== 第三章: 历史失败深度诊断 =====
    add_h1(doc, '第三章 历史失败深度诊断 (基于两步法 + 调制剂双介导)')

    add_p(doc, '下表汇总 14 条反馈, 每个失败都基于修正后的两步法机理解读。')

    add_table(doc,
              ['ID', '两步法分类', '原方案步骤 1', '原方案步骤 2', '根因', 'v8 修正'],
              [
                  ['A1', '两步法 (TAPT+A6)',
                   'TAPT + 溶剂 → 120°C × 3 天 (无苯胺/乙酸)',
                   '冷却 + A6 + 苯胺 + 6M 乙酸 → 120°C × 3 天',
                   '步骤 1 无调制剂 → 基底膜不稳定, 反复溶解-再形成',
                   '步骤 1 加苯胺 + 乙酸 (D-1/A1-1 方案)'],
                  ['A2', '两步法 (TAPB+A6)',
                   'TAPB + 溶剂 → 120°C × 1 天',
                   '冷却 + A6 + 苯胺 + 乙酸',
                   'TAPB 反应性弱, 自组装更不稳定',
                   '改用 TAPT (A1-1 方案)'],
                  ['A5', '两步法 (TAPT+A6)',
                   'TAPT + 溶剂 (氯仿 2.2 mL) → 120°C × 1 天',
                   '冷却 + A6 + 苯胺 + 乙酸 (氯仿微调)',
                   '氯仿量微调无法解决根本问题',
                   '回归侯老师 SOP (A1-1 方案)'],
                  ['A8', '新加料序一锅法',
                   'A6 + 苯胺 + 6M 乙酸 + TAPT (颠倒顺序)',
                   '—',
                   '颠倒加料序破坏苯胺-醛优先反应机制',
                   '回归"醛+苯胺→乙酸→第二单体"'],
                  ['D', '一锅法',
                   'TFPT + TFMB + 苯胺 + 18M 乙酸',
                   '—',
                   '【乙酸 18M 过量】破坏可逆性',
                   '严格 6M 乙酸 (D-1)'],
                  ['D3', '两步法 (TFPT+TFMB)',
                   'TFPT + 溶剂 → 120°C × 3 天 (无苯胺/乙酸)',
                   '冷却 + TFMB + 苯胺 + 12M 乙酸 → 120°C × 3 天',
                   '步骤 1 无调制剂 + 步骤 2 乙酸过量 + 操作爆沸',
                   '步骤 1 加苯胺+乙酸, 步骤 2 用 6M 乙酸 (D-1)'],
                  ['D4', '两步法 (TFPB+TFMB)',
                   'TFPB + 溶剂 → 120°C × 1 天',
                   '未冷却开盖 → 氯仿爆沸 → 数据作废',
                   '操作错: 违反温度安全规程',
                   '【废除】TFPB 节点, 用 TFPT'],
                  ['D7', '两步法 (TFPT+TFMB)',
                   'TFPT + 溶剂 → 120°C × 1 天 (无苯胺/乙酸)',
                   '冷却 + TFMB + 苯胺 + 12M 乙酸',
                   '与 D3 类似',
                   '严格按 v8 D-1'],
                  ['D9', '放大 3 倍两步法',
                   'TFPT (3x) + 溶剂 → 120°C × 16h',
                   '新序 + TFMB (3x) + 苯胺 + 乙酸',
                   '量放大破坏夹带动力学 → 碳化',
                   '【严禁放大】按侯老师原文量级'],
              ],
              col_widths=[Cm(1.0), Cm(2.5), Cm(4.5), Cm(3.5), Cm(3.5), Cm(2.5)])

    add_p(doc, '表 1 14 条失败的两步法机理解读')

    # ===== 第四章: 化学计量与单体量 =====
    add_h1(doc, '第四章 化学计量与单体量 (修正方案)')

    add_h2(doc, '4.1 单体物性')

    add_table(doc,
              ['代号', '中文名', '英文名', 'CAS', '分子量', '官能团', '作用'],
              [
                  ['TFPT', '1,3,5-三(4-甲酰基苯基)-1,3,5-三嗪', '4,4\',4\'\'-(1,3,5-Triazine-2,4,6-triyl)tribenzaldehyde',
                   '443922-06-3', '393.4', '-CHO × 3', '三醛节点 (D系)'],
                  ['TAPT', '1,3,5-三(4-氨基苯基)-1,3,5-三嗪', '2,4,6-Tris(4-aminophenyl)-1,3,5-triazine',
                   '14544-47-9', '354.4', '-NH₂ × 3', '三胺节点 (A1系)'],
                  ['TFMB / BD-CF3', '2,2\'-双(三氟甲基)-4,4\'-联苯二胺', '2,2\'-Bis(trifluoromethyl)benzidine',
                   '341-58-2', '320.2', '-NH₂ × 2', '线性二胺 (D系连接臂)'],
                  ['A6', '4,4\'\'-双(三氟甲基)-2\',5\'-二甲酰基对三联苯', '4,4\'\'-Bis(trifluoromethyl)-2\',5\'-diformyl-p-terphenyl',
                   '1300701-03-4', '422.4', '-CHO × 2', '线性二醛 (A1系连接臂)'],
                  ['苯胺 (C₆H₅NH₂)', '苯胺', 'Aniline', '62-53-3', '93.1', '-NH₂ × 1', '★ 调制剂'],
                  ['乙酸 (HOAc)', '乙酸', 'Acetic acid', '64-19-7', '60.1', '—', '★ 催化剂 (6 M)'],
              ],
              col_widths=[Cm(1.5), Cm(3.5), Cm(4.0), Cm(2.0), Cm(1.5), Cm(2.0), Cm(2.0)])

    add_p(doc, '表 2 单体物性与作用')

    add_h2(doc, '4.2 化学计量')

    add_p(doc, '本方案严格按侯老师 AMCOF-1 化学计量 (TFPT : BD-CF3 = 1 : 1.5), 推导如下:')
    add_p(doc, '• TFPT 含 3 个醛基, TFMB 含 2 个氨基')
    add_p(doc, '• 设 TFPT = 0.03 mmol (醛基 0.09 mmol), TFMB = 0.045 mmol (氨基 0.09 mmol)')
    add_p(doc, '• 醛基 : 氨基 = 0.09 : 0.09 = 1 : 1 ✓ 化学计量')
    add_p(doc, '• 苯胺 = 3 × 醛基数 = 3 × 0.09 = 0.27 mmol... 但侯老师用 0.09 mmol = 3 × TFPT (0.03)')
    add_p(doc, '  → 这是因为苯胺当量是相对 TFPT 摩尔数, 不是醛基数')

    add_table(doc,
              ['项目', '摩尔量', '质量/体积', '说明'],
              [
                  ['TFPT', '0.03 mmol', '11.8 mg', '三醛节点'],
                  ['TFMB (BD-CF3)', '0.045 mmol', '14.4 mg', '二胺连接臂'],
                  ['苯胺', '0.09 mmol (3 × TFPT)', '8.2 μL (ρ=1.022 g/mL)', '★ 调制剂 (3 eq)'],
                  ['乙酸 (6 M)', '1.2 mmol', '0.20 mL', '★ 催化剂'],
                  ['甲苯', '—', '0.6 mL', '润湿剂 (玻璃接触角小)'],
                  ['氯仿', '—', '0.4 + 1.0 = 1.4 mL', '夹带剂 (沸点 61.2°C)'],
                  ['总体积', '—', '≈ 2.2 mL', '—'],
              ],
              col_widths=[Cm(3.5), Cm(3.0), Cm(4.5), Cm(4.5)])

    add_p(doc, '表 3 化学计量与试剂用量 (D-1 / A1-1 标准方案)')

    # ===== 第五章: 通用两步法 11 步流程 =====
    add_h1(doc, '第五章 通用两步法 11 步流程 (修正版)')

    add_p(doc, '以下流程是本方案所有 6 个对照实验的通用步骤, 各组具体参数差异化。')

    add_table(doc,
              ['步骤', '操作', '化学目的'],
              [
                  ['1', '35 mL 派热克斯玻璃管清洗',
                   '丙酮/乙醇/水各超声 10 min, 120°C 干燥 30 min。目的: 暴露极性 -OH 基团, 为后续氢键预排列做准备'],
                  ['2 [步骤 1 加料 - 关键]',
                   '向管中加入【三嗪节点 (TFPT/TAPT)】+ 甲苯 + 氯仿 + 【苯胺 (3 eq)】+ 【6 M 乙酸】',
                   '★ 核心创新: 步骤 1 加入苯胺+乙酸, 让基底膜以"苯胺-三嗪希夫碱"形式稳定形成, 而非无序聚集体'],
                  ['3', '超声混合 10 min',
                   '均匀分散单体, 避免局部浓度过高'],
                  ['4 [步骤 1 反应]',
                   '密封, 120°C 油浴 × 12-24 h',
                   '★ 步骤 1 时间: 比用户原 3 天短 (12-24h 即可形成稳定基底膜), 避免步骤 1 过度反应'],
                  ['5 [步骤 2 冷却]',
                   '取出 Pyrex 管, 冷却至 30-40°C',
                   '★ 安全: 防止氯仿爆沸 (D3/D4 错); 让基底膜稳定不溶解'],
                  ['6 [步骤 2 加连接臂]',
                   '沿管壁缓慢加入【线性连接臂 (TFMB/A6)】溶于氯仿 1.0 mL',
                   '新加的连接臂与基底膜的希夫碱交换反应, 在基底膜上生长 COF'],
                  ['7', '再超声 5 min',
                   '混合新加的连接臂, 避免局部浓度过高'],
                  ['8 [步骤 2 反应]',
                   '重新密封, 120°C 油浴 × 48 h',
                   '★ 步骤 2 时间: 48 h 是侯老师监测到的最优反应终点'],
                  ['9', '刮刀剥离',
                   'COF 膜 + 基底膜一起从玻璃壁上脱落, 形成"易脱落基底膜"结构'],
                  ['10', '丙酮 + THF 各洗 3 次',
                   '去除未反应单体和低聚物'],
                  ['11', '室温干燥, 称量, 算产率',
                   '对照预期产率 (D-1: ~88%, A1-1: ~70%)'],
              ],
              col_widths=[Cm(2.5), Cm(6.0), Cm(8.0)])

    add_p(doc, '表 4 通用两步法 11 步流程 (修正版)')

    add_callout(doc, '流程与原方案的关键差异',
                '【步骤 2 加入苯胺+乙酸】这是本 v8 方案与用户原方案 (A1/D3) 的核心差异。原方案步骤 1 没加调制剂, 导致基底膜不稳定。本方案步骤 1 加苯胺+乙酸, 让基底膜以稳定希夫碱形式形成。',
                color_hex='FFF8E1')

    # ===== 第六章: D 系新方案 =====
    add_h1(doc, '第六章 D 系新方案 (TFPT + TFMB)')

    add_p(doc, 'D 系 3 个对照方案: D-1 (推荐) / D-2 (对照) / D-3 (基线)')

    add_h2(doc, '6.1 D-1: 修正两步法 + 调制剂双介导 (★ 推荐)')

    add_h3(doc, '设计依据')
    add_p(doc, '基于 v8 修正策略: 步骤 1 加苯胺 + 乙酸让基底膜稳定; 步骤 2 加 TFMB 反应 48 h。预期重现侯老师 AMCOF-1 的 88.36% 产率, 同时实现用户的"易脱落基底膜"设想。')

    add_h3(doc, '详细操作')

    add_p(doc, '【步骤 1: 基底膜形成】(调制剂双介导)')
    add_p(doc, '1.1  Pyrex 管清洗 (按通用步骤 1)')
    add_p(doc, '1.2  向管中加入:')
    add_p(doc, '      • TFPT: 11.8 mg (0.03 mmol)')
    add_p(doc, '      • 甲苯: 0.6 mL')
    add_p(doc, '      • 氯仿: 0.4 mL')
    add_p(doc, '      • 苯胺 (★ 关键): 8.2 μL (0.09 mmol, 3 eq)')
    add_p(doc, '      • 6.0 M 乙酸 (★ 关键): 0.20 mL')
    add_p(doc, '1.3  超声 10 min')
    add_p(doc, '1.4  密封, 120°C 油浴 × 12 h (★ 比原方案 3 天短)')
    add_p(doc, '      【目的】让 TFPT 与苯胺形成苯胺-三嗪希夫碱 (m/z ≈ 469), 在玻璃壁上稳定自组装成基底膜')

    add_p(doc, '【步骤 2: COF 膜生长】')
    add_p(doc, '2.1  ★ 取出 Pyrex 管, 冷却至 30-40°C (★ 防止氯仿爆沸)')
    add_p(doc, '2.2  沿管壁缓慢加入: TFMB (14.4 mg, 0.045 mmol) 溶于氯仿 1.0 mL')
    add_p(doc, '      【目的】TFMB 与基底膜的希夫碱交换, 形成 COF')
    add_p(doc, '2.3  再超声 5 min')
    add_p(doc, '2.4  重新密封, 120°C 油浴 × 48 h')

    add_p(doc, '【步骤 3: 收集与表征】(按通用步骤 9-11)')

    add_h3(doc, '预期结果与失败排查')

    add_table(doc,
              ['预测现象', '原因', '修正'],
              [
                  ['★ 预期: 易脱落黄色连续膜, 产率 ~88%',
                   '严格按修正两步法 → 重现 AMCOF-1 + 实现易脱落',
                   '—'],
                  ['步骤 1 后基底膜不存在',
                   '苯胺/乙酸量不足 → 希夫碱未充分形成',
                   '严格 3 eq 苯胺 + 6 M 乙酸'],
                  ['步骤 2 加入 TFMB 后基底膜被冲刷掉',
                   '新加溶剂过多 (1.0 mL 氯仿)',
                   '改用更少溶剂 (0.5 mL) 或预热 TFMB 溶液'],
                  ['玻璃上有粉末, 无膜',
                   '步骤 1 时间过短 (< 12 h), 基底膜未形成',
                   '延长步骤 1 至 24 h'],
                  ['步骤 2 反应 48 h 后膜仍然不连续',
                   'TFMB 反应不充分 → 延长至 72 h'],
              ],
              col_widths=[Cm(4.5), Cm(5.5), Cm(5.0)])

    add_p(doc, '表 5 D-1 预期结果与失败排查')

    add_h2(doc, '6.2 D-2: 用户原两步法 (验证失败原因)')

    add_h3(doc, '设计依据')
    add_p(doc, '严格按用户 D3 实验流程, 不做任何修正, 验证"步骤 1 没加调制剂"是失败的根本原因。预期: 与 D3 类似的失败结果。')

    add_h3(doc, '详细操作')

    add_p(doc, '【步骤 1: 自组装 (无调制剂, 与用户原方案一致)】')
    add_p(doc, '1.1  Pyrex 管清洗')
    add_p(doc, '1.2  加入 TFPT (11.8 mg, 0.03 mmol) + 甲苯 (0.6 mL) + 氯仿 (0.4 mL)')
    add_p(doc, '      ★ 注意: 步骤 1 不加苯胺/乙酸 (这是与 D-1 的核心区别)')
    add_p(doc, '1.3  超声 10 min')
    add_p(doc, '1.4  120°C 油浴 × 12 h')

    add_p(doc, '【步骤 2: 加 TFMB 反应 (与 D-1 相同)】')
    add_p(doc, '2.1  冷却至 30-40°C')
    add_p(doc, '2.2  加 TFMB 14.4 mg 溶于氯仿 1.0 mL + 苯胺 8.2 μL + 6 M 乙酸 0.20 mL')
    add_p(doc, '2.3  再超声 5 min')
    add_p(doc, '2.4  重新 120°C 油浴 × 48 h')

    add_callout(doc, '预期对比',
                'D-1 与 D-2 的唯一区别是步骤 1 是否加苯胺+乙酸。如果 D-1 成功而 D-2 失败, 直接证明"调制剂双介导"是关键。如果两者都成功或都失败, 说明还有其他因素 (如步骤 2 加料方式) 需要调整。',
                color_hex='FFF8E1')

    add_h2(doc, '6.3 D-3: 一锅法 (侯老师 AMCOF-1 原文, 作为基线)')

    add_h3(doc, '设计依据')
    add_p(doc, '严格按侯老师 AMCOF-1 Scheme S1 原文操作, 作为基线对照。如果 D-3 成功而 D-1/D-2 失败, 说明两步法不适合 TFPT+TFMB; 如果 D-3 失败而 D-1 成功, 说明两步法是创新方向。')

    add_h3(doc, '详细操作')

    add_p(doc, '1. Pyrex 管清洗')
    add_p(doc, '2. 加 TFPT (11.8 mg, 0.03 mmol) + 甲苯 (0.6 mL) + 氯仿 (0.4 mL) + 苯胺 (8.2 μL)')
    add_p(doc, '3. 加 6 M 乙酸 (0.20 mL) + TFMB (14.4 mg, 0.045 mmol) 溶于氯仿 (1.0 mL)')
    add_p(doc, '4. 超声 10 min')
    add_p(doc, '5. 120°C 油浴 × 48 h (★ 一次性, 无中间步骤)')
    add_p(doc, '6. 刮刀剥离, 丙酮+THF 洗, 室温干燥')

    add_h2(doc, '6.4 D 系对照实验汇总')

    add_table(doc,
              ['编号', '步骤 1 含调制剂?', '步骤 1 时间', '步骤 2 含调制剂?', '步骤 2 时间', '核心设计'],
              [
                  ['D-1 ★', '★ 含 (苯胺+乙酸)', '12 h', '加 TFMB', '48 h', '修正两步法'],
                  ['D-2', '无 (用户原方案)', '12 h', '加 TFMB+苯胺+乙酸', '48 h', '对照原方案'],
                  ['D-3', '— (无步骤 1)', '—', '一次性加所有', '48 h', '侯老师基线'],
              ],
              col_widths=[Cm(1.5), Cm(2.5), Cm(1.8), Cm(2.5), Cm(1.8), Cm(4.5)])

    add_p(doc, '表 6 D 系 3 个对照方案汇总')

    # ===== 第七章: A1 系新方案 =====
    add_h1(doc, '第七章 A1 系新方案 (TAPT + A6)')

    add_p(doc, 'A1 系 3 个对照方案: A1-1 (推荐) / A1-2 (对照) / A1-3 (基线)')

    add_h2(doc, '7.1 A1-1: 修正两步法 + 调制剂双介导 (★ 推荐)')

    add_h3(doc, '设计依据')
    add_p(doc, '与 D-1 类似, 但节点用 TAPT (三胺) + A6 (三联苯二醛, 含 -CF₃, 大位阻)。苯胺优先与 A6 醛基反应 (因为 TAPT 是胺, 不是醛), 形成苯胺-A6 希夫碱作为基底膜。')

    add_h3(doc, '详细操作')

    add_p(doc, '【步骤 1: 基底膜形成 (TAPT 自组装 + 苯胺-A6 希夫碱)】')
    add_p(doc, '1.1  Pyrex 管清洗')
    add_p(doc, '1.2  向管中加入:')
    add_p(doc, '      • TAPT: 10.6 mg (0.03 mmol)')
    add_p(doc, '      • A6 (★ 步骤 1 加 A6): 19.0 mg (0.045 mmol)')
    add_p(doc, '      • 甲苯: 0.6 mL')
    add_p(doc, '      • 氯仿: 0.4 mL')
    add_p(doc, '      • 苯胺 (★): 8.2 μL (3 eq)')
    add_p(doc, '      • 6.0 M 乙酸 (★): 0.20 mL')
    add_p(doc, '【★ 与 D-1 的关键差异】步骤 1 同时加 TAPT 和 A6 (因为 TAPT 是胺, 苯胺必须与 A6 醛基反应形成希夫碱作为基底膜)')
    add_p(doc, '1.3  超声 10 min')
    add_p(doc, '1.4  120°C 油浴 × 12-24 h')

    add_callout(doc, '为什么 A1-1 步骤 1 同时加 TAPT 和 A6',
                '在 A1 体系中, TAPT 是胺, A6 是醛。苯胺必须先与醛基反应形成希夫碱。所以步骤 1 必须同时含 TAPT+A6+苯胺+乙酸, 让: (a) 苯胺与 A6 形成希夫碱作为基底膜; (b) TAPT 在基底膜上自组装。\n'
                '这与 D-1 不同 (D-1 步骤 1 只加 TFPT+苯胺+乙酸, 因为 TFPT 是醛, 苯胺直接与 TFPT 醛基反应即可)。',
                color_hex='FFF8E1')

    add_p(doc, '【步骤 2: COF 膜生长 (TAPT 与 A6 充分反应)】')
    add_p(doc, '2.1  冷却至 30-40°C')
    add_p(doc, '2.2  沿管壁加入: 额外 TAPT (★) 10.6 mg 溶于氯仿 1.5 mL')
    add_p(doc, '【★ 与 D-1 不同】步骤 2 加更多 TAPT (因为步骤 1 已经消耗部分 TAPT 用于形成基底膜)')
    add_p(doc, '2.3  再超声 5 min')
    add_p(doc, '2.4  重新 120°C 油浴 × 48-72 h (★ 延长, 因为 A6 大位阻)')

    add_h2(doc, '7.2 A1-2: 用户原两步法 (验证失败原因)')

    add_h3(doc, '详细操作')
    add_p(doc, '【步骤 1: 仅 TAPT 自组装 (无 A6/苯胺/乙酸, 与用户 A1 一致)】')
    add_p(doc, '1.1  Pyrex 管清洗')
    add_p(doc, '1.2  加 TAPT (10.6 mg, 0.03 mmol) + 甲苯 (0.3 mL) + 氯仿 (2.2 mL)')
    add_p(doc, '      ★ 步骤 1 不加 A6 / 苯胺 / 乙酸')
    add_p(doc, '1.3  超声 10 min')
    add_p(doc, '1.4  120°C 油浴 × 12 h')

    add_p(doc, '【步骤 2: 加 A6 反应 (与用户 A1 一致)】')
    add_p(doc, '2.1  冷却至 30-40°C')
    add_p(doc, '2.2  加 A6 (19.0 mg, 0.045 mmol) + 甲苯 (0.12 mL) + 氯仿 (0.88 mL) + 苯胺 (8.2 μL) + 6 M 乙酸 (0.28 mL)')
    add_p(doc, '2.3  超声 10 min')
    add_p(doc, '2.4  重新 120°C 油浴 × 72 h')

    add_h2(doc, '7.3 A1-3: 一锅法 (基线对照)')

    add_p(doc, '严格按 AMCOF-1 逻辑, 但用 TAPT+A6 替代 TFPT+TFMB:')
    add_p(doc, '1. 加 TAPT (10.6 mg) + A6 (19.0 mg) + 甲苯 (0.6 mL) + 氯仿 (0.4 mL) + 苯胺 (8.2 μL)')
    add_p(doc, '2. 加 6 M 乙酸 (0.20 mL) + 额外 TAPT (★) 溶于氯仿 (1.5 mL)')
    add_p(doc, '3. 超声 10 min, 120°C 油浴 × 72 h')

    add_h2(doc, '7.4 A1 系对照实验汇总')

    add_table(doc,
              ['编号', '步骤 1 含调制剂?', '步骤 1 节点/醛?', '步骤 2 加什么?', '时间', '核心设计'],
              [
                  ['A1-1 ★', '★ 含', 'TAPT + A6', '额外 TAPT', '步骤 1: 12-24 h, 步骤 2: 48-72 h',
                   '修正两步法 (TAPT+A6 同步加)'],
                  ['A1-2', '无 (用户原方案)', '仅 TAPT', 'A6 + 苯胺 + 乙酸', '步骤 1: 12 h, 步骤 2: 72 h',
                   '用户原方案'],
                  ['A1-3', '— (无步骤 1)', '一次性加', '—', '72 h', '一锅法基线'],
              ],
              col_widths=[Cm(1.5), Cm(2.0), Cm(2.5), Cm(2.5), Cm(2.5), Cm(4.5)])

    add_p(doc, '表 7 A1 系 3 个对照方案汇总')

    # ===== 第八章: 加料顺序的化学解释 =====
    add_h1(doc, '第八章 加料顺序的化学逻辑 (每步为什么)')

    add_table(doc,
              ['步骤', '操作', '化学目的'],
              [
                  ['1', 'Pyrex 管清洗',
                   '去除表面有机污染物; 暴露极性 -OH 基团, 为后续氢键预排列做准备 (DFT 证实)'],
                  ['2 [步骤 1 加料]',
                   '三嗪节点 + 线性醛 + 苯胺 + 乙酸 + 溶剂 (D-1 不加醛, A1-1 加醛)',
                   '★ 让苯胺与醛基反应形成稳定的希夫碱基底膜, 避免步骤 1 形成无序聚集体'],
                  ['3', '超声 10 min',
                   '均匀分散, 避免局部浓度过高'],
                  ['4 [步骤 1 反应]',
                   '120°C × 12-24 h',
                   '★ 形成稳定的基底膜 (不是用户的 3 天, 而是 12-24 h 即可, 避免步骤 1 过度反应)'],
                  ['5 [冷却]',
                   '取出 Pyrex 管, 冷却至 30-40°C',
                   '★ 防止氯仿爆沸 (D3/D4 错); 让基底膜稳定不溶解'],
                  ['6 [步骤 2 加连接臂]',
                   '沿管壁缓慢加入线性连接臂',
                   '★ 新加的连接臂与基底膜的希夫碱交换, 在基底膜上生长 COF'],
                  ['7', '再超声 5 min',
                   '混合新加的连接臂, 避免局部浓度过高 (但不要超声太久, 否则会冲刷基底膜)'],
                  ['8 [步骤 2 反应]',
                   '重新 120°C × 48-72 h',
                   '48 h 是侯老师监测到的最优反应终点'],
                  ['9', '刮刀剥离',
                   'COF 膜 + 基底膜一起从玻璃壁上脱落, 形成"易脱落基底膜"结构'],
                  ['10', '丙酮 + THF 各洗 3 次',
                   '去除未反应单体和低聚物'],
                  ['11', '室温干燥',
                   '低温避免膜氧化或结构变化'],
              ],
              col_widths=[Cm(2.0), Cm(6.5), Cm(8.0)])

    add_p(doc, '表 8 通用两步法 11 步流程的化学解释')

    # ===== 第九章: 失败排查 =====
    add_h1(doc, '第九章 失败排查与修正')

    add_h2(doc, '9.1 步骤 1 相关排查')

    add_h3(doc, '症状: 步骤 1 后基底膜不存在')
    add_p(doc, '【原因】苯胺/乙酸量不足 → 希夫碱未充分形成, 三嗪单体自聚成无序聚集体')
    add_p(doc, '【修正】严格 3 eq 苯胺 + 6 M 乙酸')

    add_h3(doc, '症状: 步骤 1 后基底膜粗糙不连续')
    add_p(doc, '【原因】步骤 1 时间不足 (< 12 h), 或温度不达标 (< 120°C)')
    add_p(doc, '【修正】延长至 24 h, 严格 120°C 油浴')

    add_h3(doc, '症状: 步骤 1 后基底膜过厚 (>1 μm)')
    add_p(doc, '【原因】步骤 1 时间过长 (>48 h), 或三嗪浓度过高')
    add_p(doc, '【修正】严格 12-24 h, 浓度按侯老师原文 (0.03 mmol TFPT/2.2 mL)')

    add_h2(doc, '9.2 步骤 2 相关排查')

    add_h3(doc, '症状: 步骤 2 加入 TFMB/A6 后基底膜被冲刷掉')
    add_p(doc, '【原因】新加的溶剂 (1.0 mL 氯仿) 溶解了基底膜')
    add_p(doc, '【修正】用更少溶剂 (0.5 mL) 或预热 TFMB/A6 溶液到 50°C 后再加入')

    add_h3(doc, '症状: 步骤 2 反应后玻璃上散在颗粒, 无连续膜')
    add_p(doc, '【原因】TFMB/A6 反应不充分; 或步骤 2 时间过短 (< 48 h)')
    add_p(doc, '【修正】延长步骤 2 至 72 h, 检查是否所有 TFMB/A6 完全溶解')

    add_h3(doc, '症状: 未冷却开盖 → 氯仿爆沸 (D3/D4 类)')
    add_p(doc, '【原因】违反温度安全规程, 120°C 时管内压力大, 突然开盖氯仿瞬间沸腾')
    add_p(doc, '【修正】★ 严格冷却至 30-40°C 再开盖, 用 PTFE 内衬盖而非普通玻璃塞, 绝不直接开盖!')

    add_h2(doc, '9.3 D 系特异排查')

    add_h3(doc, '症状: TFMB 加料后立即浑浊')
    add_p(doc, '【原因】TFMB 在氯仿中溶解度 + (5-20 mg/mL), 14.4 mg/1.0 mL 接近饱和')
    add_p(doc, '【修正】加料前超声 TFMB+氯仿 10 min 完全溶解, 或预热至 50°C')

    add_h2(doc, '9.4 A1 系特异排查')

    add_h3(doc, '症状: A6 在体系中析出 (颗粒物)')
    add_p(doc, '【原因】A6 在甲苯/氯仿中溶解度 + ~ ++, 但 A6 三联苯骨架大, 19 mg/1.0 mL 接近溶解度上限')
    add_p(doc, '【修正】预热 A6 溶液至 50°C, 或用 BTF 替代部分甲苯 (含氟与 A6 亲氟)')

    add_h3(doc, '症状: TAPT 在步骤 1 没完全溶解')
    add_p(doc, '【原因】TAPT 在甲苯/氯仿中溶解度 +/- (微溶)')
    add_p(doc, '【修正】步骤 1 用氯仿更多 (0.6 mL 而非 0.4 mL), 或预热超声')

    # ===== 第十章: 表征方案 =====
    add_h1(doc, '第十章 表征方案')

    add_table(doc,
              ['表征', '目的', '关键判定', '数据对比 (AMCOF-1)'],
              [
                  ['PXRD', '结晶度', 'fwhm(100) ≤ 0.16°', 'fwhm = 0.16°'],
                  ['FT-IR', '亚胺键形成', 'C=N 峰 ~1631 cm⁻¹ 出现, C=O ~1708 消失', '1631 cm⁻¹'],
                  ['SEM', '膜表面形貌', '气相侧棱面晶, 玻璃侧光滑', '棱面晶'],
                  ['AFM', '膜厚与粗糙度', '膜厚均匀, Ra < 5 nm', '4.60 μm 厚'],
                  ['BET', '比表面积', '> 500 m²/g, 孔径 ~3.9 nm', '594.87 m²/g'],
                  ['XPS', '元素 + 化学态', 'C, N, F 峰对应', '含 F'],
                  ['接触角', '疏水性', '含氟膜 > 100°', '应 > 110°'],
                  ['★ 新增: 剥离实验', '易脱落基底膜验证',
                   '刮刀剥离后, 膜能否从基底膜上分离, 形成独立 COF 膜',
                   '验证用户"易脱落基底膜"设想'],
              ],
              col_widths=[Cm(3.0), Cm(2.5), Cm(5.0), Cm(5.0)])

    add_p(doc, '表 9 表征方案 (新增"剥离实验"验证易脱落基底膜)')

    # ===== 第十一章: 参考文献 =====
    add_h1(doc, '第十一章 参考文献')

    add_h2(doc, '11.1 核心文献')

    add_p(doc, '【1】侯盛怀, 张干兵, 乔钊宇, 白玉轩, 邸昊昕, 华烨同, 郝甜甜, 徐慧* "扩散/调制剂双介导固-液/气界面合成结晶型共价有机框架膜"《德国应用化学》 2025, 64, e202421555.')
    add_p(doc, '   DOI: 10.1002/anie.202421555', indent=False)
    add_p(doc, '   本地路径: 实验\\文章\\侯老师实验\\侯盛怀德国应化.pdf', indent=False)
    add_p(doc, '   补充信息: 实验\\文章\\侯老师实验\\sl-884-anie202421555-sup-0001-misc_information.pdf', indent=False)
    add_p(doc, '   提取文本: 实验\\文章\\侯老师实验\\_main_full.txt + _si.txt', indent=False)

    add_h2(doc, '11.2 GraphRAG v2 检索到的相关文献')

    add_p(doc, '【2】L-0ca3cbf692f8 (Science Advances). 界面合成制备高结晶度 2D COF 膜. DOI 待补充.')
    add_p(doc, '【3】L-c8cc6ff417dd (J. Am. Chem. Soc.). SI-SBMAP 方法, sp²c-COF 薄膜 (表面引发法先例). DOI 待补充.')
    add_p(doc, '【4】L-a35edf0e281d (J. Am. Chem. Soc.). Boc 保护 + TFA 脱保护, 编织结构 COF-112. DOI 待补充.')

    add_h2(doc, '11.3 本方案相关本地文献')

    add_p(doc, '【5】tianxuan-seek 训练数据集. C:\\Users\\ckx\\Desktop\\tianxuan seek\\data\\structured\\')
    add_p(doc, '【6】实验 ABCDEF 反馈记录. C:\\Users\\ckx\\Desktop\\实验\\方案\\实验ABCDEF.docx')
    add_p(doc, '【7】v3.9 实验方案 (上游模板). C:\\Users\\ckx\\Desktop\\实验\\方案\\实验方案_含氟COF薄膜合成_v3.9_20260626.docx')

    # ===== 第十二章: 实验准备清单 =====
    add_h1(doc, '第十二章 实验准备清单')

    add_h2(doc, '12.1 试剂清单 (D-1 方案)')

    add_table(doc,
              ['名称', 'CAS', '用量', '来源'],
              [
                  ['TFPT', '443922-06-3', '11.8 mg', 'D&B Biotechnology'],
                  ['TFMB / BD-CF3', '341-58-2', '14.4 mg', 'D&B Biotechnology'],
                  ['苯胺', '62-53-3', '8.2 μL', 'Sinopharm 国药'],
                  ['乙酸 (冰醋酸)', '64-19-7', '0.20 mL × 6 M', 'Sinopharm 国药'],
                  ['甲苯', '108-88-3', '0.6 mL', 'Sinopharm 国药'],
                  ['氯仿', '67-66-3', '0.4 + 1.0 = 1.4 mL', 'Kelong 科龙'],
                  ['丙酮 (洗涤)', '67-64-1', '5 mL × 3', '—'],
                  ['THF (洗涤)', '109-99-9', '5 mL × 3', '—'],
                  ['去离子水 (配乙酸)', '—', '适量', 'Milli-Q'],
              ],
              col_widths=[Cm(4.0), Cm(3.0), Cm(4.0), Cm(4.5)])

    add_p(doc, '表 10 D-1 试剂清单')

    add_h2(doc, '12.2 试剂清单 (A1-1 方案)')

    add_table(doc,
              ['名称', 'CAS', '用量', '来源'],
              [
                  ['TAPT', '14544-47-9', '10.6 + 10.6 = 21.2 mg', '—'],
                  ['A6', '1300701-03-4', '19.0 mg', '—'],
                  ['苯胺', '62-53-3', '8.2 μL', 'Sinopharm 国药'],
                  ['乙酸 (冰醋酸)', '64-19-7', '0.20 mL × 6 M', 'Sinopharm 国药'],
                  ['甲苯', '108-88-3', '0.6 mL', 'Sinopharm 国药'],
                  ['氯仿', '67-66-3', '0.4 + 1.5 = 1.9 mL', 'Kelong 科龙'],
                  ['丙酮 (洗涤)', '67-64-1', '5 mL × 3', '—'],
                  ['THF (洗涤)', '109-99-9', '5 mL × 3', '—'],
              ],
              col_widths=[Cm(4.0), Cm(3.0), Cm(4.0), Cm(4.5)])

    add_p(doc, '表 11 A1-1 试剂清单 (TAPT 用量是 D-1 的 2 倍, 因为步骤 1+2 各加 1 份)')

    add_h2(doc, '12.3 设备清单')

    add_bullet(doc, '35 mL 派热克斯玻璃管 (Pyrex tube)')
    add_bullet(doc, '聚四氟乙烯 (PTFE) 内衬螺旋盖 (防止氯仿爆沸)')
    add_bullet(doc, '120°C 油浴 (带磁力搅拌或不搅拌均可)')
    add_bullet(doc, '超声清洗仪')
    add_bullet(doc, '金属刮刀')
    add_bullet(doc, '微量注射器 (10 μL / 100 μL / 1 mL)')
    add_bullet(doc, '分析天平 (0.1 mg 精度)')

    add_p(doc, '')
    add_p(doc, '--- 本方案结束 ---')

    doc.save(OUT)
    print(f'✓ 已生成: {OUT}')
    print(f'  大小: {os.path.getsize(OUT)/1024:.1f} KB')


if __name__ == '__main__':
    main()