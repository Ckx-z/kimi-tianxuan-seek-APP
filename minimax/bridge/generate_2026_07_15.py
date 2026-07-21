"""bridge/generate_2026_07_13.py — JACS 2023 原文修正版"""
import os, sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from _docx_styles import *
from docx import Document

PROJ = r'C:\Users\ckx\Desktop\minimax'
OUT = os.path.join(PROJ, 'experiment', 'proposals\COF-TFPT-2026-07-15-jelly-v8.docx', 'COF-TFPT-2026-07-15-jelly-v8.docx')

HOU = """\
【步骤1: Pyrex管准备】35mL派热克斯玻璃管,丙酮/乙醇/水各超声10min,120°C干燥30min。

【步骤2: 投料(一次性加所有)】向管中依次加入:(1)TFPT 23.6mg;(2)甲苯 1.2mL;(3)氯仿 0.8mL;(4)苯胺 {aniline_ul}μL({aniline_eq}eq);(5)6.0M乙酸水溶液 {hac_ml}mL;(6)TFMB 28.8mg(预先溶于氯仿2.0mL)。加料目的:苯胺优先与TFPT醛基形成希夫碱中间体(m/z≈469),乙酸催化,TFMB最后加入交换形成COF。

【步骤3: 超声】超声10min至完全溶解或微浊。

【步骤4: 加热反应】PTFE内衬螺旋盖密封,120°C油浴×2天(48h)。低沸点溶剂蒸发夹带单体向上→玻璃壁超饱和液层→限域微反应器。

【步骤5: 冷却+收集】冷却至室温,金属刮刀剥离产物。凝胶状用药勺取出。

【步骤6: 洗涤】丙酮+THF各洗3次。

【步骤7: 干燥与称量】室温干燥,称量,记录形态。"""

USER = """\
【步骤1: Pyrex管准备】丙酮/乙醇/水各超声10min,120°C干燥。

【步骤2: 第一次投料(自组装,★加少量TFMB模拟瓶壁残留)】加入:TFPT 23.6mg+甲苯1.2mL+氯仿0.8mL+★TFMB 3.0mg(约0.009mmol,为正常量的10%)。目的:步骤1中引入少量TFMB,让TFPT与微量胺在玻璃壁上部分反应形成松散网络。

【步骤3: 第一次超声】超声10min至完全溶解或微浊。

【步骤4: 自组装反应(★形成果冻状的阶段)】密封,120°C油浴×1天(24h)。预期:果冻状初始产物。

【步骤5: 冷却】取出Pyrex管,冷却至30-40°C(防止开盖氯仿爆沸)。

【步骤6: 第二次投料(COF形成)】沿管壁缓慢加入:苯胺16.4μL+TFMB 25.8mg(预先溶于氯仿2.0mL,补足剩余量,总TFMB=28.8mg)+6.0M乙酸0.4mL。

【步骤7: 第二次超声】再超声5min(不宜过久以免冲刷初始产物)。

【步骤8: COF形成反应】密封,120°C油浴×2天(48h)。

【步骤9: 收集】冷却至室温,金属刮刀或药勺取出。

【步骤10: 洗涤】丙酮+THF各3次。

【步骤11: 干燥与称量】室温干燥,称量,记录形态。"""

BOC = """\
【步骤1: 密封釜准备】取25mL聚四氟乙烯内衬不锈钢密封釜(若无,可用35mL Pyrex管+PTFE盖代替),丙酮/乙醇/水各超声10min,120°C干燥。

【步骤2: 投料(★严格按JACS 2023 SI Figure S22配比)】向密封釜中依次加入:
  (1)TFPT(三醛节点):23.6mg(0.06mmol)
  (2)★Boc-TFMB(Boc保护的TFMB):约47mg(0.09mmol)
     Boc-TFMB分子量≈520g/mol(TFMB 320+2×Boc 101-2×1)
     JACS原文:Pa-NHBoc 27.8mg(0.09mmol)+Tp 12.6mg(0.06mmol)
     →本方案按相同醛:胺摩尔比换算(1:1.5)
     ★如不可购买,可按附录制备:TFMB+Boc₂O→Boc-TFMB
  (3)★1,4-二氧六环(有机溶剂):3.4mL
  (4)★水(★凝胶形成核心!提供氢键网络):0.6mL
     JACS原文:\"dioxane 3.4mL+H₂O 0.6mL\"(二元溶剂体系)
  (5)★TFA(三氟乙酸):60μL(约0.78mmol)
     JACS原文SI Table S2:TFA用量优化0/30/60/90/120/150μL,60μL为最佳
     TFA的双重作用:(a)脱保护—催化Boc裂解,缓慢释放游离-NH₂;
     (b)催化剂—催化希夫碱缩合反应(替代乙酸)
  ★不加苯胺和乙酸(JACS原文无此两项)

【步骤3: 密封】将密封釜盖紧(★水在120°C产生蒸汽压,必须密封)。

【步骤4: 加热反应】将密封釜置于120°C烘箱(或油浴)×3天(72h)。
  机理:Boc在TFA催化下逐步裂解→游离-NH₂缓慢释放→与TFPT醛基反应形成纳米纤维/纳米片;
  水与二氧六环形成二元氢键网络→被COF骨架固定→形成有机水凝胶。

【步骤5: 冷却+取出】冷却至室温,小心取出。预期:圆柱形有机水凝胶(弹性好)。
  JACS原文描述:\"a cylindrical organohydrogel with good appearance and elasticity\"

【步骤6: ★不要用丙酮/THF洗】有机水凝胶含水和有机溶剂,传统洗涤会破坏凝胶。
  如需要:冷冻干燥→气凝胶;或溶剂交换(水→乙醇→丙酮)。

【步骤7: 表征】拍照+压缩弹性测试+PXRD/FT-IR/SEM。"""


def main():
    doc = Document()
    setup_styles(doc)
    s = doc.sections[0]; s.top_margin = Cm(2); s.bottom_margin = Cm(2)
    s.left_margin = Cm(2); s.right_margin = Cm(2)

    # 封面
    tp = doc.add_paragraph(); tp.alignment = 1  # CENTER
    tp.paragraph_format.space_before = Pt(36); tp.paragraph_format.space_after = Pt(12)
    add_run(tp, '果冻状 COF 膜重现与调制剂优化方案', size=22, bold=True, color=RGBColor(0x1F,0x3A,0x5F), font='黑体')

    sp = doc.add_paragraph(); sp.alignment = 1
    sp.paragraph_format.space_after = Pt(6)
    add_run(sp, '含 B-6 Boc 保护策略 (JACS 2023 SI 精确配比)', size=14, bold=True, color=RGBColor(0x2E,0x5C,0x8A), font='黑体')

    s2 = doc.add_paragraph(); s2.alignment = 1
    s2.paragraph_format.space_after = Pt(36)
    add_run(s2, 'B-1/B-3~B-5: 侯老师一锅法 | B-2: 用户两步法 | B-6: Boc+TFA (Jia 2023)', size=11, italic=True, color=RGBColor(0x6C,0x75,0x7D))

    ip = doc.add_paragraph(); ip.alignment = 1
    add_run(ip, '方案编号: COF-TFPT-2026-07-15-jelly-v8 | 生成日期: 2026-07-15', size=10.5)
    ip.add_run().add_break()
    add_run(ip, '核心文献: Hou 2025 (DOI 10.1002/anie.202421555) | Jia 2023 (DOI 10.1021/jacs.3c09284)', size=10, italic=True)

    add_page_break(doc)

    # 摘要
    add_h1(doc, '摘要')
    add_callout(doc, '方案设计思路',
        '本方案基于用户在自组装探究中的意外发现(瓶壁残留TFMB→果冻状产物),设计6组对照实验。'
        'B-2模拟用户的瓶壁残留效应;B-6基于贾淑萍等JACS 2023原文(DOI 10.1021/jacs.3c09284,已读全文+SI)的精确配比,'
        '用Boc保护TFMB的氨基+TFA脱保护→缓慢释放游离-NH₂+水/二氧六环二元溶剂氢键网络→有机水凝胶。'
        '其余组为侯老师一锅法对照。', 'FFF8E1')

    add_p(doc, '6个对照:', bold=True)
    for t in [
        'B-1基线(侯老师一锅法):TFPT+甲苯+氯仿+苯胺+6M乙酸+TFMB→120°C×2天。预期:普通膜/粉末。',
        'B-2★果冻重现(用户两步法):步骤1 TFPT+少量TFMB(3mg)→自组装1天;步骤2+剩余TFMB+苯胺+乙酸→反应2天。',
        'B-3减苯胺:一锅法,苯胺16.4→8.2μL。',
        'B-4减乙酸:一锅法,乙酸0.4→0.2mL(侯老师最佳量)。',
        'B-5双重减:一锅法,苯胺减半+乙酸减半。',
        'B-6★Boc保护(JACS 2023 SI精确配比):TFPT+Boc-TFMB+二氧六环3.4mL+水0.6mL+TFA60μL→120°C×3天。不加苯胺/乙酸。']:
        add_bullet(doc, t)

    # 一、背景
    add_h1(doc, '一、实验背景')
    add_h2(doc, '1.1 操作失误与果冻状发现')
    add_p(doc, '用户在探究TFPT自组装时发生操作失误:称量TFMB时错误加入Pyrex管,倒掉后瓶身未清洗(残留少量TFMB),随后加入TFPT继续自组装。残留的TFMB与TFPT在步骤1中反应,形成了有弹性、果冻状柔软的产物。B-2在步骤1加入3mg TFMB模拟此效应。')

    add_h2(doc, '1.2 JACS 2023 水凝胶文献支撑(B-6的科学基础)')
    add_p(doc, '贾淑萍等JACS 2023(DOI 10.1021/jacs.3c09284,本地SI已读61页)提出了通用基团保护合成策略制备COF有机水凝胶:', bold=True)
    add_bullet(doc, '核心方法:用Boc保护胺单体(如Pa-NHBoc)+醛(如Tp)+二氧六环/水(3.4/0.6mL)+TFA(60μL)→120°C×72h→圆柱形有机水凝胶')
    add_bullet(doc, '机理:(1)Boc保护胺无活性,TFA催化Boc逐步裂解,缓慢释放游离-NH₂→反应速率可控;(2)水与二氧六环形成二元氢键网络→被COF纳米纤维/纳米片骨架固定→凝胶化')
    add_bullet(doc, 'TFA用量优化:0/30/60/90/120/150μL,60μL为最佳(SI Table S2)')
    add_bullet(doc, '★全文未使用苯胺和乙酸(TFA同时做脱保护+催化)')
    add_bullet(doc, '★溶剂必须含水(无水的纯有机溶剂得不到凝胶)')

    add_h2(doc, '1.3 单体物性')
    add_table(doc, ['代号','中文名','CAS','分子量','官能团','作用'],
        [['TFPT','1,3,5-三(4-甲酰基苯基)-1,3,5-三嗪','443922-06-3','393.4','-CHO×3','三醛节点'],
         ['TFMB','2,2-双(三氟甲基)-4,4-联苯二胺','341-58-2','320.2','-NH₂×2','二胺连接臂'],
         ['Boc-TFMB','Boc保护的TFMB','—','≈520','-NHBoc×2','★B-6保护胺'],
         ['苯胺','苯胺','62-53-3','93.1','-NH₂×1','调制剂'],
         ['乙酸','乙酸','64-19-7','60.1','—','催化剂(6M)'],
         ['TFA','三氟乙酸','76-05-1','114.0','—','★B-6脱保护+催化']],
        col_widths=[Cm(1.8),Cm(4.0),Cm(2.0),Cm(1.5),Cm(2.0),Cm(2.5)])

    # B-1~B-5
    for title, an, eq, hac, desc, is_b2 in [
        ('二、B-1 基线(侯老师一锅法)','16.4','3','0.4','严格按侯老师AMCOF-1一锅法。作为基线对照。',False),
        ('三、B-2 ★ 果冻状重现(用户两步法)','16.4','3','0.4','步骤1加3mg TFMB模拟瓶壁残留,步骤2补足。总TFMB=28.8mg(与B-1一致)。',True),
        ('四、B-3 减少苯胺(侯老师一锅法)','8.2','1.5','0.4','苯胺16.4→8.2μL(减半)。调制剂减弱→反应变快。',False),
        ('五、B-4 减少乙酸(侯老师一锅法)','16.4','3','0.2','乙酸0.4→0.2mL(侯老师最佳量)。催化减弱→可逆性恢复。',False),
        ('六、B-5 双重减少(侯老师一锅法)','8.2','1.5','0.2','苯胺减半+乙酸减半。最弱催化条件。',False)]:
        add_page_break(doc)
        add_h1(doc, title)
        add_h3(doc, '设计目的'); add_p(doc, desc, bold=True)
        if is_b2:
            add_h3(doc, '试剂'); add_p(doc, '步骤1:TFPT 23.6mg|甲苯 1.2mL|氯仿 0.8mL|★TFMB 3.0mg。步骤2:苯胺16.4μL|★TFMB 25.8mg|6M乙酸0.4mL|氯仿2.0mL')
            add_h3(doc, '完整操作步骤'); add_p(doc, USER)
            add_h3(doc, '预期'); add_p(doc, '步骤1形成果冻状基底,步骤2继续生长。预期重现用户观察到的柔软果冻状产物。')
        else:
            add_h3(doc, '试剂'); add_p(doc, f'TFPT 23.6mg|TFMB 28.8mg|★苯胺{an}μL|★6M乙酸{hac}mL|甲苯1.2mL|氯仿0.8+2.0mL')
            add_h3(doc, '完整操作步骤'); add_p(doc, HOU.format(aniline_ul=an, aniline_eq=eq, hac_ml=hac))
            add_h3(doc, '预期'); add_p(doc, desc)

    # B-6 Boc
    add_page_break(doc)
    add_h1(doc, '七、B-6 ★ Boc保护策略(JACS 2023 SI精确配比,已读原文+SI)')
    add_callout(doc, '文献依据',
        '贾淑萍等JACS 2023,DOI 10.1021/jacs.3c09284。本地SI文件:ja3c09284_si_001.pdf(61页,已读)。\n'
        'SI Figure S22:Pa-NHBoc 27.8mg+Tp 12.6mg+二氧六环3.4mL+水0.6mL+TFA 60μL\n'
        '→25mL密封釜→120°C×72h→\"a cylindrical organohydrogel with good appearance and elasticity\"\n'
        '本B-6严格按此配比换算到TFPT+Boc-TFMB(醛:胺=0.06:0.09mmol,与原文一致)。', 'FFF8E1')
    add_h3(doc, '设计目的'); add_p(doc, '直接复现JACS 2023的Boc保护策略。预期得到有机水凝胶/果冻状产物。',bold=True)
    add_h3(doc, '试剂'); add_p(doc, 'TFPT 23.6mg|★Boc-TFMB~47mg(0.09mmol)|★1,4-二氧六环3.4mL|★水0.6mL|★TFA 60μL|★不加苯胺/乙酸')
    add_h3(doc, '完整操作步骤'); add_p(doc, BOC)
    add_h3(doc, 'Boc-TFMB制备(实验前准备)')
    add_p(doc, '若不可购买:1.TFMB(320mg,1.0mmol)+DCM(10mL)+Et₃N(0.3mL);2.冰浴0°C,滴加Boc₂O(480mg,2.2mmol)的DCM溶液;3.室温12h,TLC;4.水洗3次,Na₂SO₄干燥,旋蒸(预期产率~85%);5.¹H NMR验证(CDCl₃):δ7.3-7.5(芳环H,6H),δ1.52(Boc-CH₃,18H)')
    add_h3(doc, '预期'); add_p(doc, '★参照JACS 2023:圆柱形有机水凝胶,弹性好。若成功,这是首个TFPT+TFMB体系的Boc保护COF凝胶。')
    add_h3(doc, '失败排查')
    add_bullet(doc, '粉末而非凝胶:Boc脱保护太快或氢键不足。减TFA至30μL或增水至1.0mL。')
    add_bullet(doc, 'Boc-TFMB不溶:换二氧六环/DMF(3:1)或预热至80°C。')
    add_bullet(doc, '无产物:TFA不够。增TFA至90μL,延长至96h。')
    add_bullet(doc, '凝胶太软:增TFPT浓度(+20%)或延长至96h。')

    # 八、汇总
    add_page_break(doc)
    add_h1(doc, '八、六组汇总+时间表')
    add_table(doc, ['编号','方法','TFPT','胺单体','苯胺','酸','核心变量'],
        [['B-1基线','一锅法','23.6mg','TFMB 28.8mg','16.4μL(3eq)','6M×0.4mL HOAc','—'],
         ['B-2★果冻','两步法','23.6mg','TFMB 3+25.8=28.8','16.4μL(3eq)','6M×0.4mL HOAc','★步1加少量胺'],
         ['B-3减苯胺','一锅法','23.6mg','TFMB 28.8mg','★8.2μL(1.5eq)','6M×0.4mL HOAc','★苯胺减半'],
         ['B-4减乙酸','一锅法','23.6mg','TFMB 28.8mg','16.4μL(3eq)','★6M×0.2mL HOAc','★乙酸减量'],
         ['B-5双重减','一锅法','23.6mg','TFMB 28.8mg','★8.2μL(1.5eq)','★6M×0.2mL HOAc','★苯胺+乙酸减'],
         ['B-6★Boc','一锅法','23.6mg','★Boc-TFMB~47mg','—(不加)','★TFA60μL(无HOAc)','★JACS2023SI配比']],
        col_widths=[Cm(1.5),Cm(1.5),Cm(2.0),Cm(2.5),Cm(2.0),Cm(3.0),Cm(2.5)])
    add_p(doc, 'B-6与其他组的关键差异:(1)溶剂不是甲苯/氯仿,是二氧六环/水;(2)不加苯胺/乙酸;(3)72h不是48h;(4)用密封釜不是Pyrex管;(5)不要丙酮/THF洗。',bold=True)

    add_h2(doc, '时间表')
    add_table(doc, ['日期','B-2(两步法)','B-1/B-3/B-4/B-5(一锅法)','B-6(Boc保护)'],
        [['Day1','管清洗+步骤1→120°C','—','—'],
         ['Day2','步骤1完成→冷却+步骤2→120°C','管清洗+投料→120°C','管清洗+投料(二氧六环/水+TFA)→密封→120°C烘箱'],
         ['Day2-4','120°C×48h','120°C×48h','120°C×72h'],
         ['Day4-5','冷却+收集+洗涤+干燥','同B-2','—(B-6还在反应)'],
         ['Day5','—','—','冷却+取出+表征(不洗涤)']],
        col_widths=[Cm(2.0),Cm(4.5),Cm(4.5),Cm(5.0)])

    # 九、参考文献
    add_page_break(doc)
    add_h1(doc, '九、参考文献(含DOI与本地路径)')
    add_p(doc, '【1】侯盛怀等.Angew.Chem.Int.Ed.2025,64,e202421555.DOI:10.1002/anie.202421555.本地:实验\\文章\\侯老师实验\\侯盛怀德国应化.pdf',bold=True)
    add_p(doc,'')
    add_p(doc, '【2】贾淑萍,刘玉洁等.J.Am.Chem.Soc.2023,145,26266-26278.DOI:10.1021/jacs.3c09284.★已读原文+SI(61页).',bold=True)
    add_p(doc,'  SI Figure S22(TpPa gel精确配比):Pa-NHBoc 27.8mg+Tp 12.6mg+dioxane 3.4mL+H₂O 0.6mL+TFA 60μL→120°C×72h→圆柱形有机水凝胶')
    add_p(doc,'')
    add_p(doc, '【3】COF-112编织结构.J.Am.Chem.Soc.2017.DOI:10.1021/jacs.7b07457.Boc+TFA脱保护,编织结构COF晶体制备.')
    add_p(doc,'')
    add_p(doc, '【4】R-62系列(GraphRAG索引).Tp+Boc保护胺,dioxane/water,120°C.本地:tianxuan seek/data/structured/.')
    add_p(doc,'')
    add_p(doc, '【5】v8方案.experiment\\proposals\COF-TFPT-2026-07-15-jelly-v8.docx\\COF-TFPT-TAPT-2026-07-15-D-A1-v8.docx')

    # 附录
    add_h1(doc, '附录:JACS 2023 vs 本方案B-6配比对照')
    add_table(doc, ['参数','JACS 2023 TpPa gel','本方案B-6','是否一致'],
        [['醛','Tp 12.6mg(0.06mmol)','TFPT 23.6mg(0.06mmol)','✓摩尔数一致'],
         ['胺','Pa-NHBoc 27.8mg(0.09mmol)','Boc-TFMB~47mg(0.09mmol)','✓摩尔数一致'],
         ['醛:胺摩尔比','1:1.5','1:1.5','✓'],
         ['有机溶剂','二氧六环3.4mL','二氧六环3.4mL','✓'],
         ['水','0.6mL','0.6mL','✓'],
         ['TFA','60μL','60μL','✓'],
         ['温度','120°C','120°C','✓'],
         ['时间','72h','72h','✓'],
         ['容器','25mL密封釜','25mL密封釜','✓'],
         ['苯胺','无','无','✓'],
         ['乙酸','无','无','✓']],
        col_widths=[Cm(3.0),Cm(4.5),Cm(4.5),Cm(4.0)])

    add_p(doc,''); add_p(doc,'---本方案结束---')
    add_p(doc,'生成工具:bridge/generate_2026_07_13.py|JACS 2023 SI原文已读|配比精确对照')

    doc.save(OUT)
    print(f'✓已生成:{OUT}')
    print(f'  大小:{os.path.getsize(OUT)/1024:.1f}KB')

if __name__=='__main__':
    main()