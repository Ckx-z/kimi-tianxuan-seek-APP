"""提取实验方案 v3.9 与实验 ABCDEF 记录的文本与表格（只读副本）。"""
import json
from pathlib import Path
from docx import Document

BASE = Path(__file__).resolve().parent.parent
REF = BASE / "data" / "experimental_refs"

def dump_docx(path):
    doc = Document(str(path))
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    tables = []
    for ti, t in enumerate(doc.tables):
        rows = []
        for r in t.rows:
            rows.append([c.text.strip() for c in r.cells])
        tables.append(rows)
    return {"paragraphs": paras, "tables": tables}

out = {}
for name in ["实验方案_含氟COF薄膜合成_v3.9_20260626.docx", "实验ABCDEF.docx"]:
    p = REF / name
    out[name] = dump_docx(p)
    print(name, "paras:", len(out[name]["paragraphs"]), "tables:", len(out[name]["tables"]))

with open(BASE / "data" / "experimental_refs" / "_extracted.json", "w", encoding="utf-8") as f:
    json.dump(out, f, ensure_ascii=False, indent=1)
print("saved")
