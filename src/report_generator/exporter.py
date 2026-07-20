"""报告生成器：生成 Word 实验报告。

输入预测结果和推荐条件，填充 Word 模板生成完整实验报告。
MVP 阶段先创建简化报告，后续可复用 main_template.docx 的格式。
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches

from utils.molecule_viz import png_temp_file

PROJECT_ROOT = Path(__file__).resolve().parents[2]
REPORTS_DIR = PROJECT_ROOT / "reports"
REPORTS_DIR.mkdir(parents=True, exist_ok=True)


def _safe_text(value) -> str:
    if value is None:
        return "N/A"
    return str(value)


def _add_structure_images(doc: Document, ald_smiles: str, amine_smiles: str) -> None:
    """在报告中嵌入醛/胺单体 2D 结构图。

    SMILES 解析失败时写降级提示，不影响报告其余部分；
    临时 PNG 文件用后清理。
    """
    entries = [("醛单体", ald_smiles), ("胺单体", amine_smiles)]
    temp_paths: list[Path] = []
    try:
        for label, smiles in entries:
            png = png_temp_file(smiles, prefix="report_mol")
            if png is None:
                doc.add_paragraph(f"{label}结构图：SMILES 无法解析，结构图不可用。")
                continue
            temp_paths.append(png)
            doc.add_paragraph(f"{label}结构图：")
            doc.add_picture(str(png), width=Inches(2.8))
    finally:
        for p in temp_paths:
            try:
                p.unlink(missing_ok=True)
            except OSError:
                pass


def generate_report(
    ald_smiles: str,
    amine_smiles: str,
    prediction: dict,
    conditions: dict,
    output_path: str | Path | None = None,
) -> Path:
    """生成 Word 实验报告。

    Args:
        ald_smiles: 醛单体 SMILES
        amine_smiles: 胺单体 SMILES
        prediction: 预测结果字典（包含 gnn_probability/tree_probability/ensemble_probability）
        conditions: 推荐条件字典
        output_path: 输出路径，默认 reports/报告_时间戳.docx

    Returns:
        生成的报告路径
    """
    if output_path is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = REPORTS_DIR / f"COF实验报告_{timestamp}.docx"
    else:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # 标题
    doc.add_heading("COF 成膜实验方案报告", 0)

    # 1. 单体信息
    doc.add_heading("1. 单体信息", level=1)
    table = doc.add_table(rows=2, cols=2)
    table.style = "Light Grid Accent 1"
    table.rows[0].cells[0].text = "醛单体 SMILES"
    table.rows[0].cells[1].text = ald_smiles
    table.rows[1].cells[0].text = "胺单体 SMILES"
    table.rows[1].cells[1].text = amine_smiles

    # 单体 2D 结构图（解析失败自动降级为文字提示）
    _add_structure_images(doc, ald_smiles, amine_smiles)

    # 2. 成膜概率预测
    doc.add_heading("2. 成膜概率预测", level=1)
    pred_table = doc.add_table(rows=4, cols=2)
    pred_table.style = "Light Grid Accent 1"
    pred_table.rows[0].cells[0].text = "模型"
    pred_table.rows[0].cells[1].text = "成膜概率"
    pred_table.rows[1].cells[0].text = "GNN v5.3"
    pred_table.rows[1].cells[1].text = _safe_text(prediction.get("gnn_probability", "N/A"))
    pred_table.rows[2].cells[0].text = "树模型"
    pred_table.rows[2].cells[1].text = _safe_text(prediction.get("tree_probability", "N/A"))
    pred_table.rows[3].cells[0].text = "综合概率"
    pred_table.rows[3].cells[1].text = _safe_text(prediction.get("ensemble_probability", "N/A"))

    # 3. 推荐实验条件
    doc.add_heading("3. 推荐实验条件", level=1)
    cond_table = doc.add_table(rows=8, cols=2)
    cond_table.style = "Light Grid Accent 1"
    fields = [
        ("合成方法", "method"),
        ("溶剂体系", "solvent_system"),
        ("溶剂比例", "solvent_ratio"),
        ("反应温度", "temperature"),
        ("反应时间", "time"),
        ("催化剂", "catalyst"),
        ("当量比", "stoichiometry"),
        ("备注", "notes"),
    ]
    for i, (label, key) in enumerate(fields):
        cond_table.rows[i].cells[0].text = label
        cond_table.rows[i].cells[1].text = _safe_text(conditions.get(key, "N/A"))

    # 4. 相似历史案例
    doc.add_heading("4. 相似历史案例", level=1)
    doc.add_paragraph(
        f"匹配案例：{conditions.get('case_description', 'N/A')} "
        f"（ID: {conditions.get('matched_case', 'N/A')}）\n"
        f"相似度：{conditions.get('case_similarity_score', 0):.2f}\n"
        f"案例备注：{conditions.get('case_notes', 'N/A')}"
    )

    # 5. 风险提示
    doc.add_heading("5. 风险提示", level=1)
    doc.add_paragraph(
        "1. 本报告由机器学习模型和规则引擎生成，仅供参考。\n"
        "2. 成膜概率受单体纯度、反应条件、操作细节等多因素影响。\n"
        "3. 建议先做小试验证膜连续性，再放大。\n"
        "4. 含 F/CF3 单体疏水性强，需特别注意溶剂溶解度和界面稳定性。"
    )

    # 6. 附录：单体拓扑分类
    doc.add_heading("6. 单体拓扑分类", level=1)
    cls = conditions.get("classification", {})
    doc.add_paragraph(
        f"拓扑类型：{cls.get('topology', 'N/A')}\n"
        f"含氟/CF3：{'是' if cls.get('fluorinated') else '否'}\n"
        f"醛反应位点数：{cls.get('n_ald_sites', 'N/A')}\n"
        f"胺反应位点数：{cls.get('n_amine_sites', 'N/A')}"
    )

    # 保存
    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    pred = {
        "gnn_probability": 0.92,
        "tree_probability": 0.85,
        "ensemble_probability": 0.885,
    }
    cond = {
        "method": "溶剂热法",
        "solvent_system": "均三甲苯 / 二氧六环",
        "solvent_ratio": "1:1",
        "temperature": "120 °C",
        "time": "3 天",
        "catalyst": "乙酸 6.0 M, 0.2 mL",
        "stoichiometry": "1:1",
        "notes": "常规亚胺 COF 合成",
        "matched_case": "HOU_BASE",
        "case_description": "侯老师基准",
        "case_similarity_score": 0.8,
        "case_notes": "经典方法",
        "classification": {"topology": "C3+C2", "fluorinated": False, "n_ald_sites": 3, "n_amine_sites": 2},
    }
    path = generate_report(
        "O=CC1=C(C=O)C(=O)C(C=O)=C1O",
        "Nc1ccc(N)cc1",
        pred,
        cond,
    )
    print(f"报告已生成：{path}")
