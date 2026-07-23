"""方案卡模板系统（P4b-D 后端）。

模板 schema：
    {id, name, source, conditions{}, steps[], checklist[], hints_rules[]}

- 内置模板：data/plan_templates/builtin_hou.json（侯老师界面法 v3.9，入库）
- 用户模板：data/plan_templates/*.json（gitignored，builtin_ 前缀保留给内置）
- extract_template_from_docx：python-docx 读全文 → LLM 提取为 schema JSON
  → 解析校验后返回预览 dict，**不落盘**；确认后由 save_template 落盘。
- LLM 未配置时提取优雅降级为 None 语义（抛 TemplateError 由上层捕获提示），
  list/get/save 不依赖 LLM。
"""

from __future__ import annotations

import json
import logging
import re
from pathlib import Path

logger = logging.getLogger(__name__)

try:
    from src import runtime_config
except ImportError:
    import runtime_config  # type: ignore

PROJECT_ROOT = runtime_config.resource_root()
TEMPLATES_DIR = PROJECT_ROOT / "data" / "plan_templates"
BUILTIN_ID = "builtin_hou_v3_9"

_REQUIRED_KEYS = ("id", "name", "source", "conditions", "steps", "checklist", "hints_rules")


class TemplateError(Exception):
    """模板相关可读错误（LLM 未配置、docx 解析失败、schema 校验失败等）。"""


# ---------------------------------------------------------------- 校验

def validate_template(tpl: dict) -> dict:
    """校验并规范化模板 dict；不合法抛 TemplateError。"""
    if not isinstance(tpl, dict):
        raise TemplateError("模板必须是 JSON 对象")
    missing = [k for k in _REQUIRED_KEYS if k not in tpl]
    if missing:
        raise TemplateError(f"模板缺少字段: {', '.join(missing)}")
    out = {k: tpl[k] for k in _REQUIRED_KEYS}
    for k in ("id", "name", "source"):
        if not isinstance(out[k], str) or not out[k].strip():
            raise TemplateError(f"模板字段 {k} 必须是非空字符串")
        out[k] = out[k].strip()
    if not isinstance(out["conditions"], dict):
        raise TemplateError("conditions 必须是对象")
    if not isinstance(out["steps"], list) or not all(isinstance(s, str) for s in out["steps"]):
        raise TemplateError("steps 必须是字符串数组")
    for c in out["checklist"]:
        if not isinstance(c, dict) or "item" not in c:
            raise TemplateError("checklist 项必须含 item 字段")
    for r in out["hints_rules"]:
        if not isinstance(r, dict) or "hint" not in r:
            raise TemplateError("hints_rules 项必须含 hint 字段")
    return out


# ---------------------------------------------------------------- CRUD

def _template_path(tpl_id: str) -> Path:
    safe = re.sub(r"[^A-Za-z0-9_\-]", "_", tpl_id)
    return TEMPLATES_DIR / f"{safe}.json"


def list_templates() -> list[dict]:
    """列出全部模板（内置 + data/plan_templates/ 下用户模板）。"""
    templates: list[dict] = []
    if not TEMPLATES_DIR.exists():
        return templates
    for p in sorted(TEMPLATES_DIR.glob("*.json")):
        try:
            tpl = json.loads(p.read_text(encoding="utf-8"))
            validate_template(tpl)
            tpl["builtin"] = tpl["id"].startswith("builtin_")
            templates.append(tpl)
        except Exception as exc:
            logger.warning("跳过非法模板 %s: %s", p.name, exc)
    return templates


def get_template(tpl_id: str) -> dict:
    """按 id 取模板；不存在抛 TemplateError。"""
    p = _template_path(tpl_id)
    if not p.exists():
        raise TemplateError(f"模板不存在: {tpl_id}")
    try:
        tpl = json.loads(p.read_text(encoding="utf-8"))
        return validate_template(tpl)
    except TemplateError:
        raise
    except Exception as exc:
        raise TemplateError(f"模板读取失败: {tpl_id}: {exc}")


def save_template(tpl: dict) -> dict:
    """校验后落盘 data/plan_templates/<id>.json，返回规范化模板。"""
    tpl = validate_template(tpl)
    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
    _template_path(tpl["id"]).write_text(
        json.dumps(tpl, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return tpl


# ---------------------------------------------------------------- docx 提取

def _read_docx_text(file_path: str | Path) -> str:
    try:
        import docx
    except Exception as exc:
        raise TemplateError(f"python-docx 不可用: {exc}")
    try:
        doc = docx.Document(str(file_path))
    except Exception as exc:
        raise TemplateError(f"docx 打开失败: {exc}")
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts).strip()
    if not text:
        raise TemplateError("docx 中没有可提取的文本")
    return text


def _extract_prompt(text: str, name: str) -> list[dict]:
    user = (
        "你是化学实验方案结构化助手。下面是一份文献/实验方案全文，"
        "请提取为 JSON 模板，字段严格为：\n"
        '{"name": 模板名称, "source": 来源说明, "conditions": {条件键值对}, '
        '"steps": [操作步骤...], "checklist": [{"item":..., "detail":...}...], '
        '"hints_rules": [{"rule":..., "match":..., "hint":...}...]}\n'
        "要求：steps 保持原文操作顺序；checklist 提炼易错点；hints_rules 提炼"
        "针对特定单体结构（如含氟、大芳环）的调整建议；不要编造原文没有的条件数值；"
        "只输出 JSON 本体，不要 markdown 代码围栏，不要多余文字。\n"
        f"模板名称建议：{name or '（请根据内容命名）'}\n"
        "---- 方案全文 ----\n"
        f"{text}"
    )
    return [
        {"role": "system", "content": "你是严谨的化学实验方案结构化助手，只输出合法 JSON。"},
        {"role": "user", "content": user},
    ]


def _parse_llm_json(raw: str) -> dict:
    text = (raw or "").strip()
    # 容忍模型裹了 ```json 围栏
    m = re.search(r"```(?:json)?\s*(.*?)```", text, re.S)
    if m:
        text = m.group(1).strip()
    try:
        return json.loads(text)
    except Exception as exc:
        raise TemplateError(f"LLM 输出不是合法 JSON: {exc}")


def extract_template_from_docx(file_path: str | Path, name: str = "") -> dict:
    """docx → LLM → 模板预览 dict（含新 id），**不落盘**。

    LLM 未配置时抛 TemplateError（上层据此提示"未配置 LLM"）。
    """
    try:
        from src.llm import client as llm_client
        configured = llm_client.is_configured()
    except Exception:
        configured = False
    if not configured:
        raise TemplateError("未配置 LLM，无法从 docx 自动提取模板（请在设置页配置后重试）")

    text = _read_docx_text(file_path)
    raw = llm_client.chat_completion(
        _extract_prompt(text, name), max_tokens=4000, temperature=0.1
    )
    if not raw:
        raise TemplateError("LLM 提取失败（空响应），可重试或手动编辑模板")
    data = _parse_llm_json(raw)
    data.setdefault("name", name or Path(file_path).stem)
    data.setdefault("source", f"用户上传提取自 {Path(file_path).name}")
    data["id"] = "user_" + re.sub(r"[^A-Za-z0-9_\-]", "_", Path(file_path).stem)
    return validate_template(data)
