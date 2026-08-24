"""科研助手附件：上传落盘 + 文档文本提取 + 图片 vision data URL。

落盘（runtime_config.user_data_root()/assistant/uploads/，frozen 时自动落
%APPDATA%/COF-Film-Recommend/data/assistant/uploads/，绝不写安装目录）：
    <upload_id>.<ext>   文件本体
    <upload_id>.json    元信息 {upload_id, filename, ext, kind, size, created_at}

类型白名单：
- 图片：.png .jpg .jpeg .webp（→ OpenAI vision base64 data URL）
- 文档：.txt .md .json .csv .docx .pdf（→ 提取文本注入用户消息上下文）

单文件 ≤10MB。docx 用 python-docx；pdf 优先 pypdf，缺失时回退 PyMuPDF(fitz)；
所有解析失败都抛 AttachmentError（可读中文错误），由路由层转 400。
"""

from __future__ import annotations

import base64
import json
import logging
import re
import uuid
from datetime import datetime
from pathlib import Path

logger = logging.getLogger(__name__)

try:
    from src import runtime_config
except ImportError:  # pragma: no cover
    import runtime_config  # type: ignore

UPLOADS_DIR = runtime_config.user_data_root() / "assistant" / "uploads"

IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp"}
DOC_EXTS = {".txt", ".md", ".json", ".csv", ".docx", ".pdf"}
MAX_BYTES = 10 * 1024 * 1024          # 单文件 10MB
MAX_TEXT_CHARS = 8000                 # 注入 LLM 上下文的文本上限
MAX_ATTACHMENTS_PER_MESSAGE = 3       # 单条消息附件数上限

_ID_RE = re.compile(r"^u_[0-9a-f]{12}$")

_MIME = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".webp": "image/webp",
}


class AttachmentError(Exception):
    """附件相关可读错误（类型/大小超限、解析失败、不存在等）。"""


# ---------------------------------------------------------------- 基础工具

def _now() -> str:
    return datetime.now().astimezone().isoformat(timespec="seconds")


def _new_id() -> str:
    return f"u_{uuid.uuid4().hex[:12]}"


def _meta_path(upload_id: str) -> Path:
    return UPLOADS_DIR / f"{upload_id}.json"


def file_path_of(meta: dict) -> Path:
    """附件文件本体路径（由元信息推导，不信任外部传入路径）。"""
    return UPLOADS_DIR / f"{meta['upload_id']}{meta['ext']}"


def kind_of_ext(ext: str) -> str | None:
    """扩展名 → kind（image/document）；不支持返回 None。"""
    ext = (ext or "").lower()
    if ext in IMAGE_EXTS:
        return "image"
    if ext in DOC_EXTS:
        return "document"
    return None


# ---------------------------------------------------------------- 上传 / 读取 / 删除

def save_upload(filename: str, data: bytes) -> dict:
    """校验并落盘一个附件，返回元信息 dict。非法输入抛 AttachmentError。"""
    filename = (filename or "").strip() or "attachment"
    ext = Path(filename).suffix.lower()
    kind = kind_of_ext(ext)
    if kind is None:
        raise AttachmentError(
            f"不支持的附件类型「{ext or '无扩展名'}」："
            "图片仅支持 png/jpg/jpeg/webp，文档仅支持 txt/md/json/csv/docx/pdf")
    if not isinstance(data, (bytes, bytearray)) or len(data) == 0:
        raise AttachmentError("附件内容为空")
    if len(data) > MAX_BYTES:
        raise AttachmentError(
            f"附件「{filename}」超过大小限制（{MAX_BYTES // 1024 // 1024}MB）")

    meta = {
        "upload_id": _new_id(),
        "filename": filename,
        "ext": ext,
        "kind": kind,
        "size": len(data),
        "created_at": _now(),
    }
    UPLOADS_DIR.mkdir(parents=True, exist_ok=True)
    file_path_of(meta).write_bytes(bytes(data))
    _meta_path(meta["upload_id"]).write_text(
        json.dumps(meta, ensure_ascii=False, indent=1) + "\n", encoding="utf-8")
    return meta


def get_meta(upload_id: str) -> dict | None:
    """按 upload_id 取元信息；不存在 / id 非法 / 文件本体缺失返回 None。"""
    if not isinstance(upload_id, str) or not _ID_RE.match(upload_id):
        return None
    path = _meta_path(upload_id)
    if not path.is_file():
        return None
    try:
        meta = json.loads(path.read_text(encoding="utf-8"))
    except Exception as exc:
        logger.warning("附件元信息读取失败 %s: %s", path.name, exc)
        return None
    if not isinstance(meta, dict) or meta.get("upload_id") != upload_id:
        return None
    if not file_path_of(meta).is_file():
        return None
    return meta


def delete_upload(upload_id: str) -> bool:
    """删除附件（元信息 + 本体）；成功 True，不存在 False。"""
    meta = get_meta(upload_id)
    if meta is None:
        return False
    _meta_path(upload_id).unlink(missing_ok=True)
    file_path_of(meta).unlink(missing_ok=True)
    return True


# ---------------------------------------------------------------- 文本提取（文档类）

def _extract_docx(path: Path) -> str:
    try:
        import docx
    except Exception as exc:
        raise AttachmentError(f"python-docx 不可用: {exc}")
    try:
        doc = docx.Document(str(path))
    except Exception as exc:
        raise AttachmentError(f"docx 打开失败: {exc}")
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def _extract_pdf(path: Path) -> str:
    """pdf 文本提取：优先 pypdf，缺失时回退 PyMuPDF(fitz)。"""
    try:
        from pypdf import PdfReader
        try:
            reader = PdfReader(str(path))
            return "\n".join(
                (page.extract_text() or "") for page in reader.pages).strip()
        except Exception as exc:
            raise AttachmentError(f"pdf 解析失败: {exc}")
    except ImportError:
        pass
    try:
        import fitz  # PyMuPDF
    except Exception as exc:
        raise AttachmentError(
            f"pdf 解析库不可用（pypdf 与 PyMuPDF 均未安装）: {exc}")
    try:
        with fitz.open(str(path)) as doc:
            return "\n".join(page.get_text() for page in doc).strip()
    except Exception as exc:
        raise AttachmentError(f"pdf 解析失败: {exc}")


def extract_text(meta: dict) -> str:
    """提取文档类附件文本（截断到 MAX_TEXT_CHARS）；图片/未知类型返回 ""。

    解析失败抛 AttachmentError；提取结果为空抛 AttachmentError（提示用户）。
    """
    ext = str(meta.get("ext") or "").lower()
    if kind_of_ext(ext) != "document":
        return ""
    path = file_path_of(meta)
    if ext in (".txt", ".md", ".json", ".csv"):
        try:
            text = path.read_bytes().decode("utf-8", errors="replace").strip()
        except Exception as exc:
            raise AttachmentError(f"附件读取失败: {exc}")
    elif ext == ".docx":
        text = _extract_docx(path)
    elif ext == ".pdf":
        text = _extract_pdf(path)
    else:  # pragma: no cover - 白名单已拦截
        raise AttachmentError(f"不支持的文档类型: {ext}")
    if not text:
        raise AttachmentError(f"附件「{meta.get('filename')}」中没有可提取的文本")
    if len(text) > MAX_TEXT_CHARS:
        text = text[:MAX_TEXT_CHARS] + "\n……（附件过长，已截断）"
    return text


# ---------------------------------------------------------------- 图片（vision）

def image_data_url(meta: dict) -> str:
    """图片附件 → base64 data URL（OpenAI vision 格式用）。非图片抛异常。"""
    ext = str(meta.get("ext") or "").lower()
    if kind_of_ext(ext) != "image":
        raise AttachmentError(f"附件「{meta.get('filename')}」不是图片")
    raw = file_path_of(meta).read_bytes()
    mime = _MIME.get(ext, "image/png")
    return f"data:{mime};base64,{base64.b64encode(raw).decode('ascii')}"
