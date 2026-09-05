"""深度研究循环（v1.6.0 P1）：plan → execute → critic → report。

事件契约（POST /api/assistant/research 的 SSE，data 行 JSON）：
- {"type": "plan", "steps": [{"title","query","note"}], "summary"}     研究计划
- {"type": "step_start", "index", "title"}                            子问题开始
- {"type": "tool_call"/"tool_result", ...}                            检索过程（复用 loop 契约）
- {"type": "step_done", "index", "title", "summary"}                  子问题小结
- {"type": "critic_note", "text"}                                     批判/引用校验说明
- {"type": "token", "text"}                                           报告流式
- {"type": "report", "report_id", "title"}                            报告已落盘
- {"type": "done"} / {"type": "error", "message"}

执行约束：研究模式只用只读工具白名单（检索/查询类），写操作工具
（确认机制）不参与；web_search 未配置时工具表自动缺席，检索自然降级
学术/本地源。引用纪律：报告引用的每条文献必须来自本轮工具结果，
写报告前先跑引用核验（loop._emit_verified 复用），违规打回重写一轮。

报告落盘：user_data_root()/research/report_<id>.json（Markdown + 引用清单），
支持 docx 导出（md → 段落/标题，参考文献带 DOI）。
"""

from __future__ import annotations

import json
import logging
import re
import uuid
from datetime import datetime
from pathlib import Path
from typing import Iterator

try:
    from src import runtime_config
except ImportError:  # pragma: no cover
    import runtime_config  # type: ignore

from . import llm_bridge, registry, verify
from .loop import _sanitize_for_plain, _stream_text

logger = logging.getLogger(__name__)

REPORTS_DIR = runtime_config.user_data_root() / "research"

MAX_STEPS = 5          # 计划子问题上限
STEP_TOOL_BUDGET = 3   # 每子问题检索轮上限
CRITIC_GAP_RETRY = 1   # 批判轮缺口补搜次数上限
_REPORT_ID_RE = re.compile(r"^rpt_[0-9a-f]{12}$")
# v1.7.0：会话综合报告 sessrpt_<sid12>_v<n>（一个对话一份，版本递增）
_SESSION_REPORT_RE = re.compile(r"^sessrpt_[0-9a-f]{12}_v\d+$")
_MAX_EVIDENCE_CHARS = 6000  # 写报告时每类证据注入上限
_MAX_DIGEST_CHARS = 15000   # 会话报告：对话摘要注入上限（近期消息优先）

# 研究模式工具白名单（只读，无写操作确认流）
RESEARCH_TOOLS = (
    "query_graphrag", "read_experiment_records", "predict_film",
    "get_monomer_props", "cas_resolve", "lookup_paper_doi",
    "web_search", "academic_search", "fetch_page",
    "list_favorites", "list_prediction_history", "get_daily_brief",
)

_PLAN_PROMPT = (
    "你是科研深度研究助手。把用户的问题拆解为 2–5 个可独立检索的子问题，"
    "每个子问题给出检索关键词（英文优先）。只输出一个 JSON 对象（不要 "
    "markdown 围栏、不要其他文字）：\n"
    '{"summary": "一句话研究目标", "steps": [{"title": "子问题标题", '
    '"query": "检索关键词", "note": "这一步想弄清什么"}]}'
)

_PLAN_RETRY = "格式错误。请严格只输出 JSON：{\"summary\": ..., \"steps\": [...]}。"

_EXEC_PROMPT = (
    "你正在执行研究计划的第 {index} 步：{title}（目标：{note}）。\n"
    "每一轮只输出一个 JSON 对象（不要 markdown 围栏）：\n"
    "- 需要检索证据：{{\"tool\": \"工具名\", \"args\": {{...}}}}\n"
    "- 证据已足够、可以小结：{{\"done\": \"本步小结（含关键发现，"
    "引用时给出 CAS/DOI/URL）\"}}\n\n"
    "可用工具：\n{tools}\n\n"
    "工具结果会以用户消息回填。优先学术源（academic_search）与本地"
    "证据（query_graphrag / read_experiment_records）；每步最多 "
    f"{STEP_TOOL_BUDGET} 轮。"
)

_EXEC_RETRY = "格式错误。请严格只输出一个 JSON：{\"tool\":...} 或 {\"done\": ...}。"

_CRITIC_PROMPT = (
    "你是研究批判员。对照研究计划的每一步，检查已有证据是否足以支撑结论，"
    "找出关键缺口（缺文献年份/DOI、缺最新进展、缺本地实验对比等）。"
    "只输出 JSON（不要其他文字）：\n"
    '{"ok": true/false, "gaps": [{"step_index": 从1开始, "query": "补搜关键词"}]}'
)

_WRITER_PROMPT = (
    "你是科研报告撰写员。基于【证据清单】写一份 Markdown 研究报告：\n"
    "1. 标题（# 开头）+ 一句话摘要；\n"
    "2. 按研究计划的每个子问题分节（## 标题），正文只引用证据清单里的内容；\n"
    "3. 「结论与下一步」一节给出可操作建议；\n"
    "4. 「参考文献」一节：编号列出证据清单中真实用到的文献，每条格式为\n"
    "   `[n] 标题. 期刊, 年份. DOI: 10.xxxx/...（https://doi.org/...）`；\n"
    "   没有 DOI 的给 URL。\n"
    "引用纪律（红线）：证据清单里没有的文献/CAS/数字一律不得写入；"
    "证据不足的结论要注明「证据有限」。直接输出 Markdown 正文。"
)

_REPORT_FIELDS = ("report_id", "question", "title", "created_at",
                  "markdown", "refs", "allow_web")


def _now() -> str:
    return datetime.now().astimezone().isoformat(timespec="seconds")


def _new_id() -> str:
    return f"rpt_{uuid.uuid4().hex[:12]}"


def _parse_json(text: str) -> dict | None:
    """容错解析 LLM 输出的 JSON 对象（去围栏 + 截取首尾大括号）。"""
    if not text:
        return None
    s = text.strip()
    if s.startswith("```"):
        lines = [l for l in s.splitlines() if not l.strip().startswith("```")]
        s = "\n".join(lines).strip()
    i, j = s.find("{"), s.rfind("}")
    if i < 0 or j <= i:
        return None
    try:
        obj = json.loads(s[i:j + 1])
    except Exception:
        return None
    return obj if isinstance(obj, dict) else None


def _tool_description(available: set[str]) -> str:
    """白名单工具的自然语言描述（路径 B 提示词用）。"""
    lines = []
    for name, t in registry.TOOLS.items():
        if name not in available:
            continue
        fn = t["schema"]["function"]
        props = fn["parameters"].get("properties", {})
        required = set(fn["parameters"].get("required", []))
        args_desc = ", ".join(
            f"{k}{'（必填）' if k in required else '（可选）'}" for k in props)
        lines.append(f"- {name}({args_desc})：{fn['description']}")
    return "\n".join(lines) or "（无可用工具）"


def available_research_tools() -> set[str]:
    """研究模式可用工具 = 白名单 ∩ 当前环境裁剪后的工具表。"""
    from src.llm import client as _llm_client  # noqa: PLC0415
    ok, _ = _llm_client.web_search_available()
    tools = set(RESEARCH_TOOLS)
    if not ok:
        tools.discard("web_search")
    return tools


# ---------------------------------------------------------------- 计划

def plan_steps(question: str) -> dict:
    """LLM 拆解研究计划；失败返回单步兜底计划（不阻塞研究）。"""
    try:
        text = llm_bridge.chat_text(
            [{"role": "system", "content": _PLAN_PROMPT},
             {"role": "user", "content": question}],
            max_tokens=1200)
        obj = _parse_json(text or "")
        if not obj and text:
            text2 = llm_bridge.chat_text(
                [{"role": "system", "content": _PLAN_PROMPT},
                 {"role": "user", "content": question},
                 {"role": "assistant", "content": text},
                 {"role": "user", "content": _PLAN_RETRY}],
                max_tokens=1200)
            obj = _parse_json(text2 or "")
    except Exception as exc:  # LLM 失败 → 兜底计划
        logger.warning("研究计划生成失败，用兜底单步计划: %s", exc)
        obj = None
    steps = []
    if isinstance(obj, dict) and isinstance(obj.get("steps"), list):
        for s in obj["steps"]:
            if not isinstance(s, dict):
                continue
            title = str(s.get("title") or "").strip()
            query = str(s.get("query") or title or "").strip()
            if not title:
                continue
            steps.append({"title": title[:80], "query": query,
                          "note": str(s.get("note") or "").strip()[:200]})
    if not steps:
        steps = [{"title": "整体调研", "query": question, "note": ""}]
    return {
        "summary": str((obj or {}).get("summary") or "").strip()[:200]
                   if isinstance(obj, dict) else "",
        "steps": steps[:MAX_STEPS],
    }


# ---------------------------------------------------------------- 步骤执行

def _execute_step(index: int, step: dict, tools: set[str]
                  ) -> Iterator[dict]:
    """单步检索：LLM 每轮一个指令（tool/done），结果回填，预算内收尾。"""
    step_no = f"第 {index + 1} 步"
    yield {"type": "step_start", "index": index,
           "title": step.get("title", "")}
    desc = _tool_description(tools)
    work = [
        {"role": "system", "content": _EXEC_PROMPT.format(
            index=index + 1, title=step.get("title", ""),
            note=step.get("note") or "", tools=desc)},
        {"role": "user", "content": f"检索关键词：{step.get('query', '')}"},
    ]
    retried = False
    for _ in range(STEP_TOOL_BUDGET):
        text = llm_bridge.chat_text(work)
        if text is None:
            yield {"type": "step_done", "index": index,
                   "title": step.get("title", ""),
                   "summary": "（检索失败：LLM 无响应）"}
            return
        directive = _parse_json(text)
        if directive is None:
            if not retried:
                retried = True
                work += [{"role": "assistant", "content": text},
                         {"role": "user", "content": _EXEC_RETRY}]
                continue
            yield {"type": "step_done", "index": index,
                   "title": step.get("title", ""),
                   "summary": str(text)[:400]}
            return
        done = directive.get("done")
        if done is not None:
            yield {"type": "step_done", "index": index,
                   "title": step.get("title", ""),
                   "summary": str(done)[:400]}
            return
        name = str(directive.get("tool") or "").strip()
        args = directive.get("args") if isinstance(
            directive.get("args"), dict) else {}
        if name not in tools:
            result = {"text": f"研究模式不允许使用工具 {name}（可用："
                              f"{'、'.join(sorted(tools))}）",
                      "details": {}, "is_error": True}
        else:
            result = registry.execute(name, args)
        yield {"type": "tool_call", "name": name, "args": args}
        yield {"type": "tool_result", "name": name,
               "summary": registry.summary_of(result),
               "is_error": bool(result.get("is_error"))}
        yield {"type": "_result", "result": result}  # 内部事件：主循环收集
        work += [{"role": "assistant", "content": text},
                 {"role": "user", "content":
                  f"工具 {name} 返回：\n{result['text']}\n\n请继续。"
                  f"（若证据已足够，输出 {{\"done\": \"小结\"}}）"}]
    yield {"type": "step_done", "index": index,
           "title": step.get("title", ""),
           "summary": "（检索轮次用尽）"}


# ---------------------------------------------------------------- 批判与补搜

def _critic_gaps(plan: dict) -> list[dict]:
    """批判轮：检查证据覆盖度，返回需要补搜的缺口列表。"""
    try:
        text = llm_bridge.chat_text(
            [{"role": "system", "content": _CRITIC_PROMPT},
             {"role": "user",
              "content": "研究计划：" + json.dumps(
                  plan, ensure_ascii=False, default=str)}],
            max_tokens=800)
        obj = _parse_json(text or "")
    except Exception as exc:
        logger.warning("批判轮失败: %s", exc)
        return []
    if not isinstance(obj, dict):
        return []
    gaps = obj.get("gaps")
    if not isinstance(gaps, list):
        return []
    out = []
    for g in gaps:
        if isinstance(g, dict) and str(g.get("query") or "").strip():
            out.append({"step_index": g.get("step_index"),
                        "query": str(g["query"]).strip()[:200]})
    return out


# ---------------------------------------------------------------- 报告落盘

def _report_path(report_id: str) -> Path:
    return REPORTS_DIR / f"{report_id}.json"


def save_report(report_id: str, question: str, title: str,
                markdown: str, refs: list[dict],
                allow_web: bool = True, session_id: str | None = None,
                version: int | None = None) -> dict:
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    rec = {
        "report_id": report_id,
        "question": question,
        "title": title,
        "created_at": _now(),
        "markdown": markdown,
        "refs": refs,
        "allow_web": bool(allow_web),
        "kind": "session" if version is not None else "question",
        "session_id": session_id,
    }
    if version is not None:
        rec["version"] = version
    _report_path(report_id).write_text(
        json.dumps(rec, ensure_ascii=False, indent=2), encoding="utf-8")
    return rec


def load_report(report_id: str) -> dict | None:
    if not (_REPORT_ID_RE.match(report_id or "")
            or _SESSION_REPORT_RE.match(report_id or "")):
        return None
    p = _report_path(report_id)
    if not p.is_file():
        return None
    try:
        data = json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        return None
    return data if isinstance(data, dict) else None


def list_reports() -> list[dict]:
    out = []
    if not REPORTS_DIR.is_dir():
        return out
    for p in REPORTS_DIR.glob("*.json"):
        stem = p.stem
        if not (_REPORT_ID_RE.match(stem) or _SESSION_REPORT_RE.match(stem)):
            continue
        data = load_report(stem)
        if data:
            out.append({
                "report_id": data["report_id"],
                "title": data["title"],
                "created_at": data["created_at"],
                "ref_count": len(data.get("refs") or []),
                "kind": data.get("kind") or "question",
                "session_id": data.get("session_id"),
                "version": data.get("version"),
            })
    out.sort(key=lambda r: r["created_at"], reverse=True)
    return out


def delete_report(report_id: str) -> bool:
    if not (_REPORT_ID_RE.match(report_id or "")
            or _SESSION_REPORT_RE.match(report_id or "")):
        return False
    p = _report_path(report_id)
    if not p.is_file():
        return False
    p.unlink()
    return True


def _evidence_block(results: list[dict]) -> str:
    """证据清单（写报告注入用）：每条工具结果限长。"""
    lines = []
    for r in results or []:
        if not isinstance(r, dict):
            continue
        text = str(r.get("text") or "")[:_MAX_EVIDENCE_CHARS]
        lines.append(f"- {text}")
    return "\n".join(lines) or "（无证据）"


def _collect_report_refs(results: list[dict]) -> list[dict]:
    """从工具结果 details 提取结构化引用（文献/网页），去重。"""
    refs: list[dict] = []
    seen: set[str] = set()
    for r in results or []:
        details = r.get("details") if isinstance(r, dict) else None
        if not isinstance(details, dict):
            continue
        for item in (details.get("papers") or details.get("results") or []):
            if not isinstance(item, dict):
                continue
            url = item.get("url") or ""
            doi = item.get("doi") or ""
            key = doi or url
            if not key or key in seen:
                continue
            seen.add(key)
            refs.append({
                "title": item.get("title") or "",
                "doi": doi,
                "url": url,
                "source": item.get("source") or "",
            })
    return refs[:20]


def _title_from_markdown(markdown: str, question: str) -> str:
    m = re.search(r"^#\s+(.+)$", markdown or "", re.MULTILINE)
    if m:
        return m.group(1).strip()[:80]
    return (question or "").strip()[:40] or "研究报告"


def report_to_docx(report: dict) -> bytes:
    """Markdown 报告 → docx 字节（标题/段落/参考文献，DOI 原样保留）。"""
    try:
        from docx import Document
    except ImportError:  # pragma: no cover
        raise RuntimeError("python-docx 未安装，无法导出 Word")
    doc = Document()
    doc.add_heading(report.get("title") or "研究报告", level=0)
    doc.add_paragraph(f"生成时间：{report.get('created_at') or ''}")
    doc.add_paragraph(f"研究问题：{report.get('question') or ''}")
    lines = (report.get("markdown") or "").splitlines()
    for line in lines:
        s = line.strip()
        if not s:
            continue
        if s.startswith("### "):
            doc.add_heading(s[4:], level=3)
        elif s.startswith("## "):
            doc.add_heading(s[3:], level=2)
        elif s.startswith("# "):
            doc.add_heading(s[2:], level=1)
        elif re.match(r"^[-*] ", s):
            doc.add_paragraph(s[2:], style="List Bullet")
        elif re.match(r"^\d+\. ", s):
            doc.add_paragraph(re.sub(r"^\d+\. ", "", s), style="List Number")
        else:
            doc.add_paragraph(s)
    if report.get("refs"):
        doc.add_heading("参考文献", level=1)
        for i, ref in enumerate(report["refs"], 1):
            doi = ref.get("doi") or ""
            url = ref.get("url") or ""
            link = f"https://doi.org/{doi}" if doi else url
            doc.add_paragraph(
                f"[{i}] {ref.get('title') or '（无标题）'}"
                + (f"，DOI: {doi}" if doi else f"，URL: {url}"),
                style="List Number")
    from io import BytesIO
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------- 主循环

def run_research(question: str, allow_web: bool = True,
                 session_id: str | None = None) -> Iterator[dict]:
    """深度研究主入口（生成器，SSE 事件流）。LLM 未配置先抛 error 事件。

    session_id（v1.7.0）：报告落盘时关联会话；会话综合报告（一对话一报告）
    会把该报告作为来源并入。不传则为独立单问报告（历史行为）。
    """
    question = (question or "").strip()
    if not question:
        yield {"type": "error", "message": "question 不能为空"}
        return
    if not llm_bridge.is_configured():
        yield {"type": "error",
               "message": "LLM 未配置：请到设置页填写 base_url / api_key / "
                          "model 后再使用深度研究。"}
        return
    tools = available_research_tools()

    # 1. 计划
    plan = plan_steps(question)
    yield {"type": "plan",
           "steps": [{"title": s["title"], "note": s["note"]}
                     for s in plan["steps"]],
           "summary": plan.get("summary") or ""}

    # 2. 逐子问题执行（_result 内部事件收集完整工具结果供引用核验）
    all_results: list[dict] = []
    for i, step in enumerate(plan["steps"]):
        for ev in _execute_step(i, step, tools):
            if ev.get("type") == "_result":
                all_results.append(ev["result"])
                continue
            yield ev

    # 3. 批判补搜（缺口 ≤1 轮，最多补 2 个缺口）
    gaps = _critic_gaps(plan)
    for g in gaps[:2]:
        q = g.get("query") or ""
        if not q:
            continue
        yield {"type": "critic_note", "text": f"发现证据缺口，补搜：{q}"}
        name = "web_search" if "web_search" in tools else "academic_search"
        result = registry.execute(name, {"query": q})
        all_results.append(result)
        yield {"type": "tool_call", "name": name, "args": {"query": q}}
        yield {"type": "tool_result", "name": name,
               "summary": registry.summary_of(result),
               "is_error": bool(result.get("is_error"))}

    # 4. 撰写报告（含引用核验打回）
    work = [
        {"role": "system", "content": _WRITER_PROMPT},
        {"role": "user",
         "content": f"研究问题：{question}\n\n"
                    f"研究计划：{json.dumps(plan, ensure_ascii=False, default=str)}\n\n"
                    f"【证据清单】\n{_evidence_block(all_results)}"},
    ]
    text = llm_bridge.chat_text(work, max_tokens=8000)
    if text is None:
        yield {"type": "error", "message": "报告生成失败（LLM 无响应），"
                                           "请重试。"}
        return
    report_md = text.strip()
    yield from loop_verified_stream(work, report_md, all_results)

    # 5. 落盘
    title = _title_from_markdown(report_md, question)
    refs = _collect_report_refs(all_results)
    rid = _new_id()
    save_report(rid, question, title, report_md, refs, allow_web=allow_web,
                session_id=session_id)
    if session_id:
        # 会话内深度研究：报告计入会话（供一对话一报告合成），并留一条
        # 助手消息记录。失败不影响研究主流程。
        try:
            from . import sessions as sessions_module
            sessions_module.append_message(
                session_id, "assistant",
                f"深度研究完成：《{title}》（报告 ID: {rid}），"
                "已可作为本会话综合报告的来源。")
        except Exception as exc:  # pragma: no cover
            logger.warning("研究完成消息落会话失败（已跳过）: %s", exc)
    yield {"type": "report", "report_id": rid, "title": title,
           "session_id": session_id}
    yield {"type": "done"}


def loop_verified_stream(work: list[dict], text: str,
                         results: list[dict]) -> Iterator[dict]:
    """引用核验 + 流式（复用 loop._emit_verified，事件兼容）。"""
    from .loop import _emit_verified
    yield from _emit_verified(work, text, results)


# ---------------------------------------------------------------- 会话综合报告（v1.7.0）
# 「一个对话 → 一份综合报告」：整合会话内全部问答 + 深度研究产出，生成/
# 增量更新（版本递增，追加不覆盖）。

_SESSION_WRITER_PROMPT = (
    "你是科研助手。请把下面这段对话的完整内容整合成一份结构化的深度研究"
    "报告（Markdown，中文）。结构固定为：\n"
    "# <简洁标题>\n"
    "## 研究背景\n## 核心发现\n## 详细分析\n## 结论与建议\n## 参考文献\n"
    "## 附录：对话时间线\n"
    "要求：\n"
    "1. 忠实整合对话内容，不要编造对话中没有的事实；\n"
    "2. 只引用对话/工具结果中真实出现过的 DOI、URL 或 CAS；对话中确实"
    "没有可引用文献时，参考文献一节如实写「本次对话未产生外部引用」；\n"
    "3. 附录按时间顺序列用户问题与结论摘要（一问一答一行）。"
)

_UPDATE_SUFFIX = (
    "\n\n以上是第 {v} 版报告的原文。请在此基础上**增量更新**：保留既有章节"
    "与结论，把对话中新增的进展并入相应章节（不标注版本号）；参考文献取"
    "并集；附录时间线补充新增问答。输出完整的新版报告。"
)


def _session_report_id(session_id: str, version: int) -> str:
    """会话报告 ID：sessrpt_<sid 12hex>_v<n>（sid 即 sess_<12hex> 后缀）。"""
    suffix = (session_id or "").split("_", 1)[-1][:12]
    return f"sessrpt_{suffix}_v{int(version)}"


def _conversation_digest(session: dict) -> str:
    """对话摘要：近期消息优先，总量受 _MAX_DIGEST_CHARS 约束。"""
    msgs = list(session.get("messages") or [])
    parts: list[str] = []
    total = 0
    for m in reversed(msgs):
        role = "用户" if m.get("role") == "user" else "助手"
        content = str(m.get("content") or "").strip()
        block = f"{role}：{content[:_MAX_EVIDENCE_CHARS]}\n"
        for te in (m.get("tool_events") or []):
            if not isinstance(te, dict):
                continue
            if te.get("type") == "tool_call":
                args = json.dumps(te.get("args"), ensure_ascii=False,
                                  default=str)[:200]
                block += f"[调用工具 {te.get('name')}：{args}]\n"
            elif te.get("type") == "tool_result":
                summary = str(te.get("summary") or "")[:500]
                block += f"[工具结果 {te.get('name')}：{summary}]\n"
        if total + len(block) > _MAX_DIGEST_CHARS:
            block = block[:_MAX_DIGEST_CHARS - total]
            if block:
                parts.append(block)
            break
        parts.append(block)
        total += len(block)
    parts.reverse()
    return "\n".join(parts) or "（空对话）"


def _conversation_ref_results(session: dict, sub_reports: list[dict]) -> list[dict]:
    """构造引用核验的「工具结果」：对话文本 + 关联单问研究产出。

    verify.collect_refs 扫描 {text} 字段提取 CAS/DOI/URL，因此把消息正文与
    tool_result 摘要拼进 text 即可复用同一套核验口径。
    """
    results: list[dict] = []
    for m in (session.get("messages") or []):
        blob = str(m.get("content") or "")
        for te in (m.get("tool_events") or []):
            if isinstance(te, dict) and te.get("type") == "tool_result":
                blob += "\n" + str(te.get("summary") or "")
        if blob.strip():
            results.append({"text": blob, "details": {}})
    for rep in sub_reports:
        results.append({"text": rep.get("markdown") or "", "details": {}})
    return results


def _collect_refs_from_results(results: list[dict]) -> list[dict]:
    """伪工具结果 → 结构化引用清单（与 _collect_report_refs 同口径）。"""
    refs: list[dict] = []
    seen: set[str] = set()
    from . import verify as verify_module
    pooled = verify_module.collect_refs(results)
    for doi in sorted(pooled.get("doi") or []):
        if doi in seen:
            continue
        seen.add(doi)
        refs.append({"title": "", "doi": doi, "url": "", "source": ""})
    for url in sorted(pooled.get("url") or []):
        if url in seen:
            continue
        seen.add(url)
        refs.append({"title": "", "doi": "", "url": url, "source": ""})
    return refs[:20]


def _question_reports_of(session_id: str) -> list[dict]:
    """该会话关联的单问深度研究报告（kind=question，按时间升序）。"""
    if not REPORTS_DIR.is_dir():
        return []
    out = []
    for p in REPORTS_DIR.glob("rpt_*.json"):
        data = load_report(p.stem)
        if data and data.get("session_id") == session_id:
            out.append(data)
    out.sort(key=lambda r: r.get("created_at") or "")
    return out


def build_session_report(session: dict, hint: str | None = None) -> Iterator[dict]:
    """会话综合报告（生成/增量更新，SSE 事件流）。

    事件：token（流式）/ critic_note（引用核验修正）/
    report（{report_id, title, version, session_id}）/ done / error。
    """
    from . import sessions as sessions_module
    sid = session["session_id"]
    if not llm_bridge.is_configured():
        yield {"type": "error",
               "message": "LLM 未配置：请到设置页填写 base_url / api_key / "
                          "model 后再生成研究报告。"}
        return

    sub_reports = _question_reports_of(sid)
    prev = session.get("report") if isinstance(session.get("report"), dict) else None
    prev_data = load_report(prev["report_id"]) if prev else None
    base_version = int((prev or {}).get("version") or 0)

    digest = _conversation_digest(session)
    for i, rep in enumerate(sub_reports, 1):
        digest += (f"\n\n【深度研究产出 {i}：{rep.get('title') or ''}】\n"
                   + (rep.get("markdown") or "")[:_MAX_EVIDENCE_CHARS])

    user_content = (f"会话标题：{session.get('title')}\n\n"
                    f"【对话内容】\n{digest}")
    if hint:
        user_content += f"\n\n用户额外要求：{hint.strip()[:500]}"
    work = [{"role": "system", "content": _SESSION_WRITER_PROMPT}]
    if prev_data:
        user_content += _UPDATE_SUFFIX.format(v=base_version)
        work.append({"role": "user",
                     "content": prev_data.get("markdown") or ""})
    work.append({"role": "user", "content": user_content})

    text = llm_bridge.chat_text(work, max_tokens=8000)
    if text is None:
        yield {"type": "error", "message": "报告生成失败（LLM 无响应），请重试。"}
        return
    report_md = text.strip()

    # 引用核验：对话 + 既有报告中的引用均为「已核实」来源
    ref_results = _conversation_ref_results(session, sub_reports)
    if prev_data:
        ref_results.append({"text": prev_data.get("markdown") or "",
                            "details": {}})
    yield from loop_verified_stream(work, report_md, ref_results)

    version = base_version + 1
    report_id = _session_report_id(sid, version)
    title = _title_from_markdown(report_md, session.get("title") or "会话报告")
    refs = _collect_refs_from_results(ref_results)
    save_report(report_id, question=(hint or "会话综合报告"), title=title,
                markdown=report_md, refs=refs, allow_web=True,
                session_id=sid, version=version)
    pointer = {"report_id": report_id, "version": version, "updated_at": _now()}
    try:
        sessions_module.update_meta(sid, report=pointer)
    except Exception as exc:  # pragma: no cover
        logger.warning("会话报告指针更新失败（报告已落盘）: %s", exc)
    yield {"type": "report", "report_id": report_id, "title": title,
           "version": version, "session_id": sid}
    yield {"type": "done"}
