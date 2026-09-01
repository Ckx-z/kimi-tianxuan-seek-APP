"""实验记录 Word（.docx）导出。

build_record_docx(rec, version="") -> bytes：完整实验记录报告；
export_filename(rec) -> str：中文文件名（路由层走 RFC 5987 filename*）。

排版约定（与项目既有 docx 经验一致）：
- python-docx 默认样式，标题层级 0/1 清晰；
- 表格一律 'Table Grid'（带边框）；
- 中文字体宋体：Normal/标题样式 eastAsia + 正文 run 显式设置双保险。

降级路径（不抛异常，保证任何记录都能导出）：
- 无关联收藏 / 收藏被删 / 无打分快照 → 模型打分小节写「未打分」；
- LLM 未配置 / 性质卡生成失败 / RDKit 不可用 → 单体性质小节保留占位，
  写「（未配置 LLM，本节略）」。
"""

from __future__ import annotations

import io
import logging
import re
from datetime import datetime

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

logger = logging.getLogger(__name__)

# 与前端 components/records/meta.ts CONDITION_LABELS 一致的九键中文名
_CONDITION_LABELS = {
    "solvent_1": "溶剂一",
    "solvent_2": "溶剂二",
    "eluent": "洗脱剂",
    "modulator": "调制剂",
    "catalyst": "催化剂",
    "temperature_c": "温度（℃）",
    "time_days": "时间（天）",
    "vessel": "容器",
    "addition_order": "加料顺序",
}

_OUTCOME_LABELS = {
    "film": "成膜",
    "partial": "部分成膜",
    "failed": "失败",
    "": "未定",
}

_STATUS_LABELS = {"draft": "草稿", "final": "正式"}

_OOD_LABELS = {
    "none": "适用域内（无警告）",
    "warning": "警告：接近模型适用域边界，打分可信度降低",
    "out": "超出适用域（模型不适用）",
}

# RDKit 事实键中文名（src/recommend/monomer_props.py _FACT_KEYS）
_FACT_LABELS = {
    "mw": "分子量",
    "xlogp": "XlogP",
    "tpsa": "TPSA",
    "hbd": "氢键给体数",
    "hba": "氢键受体数",
    "aromatic_rings": "芳香环数",
    "f_count": "氟原子数",
    "rotatable_bonds": "可旋转键数",
}

_CN_FONT = "宋体"


# ---------------------------------------------------------------------------
# 字体 / 段落工具
# ---------------------------------------------------------------------------


def _set_style_cn_fonts(doc: Document) -> None:
    """Normal 与标题样式设置宋体（ascii 名 + eastAsia），表格单元格继承生效。"""
    for name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3",
                 "Heading 4"):
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        style.font.name = _CN_FONT
        rpr = style.element.get_or_add_rPr()
        rpr.get_or_add_rFonts().set(qn("w:eastAsia"), _CN_FONT)


def _apply_layout_spec(doc: Document) -> None:
    """v1.5.0 版式规范：
    - 中文字体宋体 / 西文 Times New Roman；正文小四 12pt、1.5 倍行距；
    - 段前/段后 0.5 行（Pt 6）；页面边距上下 2.54cm、左右 3.17cm；
    - 一级标题黑体加粗小四、二级标题宋体加粗小四；文档标题黑体 16pt 加粗。
    """
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)

    try:
        title = doc.styles["Title"]
        title.font.name = "Times New Roman"
        title.font.size = Pt(16)
        title.font.bold = True
        title.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "黑体")
    except KeyError:
        pass
    for style_name, cn_font in (("Heading 1", "黑体"), ("Heading 2", "宋体")):
        try:
            st = doc.styles[style_name]
        except KeyError:
            continue
        st.font.name = "Times New Roman"
        st.font.size = Pt(12)
        st.font.bold = True
        st.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), cn_font)

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)


def _set_cell_border(cell, **edges) -> None:
    """给单元格设置边框（edges: {top|bottom|left|right: (sz, val)}）。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, (sz, val) in edges.items():
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), val)
        if val != "none":
            el.set(qn("w:sz"), str(sz))
            el.set(qn("w:color"), "000000")


def _three_line_table(doc: Document, rows: list[list[str]],
                      header: bool = True) -> None:
    """三线表：仅顶线/底线（1.5pt）+ 表头下线（0.75pt）；表头加粗居中。"""
    n_cols = len(rows[0]) if rows else 2
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = str(val)
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                _cn_run(run)
            is_head = header and i == 0
            if is_head:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.bold = True
            _set_cell_border(
                cell,
                top=(12, "single") if (is_head or not header) else (0, "none"),
                bottom=(12, "single") if i == len(rows) - 1
                       else ((6, "single") if is_head else (0, "none")))
    return table


def _cn_run(run) -> None:
    """run 级中文字体双保险（正文段落显式指定宋体 + eastAsia）。"""
    run.font.name = _CN_FONT
    rpr = run._element.get_or_add_rPr()
    rpr.get_or_add_rFonts().set(qn("w:eastAsia"), _CN_FONT)


def _add_text(doc: Document, text: str) -> None:
    """多行纯文本 → 每行一个段落（python-docx 不渲染 run 内 \\n）。"""
    lines = str(text).splitlines() or [""]
    for line in lines:
        p = doc.add_paragraph()
        _cn_run(p.add_run(line))


def _kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    """两列「字段 | 值」三线表（表头=字段列加粗）。"""
    if not rows:
        return
    _three_line_table(doc, [[str(k), str(v)] for k, v in rows], header=False)


def _add_structure_image(doc: Document, smiles: str, caption: str,
                         fig_no: list[int], width_cm: float = 6.0) -> bool:
    """RDKit 绘制 2D 结构图并嵌入（居中 + 图号图注，图注五号宋体）。

    返回是否成功（SMILES 不可解析时返回 False，调用方写占位说明）。
    """
    if not smiles:
        return False
    try:
        from rdkit import Chem
        from rdkit.Chem import Draw
    except Exception:
        return False
    mol = Chem.MolFromSmiles(smiles)
    if mol is None:
        return False
    try:
        img = Draw.MolToImage(mol, size=(700, 320))
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)
    except Exception:
        return False
    fig_no[0] += 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(buf, width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(f"图{fig_no[0]}　{caption}")
    cr.font.size = Pt(10.5)
    _cn_run(cr)
    return True


def _format_ref_gbt7714(ref: dict) -> str:
    """GB/T 7714 统一编号制（[J] 期刊）。缺字段降级为「题名 + DOI」。"""
    ref = ref if isinstance(ref, dict) else {}
    authors = ref.get("authors") or []
    title = str(ref.get("title") or "").strip()
    journal = str(ref.get("journal") or "").strip()
    year = str(ref.get("year") or "").strip()
    volume = str(ref.get("volume") or "").strip()
    issue = str(ref.get("issue") or "").strip()
    pages = str(ref.get("pages") or "").strip()
    doi = str(ref.get("doi") or "").strip()
    if authors and title and journal and year:
        a = ", ".join(str(x) for x in authors[:3])
        if len(authors) > 3:
            a += "，等"
        vol_issue = volume + (f"({issue})" if issue else "")
        pages_part = f": {pages}" if pages else ""
        return f"{a}. {title}[J]. {journal}, {year}, {vol_issue}{pages_part}."
    if title and doi:
        return f"{title}. DOI: {doi}."
    if title:
        return f"{title}."
    return str(ref.get("paper_id") or "（未命名文献）")


def _favorite_data(rec: dict) -> dict | None:
    """读取记录关联收藏（只读 + 容错），供结构图/参考文献/DFT 快照复用。"""
    fav_id = rec.get("favorite_id")
    if not fav_id:
        return None
    try:
        from favorites import store as favorites_store
        fav = favorites_store.get_favorite(fav_id)
        return fav if isinstance(fav, dict) else None
    except Exception:
        return None


# ---------------------------------------------------------------------------
# 数据装配（均为只读 + 容错）
# ---------------------------------------------------------------------------


def _prediction_snapshot(rec: dict) -> dict | None:
    """模型打分快照：优先取关联收藏的 latest_prediction（最新），
    收藏缺失/删除时回落到记录创建时冗余的 prediction_snapshot。
    """
    snap = None
    fav_id = rec.get("favorite_id")
    if fav_id:
        try:
            from favorites import store as favorites_store

            fav = favorites_store.get_favorite(fav_id)
            if isinstance(fav, dict):
                cand = fav.get("latest_prediction")
                if isinstance(cand, dict) and cand.get("score") is not None:
                    snap = cand
        except Exception as exc:
            logger.warning("导出时读取收藏 %s 打分快照失败: %s", fav_id, exc)
    if snap is None:
        cand = rec.get("prediction_snapshot")
        if isinstance(cand, dict) and cand.get("score") is not None:
            snap = cand
    return snap


def _monomer_props_section(doc: Document, label: str, monomer: dict,
                           level: int = 2) -> None:
    """单体性质小节：RDKit 事实行 + LLM 中文解读；任何失败降级为占位说明。"""
    monomer = monomer if isinstance(monomer, dict) else {}
    smiles = str(monomer.get("smiles") or "").strip()
    name = str(monomer.get("name") or "").strip()
    doc.add_heading(f"{label}性质解读", level=level)
    if not smiles:
        doc.add_paragraph("（未填写 SMILES，本节略）")
        return
    try:
        from recommend.monomer_props import get_monomer_properties

        card = get_monomer_properties(smiles, name=name)
    except Exception as exc:
        logger.warning("导出时性质卡生成失败（%s）: %s", smiles, exc)
        card = None
    if not isinstance(card, dict):
        doc.add_paragraph("（未配置 LLM，本节略）")
        return
    facts = card.get("facts") or {}
    if facts:
        facts_text = "；".join(
            f"{_FACT_LABELS.get(k, k)}：{facts[k]}" for k in _FACT_LABELS if k in facts
        )
        doc.add_paragraph(f"RDKit 计算事实：{facts_text}")
    narrative = card.get("narrative")
    if isinstance(narrative, str) and narrative.strip():
        _add_text(doc, narrative.strip())
    else:
        doc.add_paragraph("（未配置 LLM，本节略）")


def _sorted_timeline(timeline: list) -> list[dict]:
    """时间线条目按时间排序：time_label 可解析为日期时间的按时间升序，
    不可解析的保持原顺序排在其后（time_label 为自由文本，如「第 2 天」）。
    """
    def _key(item: tuple[int, dict]):
        idx, entry = item
        label = str(entry.get("time_label") or "").strip()
        try:
            return (0, datetime.fromisoformat(label).timestamp(), idx)
        except ValueError:
            return (1, 0.0, idx)

    entries = [e for e in (timeline or []) if isinstance(e, dict)]
    return [e for _, e in sorted(enumerate(entries), key=_key)]


# ---------------------------------------------------------------------------
# 主入口
# ---------------------------------------------------------------------------


def build_record_docx(rec: dict, version: str = "") -> bytes:
    """生成单条实验记录 Word 报告，返回 docx 字节串。"""
    doc = Document()
    _set_style_cn_fonts(doc)
    _apply_layout_spec(doc)

    doc.add_heading("实验记录", 0)
    _build_record_body(doc, rec, base_level=1)

    # 页脚：导出方 + 版本号
    footer_text = "由 COF 科研助手导出"
    if version:
        footer_text += f" · 版本 v{version}"
    for section in doc.sections:
        p = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        _cn_run(p.add_run(footer_text))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_record_body(doc: Document, rec: dict, base_level: int = 1) -> None:
    """向 doc 追加一条实验记录的完整正文（不含标题与页脚）。

    单条导出与汇总导出共用此函数；base_level 为各小节标题层级
    （单条导出用 1，汇总导出在组标题/记录标题之下用 3）。
    """
    rec = rec if isinstance(rec, dict) else {}
    fig_no = [0]  # 图号计数器（结构图与附件图共享，按出现顺序编号）
    record_id = str(rec.get("record_id") or "")
    experiment_no = str(rec.get("experiment_no") or "").strip()
    status = _STATUS_LABELS.get(str(rec.get("status") or ""), str(rec.get("status") or ""))
    today = datetime.now().astimezone().date().isoformat()

    # 基本信息
    doc.add_heading("基本信息", level=base_level)
    _kv_table(doc, [
        ("实验编号", experiment_no or "（未填写）"),
        ("记录 ID", record_id),
        ("状态", status or "—"),
        ("记录日期", str(rec.get("date") or "—")),
        ("实验结果", _OUTCOME_LABELS.get(str(rec.get("outcome") or ""), "未定")),
        ("机械强度", str(rec.get("strength") or "—")),
        ("操作人", str(rec.get("operator") or "—")),
        ("导出日期", today),
    ])
    notes = str(rec.get("notes") or "").strip()
    if notes:
        p = doc.add_paragraph()
        _cn_run(p.add_run(f"备注：{notes}"))

    # 单体信息表（三线表）
    doc.add_heading("单体信息", level=base_level)
    ald = rec.get("aldehyde") if isinstance(rec.get("aldehyde"), dict) else {}
    amine = rec.get("amine") if isinstance(rec.get("amine"), dict) else {}
    mono_rows = [["", "醛单体", "胺单体"]]
    for field, label in (("name", "名称"), ("smiles", "SMILES"), ("cas", "CAS")):
        mono_rows.append([label, str(ald.get(field) or "—"),
                          str(amine.get(field) or "—")])
    _three_line_table(doc, mono_rows)

    # 结构图（v1.5.0）：单体 2D 结构式 + 二聚体（标注 X 与结合能口径）
    doc.add_heading("化学结构", level=base_level)
    fav = _favorite_data(rec)
    got_ald = _add_structure_image(
        doc, str(ald.get("smiles") or ""), "醛单体结构式", fig_no)
    got_amine = _add_structure_image(
        doc, str(amine.get("smiles") or ""), "胺单体结构式", fig_no)
    dimer_smiles = ""
    dimer_note = ""
    if isinstance(fav, dict):
        dft_entries = fav.get("dft_entries") or []
        snap = fav.get("dft_snapshot")
        if isinstance(dft_entries, list) and dft_entries:
            latest = dft_entries[-1] if isinstance(dft_entries[-1], dict) else None
        else:
            latest = snap if isinstance(snap, dict) else None
        if isinstance(latest, dict):
            dimer_smiles = str(latest.get("dimer_smiles") or "")
            x_desc = str(latest.get("x_description") or "—")
            e_bind = latest.get("e_bind_kj")
            e_text = f"{float(e_bind):.1f} kJ/mol" if isinstance(e_bind, (int, float)) else "—"
            dimer_note = (f"缩合二聚体（计算对象：二聚体与「{x_desc}」的结合能 "
                          f"{e_text}；吸附位点见 DFT 结果的 3D 结合构象与片段着色）")
    if dimer_smiles:
        _add_structure_image(doc, dimer_smiles,
                             dimer_note or "缩合二聚体结构式", fig_no)
    if not (got_ald or got_amine or dimer_smiles):
        doc.add_paragraph("（未记录单体 SMILES，结构图略）")

    # 模型打分（关联收藏 latest_prediction / 记录快照 / 未打分）
    doc.add_heading("模型打分", level=base_level)
    snap = _prediction_snapshot(rec)
    if snap is None:
        doc.add_paragraph("未打分（无关联收藏的打分快照）。")
    else:
        rows = [("综合评分（成膜倾向）", f"{float(snap['score']):.3f}")]
        if snap.get("std") is not None:
            rows.append(("不确定度（±std）", f"± {float(snap['std']):.3f}"))
        ood = str(snap.get("ood") or "")
        rows.append(("OOD 等级", _OOD_LABELS.get(ood, ood or "—")))
        if snap.get("tree_score") is not None:
            rows.append(("树模型分量", f"{float(snap['tree_score']):.3f}"))
        if snap.get("gnn_score") is not None:
            rows.append(("GNN 分量", f"{float(snap['gnn_score']):.3f}"))
        if snap.get("score_policy"):
            rows.append(("打分口径", str(snap["score_policy"])))
        if snap.get("arm"):
            rows.append(("打分臂", str(snap["arm"])))
        if snap.get("date"):
            rows.append(("打分时间", str(snap["date"])))
        _kv_table(doc, rows)

    # 单体性质（LLM 解读；未配置/失败保留小节占位）
    doc.add_heading("单体性质", level=base_level)
    _monomer_props_section(doc, "醛单体", ald, level=base_level + 1)
    _monomer_props_section(doc, "胺单体", amine, level=base_level + 1)

    # 实验条件（所有已填字段；九键用中文名，额外键原样保留）
    doc.add_heading("实验条件", level=base_level)
    cond = rec.get("conditions") if isinstance(rec.get("conditions"), dict) else {}
    filled = [
        (_CONDITION_LABELS.get(str(k), str(k)), str(v))
        for k, v in cond.items()
        if v not in (None, "")
    ]
    if filled:
        _kv_table(doc, filled)
    else:
        doc.add_paragraph("（未填写实验条件）")

    # 完整实验流程（全文）
    doc.add_heading("完整实验流程", level=base_level)
    process_notes = str(rec.get("process_notes") or "").strip()
    if process_notes:
        _add_text(doc, process_notes)
    else:
        doc.add_paragraph("（未填写）")

    # 时间线（按时间排序：时间 | 过程记录 | 附件名；三线表）
    doc.add_heading("时间线", level=base_level)
    timeline = _sorted_timeline(rec.get("timeline") or [])
    if timeline:
        rows = [["时间", "过程记录", "附件名"]]
        for entry in timeline:
            atts = [
                str(m.get("filename") or m.get("attachment_id") or "")
                for m in entry.get("attachments") or []
                if isinstance(m, dict)
            ]
            rows.append([str(entry.get("time_label") or "—"),
                         str(entry.get("description") or "—"),
                         "；".join(a for a in atts if a) or "—"])
        _three_line_table(doc, rows)
    else:
        doc.add_paragraph("（无时间点记录）")

    # 附件图片（v1.5.0）：时间线全部图片附件按时间序嵌入（居中 + 图注）
    image_atts: list[tuple[str, dict]] = []
    for entry in timeline:
        for m in entry.get("attachments") or []:
            if isinstance(m, dict) and (m.get("ext") or "").lower() in (
                    ".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"):
                image_atts.append((str(entry.get("time_label") or "—"), m))
    if image_atts:
        doc.add_heading("实验记录图片", level=base_level)
        try:
            from records import store as records_store
        except Exception:
            records_store = None
        embedded = 0
        for time_label, meta in image_atts:
            path = None
            if records_store is not None:
                try:
                    found = records_store.get_attachment_path(
                        record_id, str(meta.get("attachment_id") or ""))
                    path = found[0] if found else None
                except Exception:
                    path = None
            if not path or not path.is_file():
                continue
            try:
                fig_no[0] += 1
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(str(path), width=Cm(12))
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cr = cap.add_run(
                    f"图{fig_no[0]}　{meta.get('filename') or '附件图片'}（时间点：{time_label}）")
                cr.font.size = Pt(10.5)
                _cn_run(cr)
                embedded += 1
            except Exception:
                continue
        if not embedded:
            doc.add_paragraph("（附件图片缺失，无法嵌入）")

    # 参考文献（v1.5.0：GB/T 7714 统一编号制；无全文元数据时降级为题名+DOI）
    refs: list[dict] = []
    if isinstance(fav, dict):
        refs = [r for r in fav.get("references") or [] if isinstance(r, dict)]
    if refs:
        doc.add_heading("参考文献", level=base_level)
        for i, ref in enumerate(refs, start=1):
            full = dict(ref)
            doi = str(full.get("doi") or "").strip()
            if doi and not full.get("journal"):
                try:
                    from literature import resolver as lit_resolver
                    _, record = lit_resolver.find_by_doi(doi)
                    if isinstance(record, dict):
                        for key in ("authors", "journal", "year", "volume",
                                    "issue", "pages", "title"):
                            if record.get(key) and not full.get(key):
                                full[key] = record[key]
                except Exception:
                    pass
            p = doc.add_paragraph()
            run = p.add_run(f"[{i}] {_format_ref_gbt7714(full)}")
            _cn_run(run)

    # 自我总结 / 我认为的失误（有则）
    self_summary = str(rec.get("self_summary") or "").strip()
    if self_summary:
        doc.add_heading("自我总结", level=base_level)
        _add_text(doc, self_summary)
    mistakes = str(rec.get("mistakes") or "").strip()
    if mistakes:
        doc.add_heading("我认为的失误", level=base_level)
        _add_text(doc, mistakes)


def _group_title(fav: dict) -> str:
    """汇总导出的组标题：单体组名称（醛 + 胺），名称为空回落到 SMILES。"""
    fav = fav if isinstance(fav, dict) else {}
    ald = fav.get("aldehyde") if isinstance(fav.get("aldehyde"), dict) else {}
    amine = fav.get("amine") if isinstance(fav.get("amine"), dict) else {}
    ald_name = str(ald.get("name") or "").strip() or str(ald.get("smiles") or "未知醛单体")
    amine_name = str(amine.get("name") or "").strip() or str(amine.get("smiles") or "未知胺单体")
    return f"{ald_name} + {amine_name}"


def build_bundle_docx(groups: list[dict], version: str = "") -> bytes:
    """按收藏分组导出多组实验记录为一份 Word，返回 docx 字节串。

    groups: [{"favorite": fav_dict, "records": [rec_dict, ...]}, ...]
    （records 由调用方按时间序提供）。封面含标题/导出时间/软件版本/目录；
    每组一级标题为单体组名称（醛+胺），组内每条记录二级标题后接与单条
    导出一致的正文（记录间分页符分隔）；收藏无记录时标注「暂无实验记录」。
    """
    doc = Document()
    _set_style_cn_fonts(doc)
    _apply_layout_spec(doc)
    now = datetime.now().astimezone()

    doc.add_heading("实验记录汇总导出", 0)
    p = doc.add_paragraph()
    _cn_run(p.add_run(f"导出时间：{now.isoformat(timespec='seconds')}"))
    if version:
        p = doc.add_paragraph()
        _cn_run(p.add_run(f"软件版本：v{version}"))
    p = doc.add_paragraph()
    _cn_run(p.add_run(f"收藏分组数：{len(groups)}"))

    # 静态目录（v1.5.0）：按组列出（不含页码；Word 打开后如需页码可插入自动目录）
    doc.add_heading("目录", level=1)
    for gi, group in enumerate(groups, start=1):
        group = group if isinstance(group, dict) else {}
        fav = group.get("favorite") if isinstance(group.get("favorite"), dict) else {}
        records = [r for r in (group.get("records") or []) if isinstance(r, dict)]
        tp = doc.add_paragraph()
        _cn_run(tp.add_run(f"{gi}. {_group_title(fav)}（{len(records)} 条记录）"))

    for group in groups:
        group = group if isinstance(group, dict) else {}
        fav = group.get("favorite") if isinstance(group.get("favorite"), dict) else {}
        records = [r for r in (group.get("records") or []) if isinstance(r, dict)]
        doc.add_page_break()
        doc.add_heading(_group_title(fav), level=1)
        if not records:
            doc.add_paragraph("暂无实验记录")
            continue
        for idx, rec in enumerate(records):
            if idx > 0:
                # 每份记录之间用分页符分隔（版式一致、独立成页）
                doc.add_page_break()
            no = str(rec.get("experiment_no") or "").strip()
            rec_id = str(rec.get("record_id") or "")
            label = f"实验记录 {no}" if no else f"实验记录 {rec_id}"
            doc.add_heading(label, level=2)
            _build_record_body(doc, rec, base_level=3)

    # 页脚：导出方 + 版本号
    footer_text = "由 COF 科研助手导出"
    if version:
        footer_text += f" · 版本 v{version}"
    for section in doc.sections:
        p = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        _cn_run(p.add_run(footer_text))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def bundle_export_filename(now: datetime | None = None) -> str:
    """汇总导出文件名：实验记录汇总_YYYYMMDD_HHMM.docx。"""
    now = now or datetime.now()
    return f"实验记录汇总_{now.strftime('%Y%m%d_%H%M')}.docx"


_FILENAME_BAD_CHARS = re.compile(r'[\\/:*?"<>|\r\n]+')


def export_filename(rec: dict) -> str:
    """中文下载文件名：实验记录_<编号>_<记录ID>.docx（编号含非法字符时清洗）。"""
    rec = rec if isinstance(rec, dict) else {}
    record_id = str(rec.get("record_id") or "record")
    no = _FILENAME_BAD_CHARS.sub("_", str(rec.get("experiment_no") or "").strip())
    stem = f"实验记录_{no}_{record_id}" if no else f"实验记录_{record_id}"
    return f"{stem}.docx"
